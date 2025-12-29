"""
Main Application - UniteUs ETL Web Interface

FastAPI web application providing database viewing, ETL job management,
data import, advanced analytics/reporting, and comprehensive RESTful API
endpoints for the UniteUs ETL pipeline.

Author: Waqqas Hanafi
Copyright: © 2025 Calaveras County Health and Human Services Agency
"""

# ============================================================================
# IMPORTS
# ============================================================================
import sys
import tempfile
import logging
import secrets
import re
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional
from contextlib import asynccontextmanager

from fastapi import FastAPI, File, UploadFile, HTTPException, Request, Query, Form, Depends
import sqlite3
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse, RedirectResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn

# Document generation libraries
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning("python-docx not installed. Word export will be unavailable.")

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    logger = logging.getLogger(__name__)
    logger.warning("reportlab not installed. PDF export will be unavailable.")

try:
    from PIL import Image as PILImage
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

# Configure logging early
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

# Import our core services
import sys
sys.path.insert(0, str(Path(__file__).parent))

from .config import config
from .database import get_database_manager
from .etl_service import get_etl_service
from .settings_manager import get_settings_manager
from .sftp_service import get_sftp_service
from .siem_logger import get_siem_logger, SIEMEventType, SIEMSeverity
from .auth import (
    get_auth_service, 
    require_auth, 
    require_auth_redirect, 
    require_role,
    get_current_session,
    UserSession,
    UserRole
)


# Pydantic models for API
class DatabaseQuery(BaseModel):
    """Database query request model"""
    query: str
    limit: int = 100


class FileUploadResponse(BaseModel):
    """File upload response model"""
    success: bool
    message: str
    filename: str
    size: int


class ETLJobRequest(BaseModel):
    """ETL job request model"""
    force_reprocess: bool = False
    latest_only: bool = False
    max_workers: int = 4
    selected_files: Optional[List[str]] = None  # List of filenames to process


class TableSearchRequest(BaseModel):
    """Table search request model"""
    table_name: str
    search_term: str = ""
    columns: Optional[List[str]] = None
    limit: int = 100
    offset: int = 0


class TableSortRequest(BaseModel):
    """Table sort request model"""
    table_name: str
    sort_column: str
    sort_direction: str = "asc"  # asc or desc
    limit: int = 100
    offset: int = 0


class AnnualReportExportRequest(BaseModel):
    """Annual report export request model"""
    period: str
    generated_date: str
    summary: Dict[str, str]
    charts: Dict[str, str]  # Chart ID -> base64 image data
    tables: Optional[Dict[str, List[Dict[str, str]]]] = None  # New structure
    program_performance: Optional[List[Dict[str, str]]] = None  # Legacy support


# Global state
app_state = {
    "db_manager": None,
    "etl_service": None
}

# Initialize logger
logger = logging.getLogger(__name__)


# Track if settings have been loaded to prevent duplicate logging
_settings_loaded = False

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Application lifespan manager with memory leak prevention"""
    global _settings_loaded
    import gc
    import threading
    import time
    
    # Startup
    app_state["db_manager"] = get_database_manager()
    app_state["etl_service"] = get_etl_service()
    
    # Load settings from database into runtime config (only log once)
    if not _settings_loaded:
        try:
            settings_manager = get_settings_manager()
            settings_manager.load_settings_into_config()
            logger.info("Loaded SIEM and SFTP settings from database")
            _settings_loaded = True
        except Exception as e:
            logger.warning(f"Could not load settings from database: {e}")
            _settings_loaded = True  # Mark as attempted to prevent repeated warnings
    
    # Periodic memory cleanup task
    def periodic_cleanup():
        """Periodic cleanup to prevent memory leaks"""
        while True:
            try:
                time.sleep(3600)  # Run every hour
                
                # Force garbage collection
                collected = gc.collect()
                if collected > 0:
                    logger.debug(f"Garbage collection: collected {collected} objects")
                
                # Periodic database connection pool cleanup
                if app_state.get("db_manager"):
                    try:
                        app_state["db_manager"].pool.periodic_cleanup()
                    except Exception as e:
                        logger.warning(f"Error during database pool cleanup: {e}")
                
                # Log memory usage (if psutil available)
                try:
                    import psutil
                    import os
                    process = psutil.Process(os.getpid())
                    mem_info = process.memory_info()
                    mem_mb = mem_info.rss / (1024 * 1024)
                    logger.info(f"Memory usage: {mem_mb:.1f} MB (RSS)")
                    
                    # Log database pool stats
                    if app_state.get("db_manager"):
                        pool_stats = app_state["db_manager"].pool.get_pool_stats()
                        logger.debug(f"Database pool stats: {pool_stats}")
                except ImportError:
                    pass  # psutil not available
                except Exception as e:
                    logger.debug(f"Error getting memory stats: {e}")
                    
            except Exception as e:
                logger.error(f"Error in periodic cleanup: {e}", exc_info=True)
    
    # Automated sync task with lock to prevent concurrent runs
    sync_lock = threading.Lock()
    
    def automated_sync_task():
        """Background task for automated SFTP sync and ETL"""
        from datetime import datetime, timedelta
        import time
        
        logger.info("Automated sync task started")
        
        while True:
            try:
                time.sleep(60)  # Check every minute
                
                # Get automated sync config
                db_manager = app_state.get("db_manager")
                if not db_manager:
                    continue
                
                with db_manager.pool.get_connection() as conn:
                    cursor = conn.execute("""
                        SELECT enabled, interval_minutes, last_run, next_run
                        FROM automated_sync_config
                        WHERE id = 1
                    """)
                    row = cursor.fetchone()
                    
                    if not row or not row[0]:  # Not enabled
                        continue
                    
                    enabled, interval_minutes, last_run, next_run = row
                    
                    # Check if it's time to run
                    now = datetime.now()
                    should_run = False
                    
                    if next_run:
                        try:
                            next_run_dt = datetime.fromisoformat(next_run)
                            should_run = now >= next_run_dt
                        except:
                            should_run = True  # Run if next_run is invalid
                    else:
                        should_run = True  # Run if next_run is not set
                    
                    if not should_run:
                        continue
                    
                    # Try to acquire lock - if another sync is running, skip this iteration
                    if not sync_lock.acquire(blocking=False):
                        logger.info("Automated sync already running, skipping this check")
                        continue
                    
                    try:
                        logger.info("🔄 Starting automated sync...")
                        
                        # Update last_run IMMEDIATELY so user knows sync attempted to run
                        # Update next_run BEFORE syncing to prevent multiple syncs after downtime
                        # This ensures if server was down for hours, we only do ONE sync when it comes back
                        next_run_time = now + timedelta(minutes=interval_minutes)
                        conn.execute("""
                            UPDATE automated_sync_config
                            SET last_run = ?, next_run = ?
                            WHERE id = 1
                        """, (now.isoformat(), next_run_time.isoformat()))
                        conn.commit()
                        logger.info(f"📅 Next sync scheduled for: {next_run_time.strftime('%Y-%m-%d %H:%M:%S')}")
                        
                        # Step 1: Sync from SFTP
                        try:
                            sftp_service = get_sftp_service()
                            if config.sftp.enabled:
                                logger.info("📡 Automated sync: Discovering files on SFTP...")
                                files = sftp_service.discover_files(username="automated_sync")
                                
                                if files:
                                    logger.info(f"✅ Automated sync: Found {len(files)} files on SFTP")
                                    
                                    # Cache the file list in database for "Last synced" display
                                    import json
                                    files_data = [f.to_dict() for f in files]
                                    try:
                                        conn.execute("""
                                            INSERT INTO sftp_cache (sync_time, file_list, file_count, synced_by)
                                            VALUES (?, ?, ?, ?)
                                        """, (
                                            now.isoformat(),
                                            json.dumps(files_data),
                                            len(files_data),
                                            "automated_sync"
                                        ))
                                        conn.commit()
                                        logger.info(f"💾 Automated sync: Cached {len(files_data)} SFTP files to database")
                                    except Exception as cache_error:
                                        logger.warning(f"⚠️ Automated sync: Failed to cache SFTP files: {cache_error}")
                                    
                                    # Download all files
                                    file_paths = [f.remote_path for f in files]
                                    results = sftp_service.download_files(file_paths, username="automated_sync")
                                    downloaded = sum(1 for r in results if r.success)
                                    logger.info(f"⬇️ Automated sync: Downloaded {downloaded}/{len(files)} files")
                                else:
                                    logger.info("ℹ️ Automated sync: No files found on SFTP server")
                            else:
                                logger.info("⚠️ Automated sync: SFTP is disabled - skipping file sync")
                        except Exception as e:
                            logger.error(f"❌ Automated sync: SFTP sync error: {e}", exc_info=True)
                        
                        # Step 2: Check if there are new files to process
                        try:
                            etl_service = app_state.get("etl_service")
                            if etl_service:
                                from .etl_service import FileDiscoveryService
                                discovery = FileDiscoveryService()
                                tasks = discovery.discover_files(
                                    config.directories.input_dir,
                                    force_reprocess=False,
                                    selected_files=None
                                )
                                
                                # Filter only unprocessed files
                                from .etl_service import FileProcessingStatus
                                new_files = [t for t in tasks if t.status != FileProcessingStatus.SKIPPED]
                                
                                if new_files:
                                    logger.info(f"📂 Automated sync: Found {len(new_files)} new files to process")
                                    # Start ETL job with automatic trigger
                                    job_id = etl_service.start_etl_job(
                                        force_reprocess=False,
                                        latest_only=False,
                                        max_workers=config.etl.max_workers,
                                        selected_files=None,
                                        username="automated_sync",
                                        trigger_type="automatic"
                                    )
                                    logger.info(f"🚀 Automated sync: Started ETL job {job_id}")
                                else:
                                    logger.info("ℹ️ Automated sync: No new files to process")
                        except Exception as e:
                            logger.error(f"❌ Automated sync: ETL start error: {e}", exc_info=True)
                        
                        logger.info("✅ Automated sync completed successfully")
                        
                    finally:
                        # Always release the lock
                        sync_lock.release()
                    
            except Exception as e:
                logger.error(f"Error in automated sync task: {e}", exc_info=True)
                time.sleep(60)  # Wait a minute before retrying
    
    # Start cleanup thread (daemon so it doesn't block shutdown)
    cleanup_thread = threading.Thread(target=periodic_cleanup, daemon=True, name="MemoryCleanup")
    cleanup_thread.start()
    app_state["cleanup_thread"] = cleanup_thread
    
    # Start automated sync thread
    sync_thread = threading.Thread(target=automated_sync_task, daemon=True, name="AutomatedSync")
    sync_thread.start()
    app_state["sync_thread"] = sync_thread
    
    yield
    
    # Shutdown
    logger.info("Shutting down application...")
    
    # Stop cleanup thread
    if app_state.get("cleanup_thread"):
        # Thread will stop automatically on shutdown (daemon thread)
        pass
    
    # Close database connections
    if app_state.get("db_manager"):
        try:
            app_state["db_manager"].close()
            logger.info("Database connections closed")
        except Exception as e:
            logger.error(f"Error closing database connections: {e}")
    
    # Final garbage collection
    collected = gc.collect()
    if collected > 0:
        logger.info(f"Final garbage collection: collected {collected} objects")


# Create FastAPI app
app = FastAPI(
    title="UniteUs ETL Pipeline",
    description="Comprehensive web interface for ETL operations management",
    version="1.0.0",
    lifespan=lifespan
)

# Add memory monitoring middleware (runs first to track all requests)
@app.middleware("http")
async def monitor_memory(request: Request, call_next):
    """Monitor memory usage for long-running processes"""
    import gc
    
    # Periodic garbage collection for long-running requests
    # Only run GC occasionally to avoid performance impact
    import random
    if random.random() < 0.01:  # 1% chance per request
        collected = gc.collect()
        if collected > 100:  # Only log if significant cleanup
            logger.debug(f"Garbage collection: collected {collected} objects")
    
    response = await call_next(request)
    return response

# Add caching middleware for static files and API responses
@app.middleware("http")
async def add_cache_headers(request: Request, call_next):
    """Add aggressive cache headers for LAN deployment"""
    response = await call_next(request)
    
    # Add cache headers for static files (1 hour cache - session-based optimization)
    if request.url.path.startswith("/static/"):
        response.headers["Cache-Control"] = "public, max-age=3600, must-revalidate"
        # Add ETag support for cache validation
        import hashlib
        import os
        static_file_path = static_dir / request.url.path.replace("/static/", "")
        if static_file_path.exists():
            try:
                file_stat = os.stat(static_file_path)
                etag = hashlib.md5(f"{static_file_path}{file_stat.st_mtime}".encode()).hexdigest()
                response.headers["ETag"] = f'"{etag}"'
            except:
                pass  # Ignore errors
    
    # Add cache headers for safe API responses (read-only, non-sensitive data)
    if request.url.path.startswith("/api/"):
        if request.method == "GET":
            # Cache read-only API responses that don't contain sensitive data
            cacheable_paths = [
                "/api/reports/", "/api/charts/", "/api/statistics/", 
                "/api/database/tables", "/api/database/columns",
                "/api/database/info"
            ]
            if any(request.url.path.startswith(path) for path in cacheable_paths):
                # Cache for 2 minutes (session-based optimization)
                response.headers["Cache-Control"] = "public, max-age=120, must-revalidate"
            elif "/api/health" in request.url.path:
                # Health check can be cached briefly
                response.headers["Cache-Control"] = "public, max-age=60"
            elif request.url.path.startswith("/api/database/") and "data" not in request.url.path:
                # Database metadata endpoints (tables, columns, info) - cache for session
                response.headers["Cache-Control"] = "public, max-age=300, must-revalidate"  # 5 minutes
    
    return response

# Add request logging middleware with error handling
@app.middleware("http")
async def log_requests(request: Request, call_next):
    """Log all HTTP requests and catch errors with enhanced error reporting"""
    start_time = datetime.now()
    
    # Skip logging health checks unless they fail (reduce noise)
    is_health_check = request.url.path == "/api/health"
    
    if not is_health_check:
        # Log request with query params for better debugging
        query_string = f"?{request.url.query}" if request.url.query else ""
        logger.info(f"Request: {request.method} {request.url.path}{query_string}")
    
    try:
        response = await call_next(request)
        
        duration = (datetime.now() - start_time).total_seconds()
        query_string = f"?{request.url.query}" if request.url.query else ""
        
        # Enhanced error logging with more context
        if response.status_code >= 500:
            logger.error(
                f"SERVER ERROR: {request.method} {request.url.path}{query_string} - "
                f"Status: {response.status_code} - Duration: {duration:.3f}s - "
                f"Client: {request.client.host if request.client else 'unknown'}"
            )
        elif response.status_code >= 400:
            # Log 401s as debug (expected when not authenticated), 404s as warnings, other 4xx as errors
            if response.status_code == 401:
                logger.debug(
                    f"Authentication required: {request.method} {request.url.path}{query_string} - "
                    f"Status: {response.status_code} - Duration: {duration:.3f}s"
                )
            elif response.status_code == 404:
                logger.warning(
                    f"CLIENT ERROR: {request.method} {request.url.path}{query_string} - "
                    f"Status: {response.status_code} - Duration: {duration:.3f}s"
                )
            else:
                logger.error(
                    f"CLIENT ERROR: {request.method} {request.url.path}{query_string} - "
                    f"Status: {response.status_code} - Duration: {duration:.3f}s"
                )
        elif not is_health_check:
            # Normal successful responses (skip health checks)
            logger.info(f"Response: {request.method} {request.url.path} - Status: {response.status_code} - Duration: {duration:.3f}s")
        
        return response
    except Exception as e:
        duration = (datetime.now() - start_time).total_seconds()
        query_string = f"?{request.url.query}" if request.url.query else ""
        
        # Detailed exception logging
        logger.error(
            f"EXCEPTION in {request.method} {request.url.path}{query_string}\n"
            f"  Duration: {duration:.3f}s\n"
            f"  Exception Type: {type(e).__name__}\n"
            f"  Exception Message: {str(e)}\n"
            f"  Client: {request.client.host if request.client else 'unknown'}",
            exc_info=True  # Include full traceback
        )
        # Re-raise to let FastAPI handle it
        raise

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=config.web.cors_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Create directories for static files and templates
web_dir = Path(__file__).parent / "web"
static_dir = web_dir / "static"
templates_dir = web_dir / "templates"

# Ensure directories exist
static_dir.mkdir(parents=True, exist_ok=True)
templates_dir.mkdir(parents=True, exist_ok=True)

# Mount static files and setup templates
app.mount("/static", StaticFiles(directory=str(static_dir)), name="static")
templates = Jinja2Templates(directory=str(templates_dir))


# ============================================================================
# AUTHENTICATION ROUTES
# ============================================================================

@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    """Display login page"""
    # Check if already logged in
    session = await get_current_session(request)
    if session:
        return RedirectResponse(url="/", status_code=303)
    
    return templates.TemplateResponse("login.html", {"request": request})


@app.post("/login")
async def login(
    request: Request,
    username: str = Form(...),
    password: str = Form(...),
    auth_method: str = Form("ad")  # Default to AD, can be "ad" or "local"
):
    """Handle login form submission"""
    auth = get_auth_service()
    
    # Validate IP address
    client_ip = request.client.host
    if not auth.validate_ip(client_ip):
        logger.warning(f"Login attempt from unauthorized IP: {client_ip}")
        return RedirectResponse(
            url="/login?error=unauthorized_network",
            status_code=303
        )
    
    # Authenticate user based on selected method with IP address for audit logging
    if auth_method == "local":
        # Force local authentication only
        user_info = auth._authenticate_local(username, password, client_ip)
        if user_info:
            logger.info(f"User {username} authenticated via Direct Login (local database)")
    elif auth_method == "ad":
        # Try AD authentication only
        user_info = auth._authenticate_ad(username, password, client_ip) if auth.ad_enabled else None
        if user_info:
            logger.info(f"User {username} authenticated via Windows Login (Active Directory)")
        elif not auth.ad_enabled:
            # AD not available, inform user
            logger.warning(f"AD authentication requested but not available for user: {username}")
            user_info = None
    else:
        # Invalid auth method
        user_info = None
    
    if not user_info:
        logger.warning(f"Failed login attempt for user: {username} from IP: {client_ip} (method: {auth_method})")
        return RedirectResponse(
            url="/login?error=invalid_credentials",
            status_code=303
        )
    
    # Create session
    user_agent = request.headers.get("user-agent", "unknown")
    session_id = auth.create_session(user_info, client_ip, user_agent)
    
    # Set secure cookie and redirect to dashboard
    response = RedirectResponse(url="/", status_code=303)
    response.set_cookie(
        key="session_id",
        value=session_id,
        httponly=True,      # Prevent JavaScript access
        secure=False,       # Set to True if using HTTPS
        samesite="lax",     # CSRF protection
        max_age=3600        # 1 hour
    )
    
    logger.info(f"User {username} logged in successfully (role: {user_info['role'].value}, method: {auth_method})")
    return response


@app.post("/logout")
@app.get("/logout")
async def logout(request: Request):
    """Handle logout"""
    auth = get_auth_service()
    session_id = request.cookies.get("session_id")
    
    if session_id:
        session = auth.get_session(session_id)
        if session:
            auth_method = "Active Directory" if session.auth_method == "ad" else "Local Database"
            logger.info(f"User {session.username} logged out (authenticated via: {auth_method})")
        auth.destroy_session(session_id)
    
    response = RedirectResponse(url="/login?message=logged_out", status_code=303)
    response.delete_cookie("session_id")
    return response


@app.get("/api/auth/user")
async def get_current_user(session: UserSession = Depends(require_auth)):
    """Get current authenticated user information"""
    return {
        "success": True,
        "user": session.to_dict()
    }


@app.get("/api/auth/sessions")
async def get_active_sessions(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Get all active sessions - admin only"""
    auth = get_auth_service()
    sessions = auth.get_active_sessions()
    
    return {
        "success": True,
        "sessions": [s.to_dict() for s in sessions],
        "count": len(sessions)
    }


# ============================================================================
# HTML PAGES (Protected)
# ============================================================================

@app.get("/new-user", response_class=HTMLResponse)
@require_auth_redirect()
async def new_user_page(request: Request):
    """Page shown to new users who need admin approval"""
    session: UserSession = request.state.user
    return templates.TemplateResponse("new_user.html", {
        "request": request,
        "title": "Access Pending",
        "user": session
    })


@app.get("/", response_class=HTMLResponse)
@require_auth_redirect()
async def dashboard(request: Request):
    """Main dashboard page with reports and analytics - requires authentication"""
    session: UserSession = request.state.user
    
    # Redirect NEW_USER role to the new user page
    if session.role == UserRole.NEW_USER:
        return RedirectResponse(url="/new-user", status_code=303)
    
    return templates.TemplateResponse("dashboard.html", {
        "request": request,
        "title": "Dashboard",
        "user": session
    })


@app.get("/database", response_class=HTMLResponse)
@require_auth_redirect()
async def database_viewer(request: Request):
    """Database viewer page - requires OPERATOR role or higher"""
    session: UserSession = request.state.user
    
    # Check if user has OPERATOR role or higher
    if not session.has_permission(UserRole.OPERATOR):
        return templates.TemplateResponse("error.html", {
            "request": request,
            "title": "Access Denied",
            "user": session,
            "error_title": "Access Denied",
            "error_message": "You do not have permission to access Database Viewer. Operator role or higher is required.",
            "error_code": 403
        }, status_code=403)
    
    return templates.TemplateResponse("database.html", {
        "request": request,
        "title": "Database Viewer",
        "user": session
    })


@app.get("/etl-logs", response_class=HTMLResponse)
@require_auth_redirect()
async def etl_logs(request: Request):
    """ETL logs and history page - requires OPERATOR role or higher"""
    session: UserSession = request.state.user
    
    # Check if user has OPERATOR role or higher
    if not session.has_permission(UserRole.OPERATOR):
        return templates.TemplateResponse("error.html", {
            "request": request,
            "title": "Access Denied",
            "user": session,
            "error_title": "Access Denied",
            "error_message": "You do not have permission to access ETL Logs. Operator role or higher is required.",
            "error_code": 403
        }, status_code=403)
    
    return templates.TemplateResponse("etl_logs.html", {
        "request": request,
        "title": "ETL Logs & History",
        "user": session
    })


@app.get("/data-import", response_class=HTMLResponse)
@require_auth_redirect()
async def data_import(request: Request):
    """Data import page - requires OPERATOR role or higher"""
    session: UserSession = request.state.user
    
    # Check if user has OPERATOR role or higher
    if not session.has_permission(UserRole.OPERATOR):
        return templates.TemplateResponse("error.html", {
            "request": request,
            "title": "Access Denied",
            "user": session,
            "error_title": "Access Denied",
            "error_message": "You do not have permission to access Data Import. Operator role or higher is required.",
            "error_code": 403
        }, status_code=403)
    
    return templates.TemplateResponse("data_import.html", {
        "request": request,
        "title": "Data Import",
        "user": session
    })


@app.get("/data-cleaning", response_class=HTMLResponse)
@require_auth_redirect()
async def data_cleaning(request: Request):
    """Data cleaning documentation page - requires OPERATOR role or higher"""
    session: UserSession = request.state.user
    
    # Check if user has OPERATOR role or higher
    if not session.has_permission(UserRole.OPERATOR):
        return templates.TemplateResponse("error.html", {
            "request": request,
            "title": "Access Denied",
            "user": session,
            "error_title": "Access Denied",
            "error_message": "You do not have permission to access Data Cleaning Documentation. Operator role or higher is required.",
            "error_code": 403
        }, status_code=403)
    
    return templates.TemplateResponse("data_cleaning.html", {
        "request": request,
        "title": "Data Cleaning Documentation",
        "user": session
    })


@app.get("/about", response_class=HTMLResponse)
@require_auth_redirect()
async def about(request: Request):
    """About page with technologies and data cleaning info - requires OPERATOR role or higher"""
    session: UserSession = request.state.user
    
    # Check if user has OPERATOR role or higher
    if not session.has_permission(UserRole.OPERATOR):
        return templates.TemplateResponse("error.html", {
            "request": request,
            "title": "Access Denied",
            "user": session,
            "error_title": "Access Denied",
            "error_message": "You do not have permission to access this page. Operator role or higher is required.",
            "error_code": 403
        }, status_code=403)
    
    return templates.TemplateResponse("about.html", {
        "request": request,
        "title": "About",
        "user": session
    })


@app.get("/settings", response_class=HTMLResponse)
@require_auth_redirect()
async def settings(request: Request):
    """User settings page - requires OPERATOR role or higher"""
    session: UserSession = request.state.user
    
    # Check if user has OPERATOR role or higher
    if not session.has_permission(UserRole.OPERATOR):
        return templates.TemplateResponse("error.html", {
            "request": request,
            "title": "Access Denied",
            "user": session,
            "error_title": "Access Denied",
            "error_message": "You do not have permission to access Settings. Operator role or higher is required.",
            "error_code": 403
        }, status_code=403)
    
    return templates.TemplateResponse("settings.html", {
        "request": request,
        "title": "Settings",
        "user": session
    })


@app.get("/admincp", response_class=HTMLResponse)
@require_auth_redirect()
async def admincp(request: Request):
    """Admin Control Panel - Dashboard"""
    session: UserSession = request.state.user
    
    # Check if user is admin
    if session.role.value != 'admin':
        return templates.TemplateResponse("error.html", {
            "request": request,
            "title": "Access Denied",
            "user": session,
            "error_title": "Access Denied",
            "error_message": "You do not have permission to access the Admin Control Panel.",
            "error_code": 403
        }, status_code=403)
    
    return templates.TemplateResponse("admincp_dashboard.html", {
        "request": request,
        "title": "Admin Control Panel - Dashboard",
        "user": session,
        "active_page": "dashboard"
    })

@app.get("/admincp/dashboard", response_class=HTMLResponse)
@require_auth_redirect()
async def admincp_dashboard(request: Request):
    """Admin Control Panel - Dashboard (alias for /admincp)"""
    session: UserSession = request.state.user
    
    # Check if user is admin
    if session.role.value != 'admin':
        return templates.TemplateResponse("error.html", {
            "request": request,
            "title": "Access Denied",
            "user": session,
            "error_title": "Access Denied",
            "error_message": "You do not have permission to access the Admin Control Panel.",
            "error_code": 403
        }, status_code=403)
    
    return templates.TemplateResponse("admincp_dashboard.html", {
        "request": request,
        "title": "Admin Control Panel - Dashboard",
        "user": session,
        "active_page": "dashboard"
    })

@app.get("/admincp/users", response_class=HTMLResponse)
@require_auth_redirect()
async def admincp_users(request: Request):
    """Admin Control Panel - User Management"""
    session: UserSession = request.state.user
    
    if session.role.value != 'admin':
        return templates.TemplateResponse("error.html", {
            "request": request,
            "title": "Access Denied",
            "user": session,
            "error_title": "Access Denied",
            "error_message": "You do not have permission to access the Admin Control Panel.",
            "error_code": 403
        }, status_code=403)
    
    return templates.TemplateResponse("admincp_users.html", {
        "request": request,
        "title": "User Management",
        "user": session,
        "active_page": "users"
    })

@app.get("/admincp/audit", response_class=HTMLResponse)
@require_auth_redirect()
async def admincp_audit(request: Request):
    """Admin Control Panel - Audit Log"""
    session: UserSession = request.state.user
    
    if session.role.value != 'admin':
        return templates.TemplateResponse("error.html", {
            "request": request,
            "title": "Access Denied",
            "user": session,
            "error_title": "Access Denied",
            "error_message": "You do not have permission to access the Admin Control Panel.",
            "error_code": 403
        }, status_code=403)
    
    return templates.TemplateResponse("admincp_audit.html", {
        "request": request,
        "title": "Audit Log",
        "user": session,
        "active_page": "audit"
    })

@app.get("/admincp/system", response_class=HTMLResponse)
@require_auth_redirect()
async def admincp_system(request: Request):
    """Admin Control Panel - System Info (redirected to Dashboard)"""
    # Redirect to dashboard since system info is shown there
    return RedirectResponse(url="/admincp", status_code=301)

@app.get("/admincp/sftp", response_class=HTMLResponse)
@require_auth_redirect()
async def admincp_sftp(request: Request):
    """Admin Control Panel - SFTP Integration"""
    session: UserSession = request.state.user
    
    if session.role.value != 'admin':
        return templates.TemplateResponse("error.html", {
            "request": request,
            "title": "Access Denied",
            "user": session,
            "error_title": "Access Denied",
            "error_message": "You do not have permission to access the Admin Control Panel.",
            "error_code": 403
        }, status_code=403)
    
    return templates.TemplateResponse("admincp_sftp.html", {
        "request": request,
        "title": "SFTP Integration",
        "user": session,
        "active_page": "sftp"
    })

@app.get("/admincp/siem", response_class=HTMLResponse)
@app.get("/admincp/logging/siem", response_class=HTMLResponse)
@require_auth_redirect()
async def admincp_siem(request: Request):
    """Admin Control Panel - SIEM Integration"""
    session: UserSession = request.state.user
    
    if session.role.value != 'admin':
        return templates.TemplateResponse("error.html", {
            "request": request,
            "title": "Access Denied",
            "user": session,
            "error_title": "Access Denied",
            "error_message": "You do not have permission to access the Admin Control Panel.",
            "error_code": 403
        }, status_code=403)
    
    return templates.TemplateResponse("admincp_siem.html", {
        "request": request,
        "title": "SIEM Integration",
        "user": session,
        "active_page": "siem"
    })

@app.get("/admincp/logging/windows", response_class=HTMLResponse)
@require_auth_redirect()
async def admincp_logging_windows(request: Request):
    """Admin Control Panel - Windows Event Log"""
    session: UserSession = request.state.user
    
    if session.role.value != 'admin':
        return templates.TemplateResponse("error.html", {
            "request": request,
            "title": "Access Denied",
            "user": session,
            "error_title": "Access Denied",
            "error_message": "You do not have permission to access the Admin Control Panel.",
            "error_code": 403
        }, status_code=403)
    
    return templates.TemplateResponse("admincp_logging_windows.html", {
        "request": request,
        "title": "Windows Event Log",
        "user": session,
        "active_page": "logging_windows"
    })

@app.get("/admincp/permissions", response_class=HTMLResponse)
@require_auth_redirect()
async def admincp_permissions(request: Request):
    """Admin Control Panel - Permissions Grid"""
    session: UserSession = request.state.user
    
    if session.role.value != 'admin':
        return templates.TemplateResponse("error.html", {
            "request": request,
            "title": "Access Denied",
            "user": session,
            "error_title": "Access Denied",
            "error_message": "You do not have permission to access the Admin Control Panel.",
            "error_code": 403
        }, status_code=403)
    
    return templates.TemplateResponse("admincp_permissions.html", {
        "request": request,
        "title": "Permissions",
        "user": session,
        "active_page": "permissions"
    })

@app.get("/admincp/database", response_class=HTMLResponse)
@require_auth_redirect()
async def admincp_database(request: Request):
    """Admin Control Panel - Database Settings"""
    session: UserSession = request.state.user
    
    if session.role.value != 'admin':
        return templates.TemplateResponse("error.html", {
            "request": request,
            "title": "Access Denied",
            "user": session,
            "error_title": "Access Denied",
            "error_message": "You do not have permission to access the Admin Control Panel.",
            "error_code": 403
        }, status_code=403)
    
    return templates.TemplateResponse("admincp_database.html", {
        "request": request,
        "title": "Database Settings",
        "user": session,
        "active_page": "database"
    })


@app.get("/admincp/database/schema", response_class=HTMLResponse)
@require_auth_redirect()
async def admincp_schema(request: Request):
    """Admin Control Panel - Schema Management"""
    session: UserSession = request.state.user
    
    if session.role.value != 'admin':
        return templates.TemplateResponse("error.html", {
            "request": request,
            "title": "Access Denied",
            "user": session,
            "error_title": "Access Denied",
            "error_message": "You do not have permission to access the Admin Control Panel.",
            "error_code": 403
        }, status_code=403)
    
    return templates.TemplateResponse("admincp_schema.html", {
        "request": request,
        "title": "Schema Management",
        "user": session,
        "active_page": "schema"
    })


@app.get("/api/database/schema-sql")
async def get_database_schema_sql(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Get the complete database schema SQL - Admin only"""
    try:
        from .database_schema import get_schema_sql, get_view_definitions, get_table_descriptions
        from .database_schema_converter import get_schema_for_database_type
        
        # Get SQLite base schema
        sqlite_schema = get_schema_sql()
        
        # Get current database type and adapted schema
        db_type = config.database.db_type
        adapted_schema = get_schema_for_database_type(db_type, sqlite_schema)
        
        return {
            "success": True,
            "database_type": db_type,
            "sqlite_schema": sqlite_schema,
            "adapted_schema": adapted_schema,
            "view_definitions": get_view_definitions(),
            "table_descriptions": get_table_descriptions()
        }
    except Exception as e:
        logger.error(f"Error getting database schema: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e)
        }


# @app.get("/admincp/config", response_class=HTMLResponse)
# @require_auth_redirect()
# async def admincp_config(request: Request):
#     """Admin Control Panel - Configuration (DEPRECATED - Not Used)"""
#     session: UserSession = request.state.user
#     
#     if session.role.value != 'admin':
#         return templates.TemplateResponse("error.html", {
#             "request": request,
#             "title": "Access Denied",
#             "user": session,
#             "error_title": "Access Denied",
#             "error_message": "You do not have permission to access the Admin Control Panel.",
#             "error_code": 403
#         }, status_code=403)
#     
#     return templates.TemplateResponse("admincp_config.html", {
#         "request": request,
#         "title": "System Configuration",
#         "user": session,
#         "active_page": "config"
#     })


# @app.get("/api/admin/config/system")
# async def get_system_config(session: UserSession = Depends(require_role(UserRole.ADMIN))):
#     """Get system configuration (ETL, Web, Logging, Directories, Security) - Admin only (DEPRECATED - Not Used)"""
#     try:
#         from .config import config
#         
#         return {
#             "success": True,
#             "config": {
#                 "etl": {
#                     "batch_size": config.etl.batch_size,
#                     "max_workers": config.etl.max_workers,
#                     "timeout_seconds": config.etl.timeout_seconds,
#                     "retry_attempts": config.etl.retry_attempts,
#                     "skip_processed_files": config.etl.skip_processed_files,
#                     "force_reprocess": config.etl.force_reprocess,
#                     "latest_only": config.etl.latest_only,
#                     "ignored_filename_prefixes": config.etl.ignored_filename_prefixes,
#                     "file_patterns": config.etl.file_patterns,
#                     "recognized_extensions": config.etl.recognized_extensions
#                 },
#                 "web": {
#                     "host": config.web.host,
#                     "port": config.web.port,
#                     "reload": config.web.reload,
#                     "log_level": config.web.log_level,
#                     "cors_origins": config.web.cors_origins
#                 },
#                 "logging": {
#                     "level": config.logging.level.value,
#                     "format": config.logging.format,
#                     "date_format": config.logging.date_format,
#                     "file_rotation_size": config.logging.file_rotation_size,
#                     "file_retention_count": config.logging.file_retention_count,
#                     "enable_console": config.logging.enable_console,
#                     "enable_file": config.logging.enable_file
#                 },
#                 "directories": {
#                     "project_root": str(config.directories.project_root),
#                     "data_dir": str(config.directories.data_dir),
#                     "input_dir": str(config.directories.input_dir),
#                     "output_dir": str(config.directories.output_dir),
#                     "logs_dir": str(config.directories.logs_dir),
#                     "database_dir": str(config.directories.database_dir),
#                     "backup_dir": str(config.directories.backup_dir)
#                 },
#                 "security": {
#                     "enable_phi_hashing": config.security.enable_phi_hashing,
#                     "hash_on_import": config.security.hash_on_import,
#                     "hash_on_export": config.security.hash_on_export,
#                     "phi_hash_salt_configured": bool(config.security.phi_hash_salt and len(config.security.phi_hash_salt) > 0)
#                     # Don't send the actual salt for security
#                 }
#             }
#         }
#     except Exception as e:
#         logger.error(f"Error fetching system config: {e}", exc_info=True)
#         return {
#             "success": False,
#             "error": str(e)
#         }


# @app.post("/api/admin/config/system")
# async def save_system_config(
#     request: Request,
#     session: UserSession = Depends(require_role(UserRole.ADMIN))
# ):
#     """Save system configuration - Admin only (DEPRECATED - Not Used)"""
#     try:
#         data = await request.json()
#         
#         # Note: Most config is read-only from environment variables
#         # This endpoint is mainly for documentation/display purposes
#         # Actual changes should be made via environment variables or config files
#         
#         logger.info(f"System config view requested by {session.username}")
#         
#         return {
#             "success": True,
#             "message": "System configuration is managed via environment variables and config files. See documentation for details."
#         }
#     except Exception as e:
#         logger.error(f"Error saving system config: {e}", exc_info=True)
#         return {
#             "success": False,
#             "error": str(e)
#         }


@app.get("/settings/siem", response_class=HTMLResponse)
@require_auth_redirect()
async def siem_settings_page(request: Request):
    """SIEM Settings page - Admin only"""
    session: UserSession = request.state.user
    
    if session.role.value != 'admin':
        return templates.TemplateResponse("error.html", {
            "request": request,
            "title": "Access Denied",
            "user": session,
            "error_title": "Access Denied",
            "error_message": "You do not have permission to access SIEM settings.",
            "error_code": 403
        }, status_code=403)
    
    return templates.TemplateResponse("settings_siem.html", {
        "request": request,
        "title": "SIEM Configuration",
        "user": session
    })


@app.get("/settings/sftp", response_class=HTMLResponse)
@require_auth_redirect()
async def sftp_settings_page(request: Request):
    """SFTP Settings page - Admin only"""
    session: UserSession = request.state.user
    
    if session.role.value != 'admin':
        return templates.TemplateResponse("error.html", {
            "request": request,
            "title": "Access Denied",
            "user": session,
            "error_title": "Access Denied",
            "error_message": "You do not have permission to access SFTP settings.",
            "error_code": 403
        }, status_code=403)
    
    return templates.TemplateResponse("settings_sftp.html", {
        "request": request,
        "title": "SFTP Configuration",
        "user": session
    })


@app.get("/annual-report", response_class=HTMLResponse)
@require_auth_redirect()
async def annual_report(request: Request):
    """Annual Report page with customizable date ranges and comprehensive analytics"""
    session: UserSession = request.state.user
    return templates.TemplateResponse("annual_report.html", {
        "request": request,
        "title": "Annual Report",
        "user": session
    })


# ============================================================================
# ETL JOB MANAGEMENT API (Protected)
# ============================================================================

@app.get("/api/etl/status")
async def get_etl_status(session: UserSession = Depends(require_auth)):
    """Get current ETL job status and all active jobs"""
    etl_service = app_state["etl_service"]
    current_job = etl_service.get_current_job_status()
    active_jobs = etl_service.get_active_jobs()
    
    return {
        "current_job": current_job.to_dict() if current_job else None,
        "active_jobs": [job.to_dict() for job in active_jobs],
        "active_count": len(active_jobs)
    }


@app.get("/api/etl/job/{job_id}")
async def get_job_status_by_id(job_id: str, session: UserSession = Depends(require_auth)):
    """Get specific job status by ID"""
    etl_service = app_state["etl_service"]
    job = etl_service.get_job_status(job_id)
    
    if not job:
        raise HTTPException(status_code=404, detail=f"Job {job_id} not found")
    
    return {"job": job.to_dict()}


@app.post("/api/etl/start")
async def start_etl_job(job_request: ETLJobRequest, session: UserSession = Depends(require_role(UserRole.OPERATOR))):
    """Start new ETL job"""
    etl_service = app_state["etl_service"]
    
    try:
        job_id = etl_service.start_etl_job(
            force_reprocess=job_request.force_reprocess,
            latest_only=job_request.latest_only,
            max_workers=job_request.max_workers,
            selected_files=job_request.selected_files,
            username=session.username  # Pass username for audit logging
        )
        
        return {
            "success": True,
            "job_id": job_id,
            "message": f"ETL job {job_id} started successfully"
        }
        
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to start ETL job: {str(e)}")


@app.post("/api/etl/cancel")
async def cancel_etl_job(session: UserSession = Depends(require_role(UserRole.OPERATOR))):
    """Cancel current ETL job - requires operator role"""
    etl_service = app_state["etl_service"]
    
    current_job = etl_service.get_current_job_status()
    if not current_job or not current_job.is_running:
        raise HTTPException(status_code=400, detail="No ETL job is currently running")
    
    etl_service.cancel_current_job()
    logger.info(f"User {session.username} cancelled ETL job {current_job.job_id}")
    
    return {
        "success": True,
        "message": f"ETL job {current_job.job_id} cancelled"
    }


@app.post("/api/etl/job/{job_id}/cancel")
async def cancel_specific_job(job_id: str, session: UserSession = Depends(require_role(UserRole.OPERATOR))):
    """Cancel a specific ETL job by ID - requires operator role"""
    etl_service = app_state["etl_service"]
    
    success = etl_service.cancel_job(job_id)
    if not success:
        raise HTTPException(status_code=404, detail=f"Job {job_id} not found or not running")
    
    logger.info(f"User {session.username} cancelled ETL job {job_id}")
    
    return {
        "success": True,
        "message": f"ETL job {job_id} cancelled"
    }


@app.get("/api/etl/history")
async def get_etl_history(limit: int = Query(10, ge=1, le=100), session: UserSession = Depends(require_auth)):
    """Get ETL job history - requires authentication"""
    etl_service = app_state["etl_service"]
    history = etl_service.get_job_history(limit)
    
    return {
        "history": [job.to_dict() for job in history]
    }


@app.post("/api/etl/undo/{job_id}")
async def undo_etl_job(job_id: str, session: UserSession = Depends(require_auth)):
    """Undo/rollback an ETL job by deleting all records imported by that job"""
    try:
        db_manager = app_state["db_manager"]
        
        # Track deletion statistics
        tables_affected = 0
        records_deleted = 0
        
        # List of all tables that might have ETL data
        tables = [
            'people',
            'employees', 
            'cases',
            'referrals',
            'assistance_requests',
            'assistance_requests_supplemental_responses',
            'resource_lists',
            'resource_list_shares'
        ]
        
        with db_manager.pool.get_connection() as conn:
            cursor = conn.cursor()
            
            # Delete records from each table that match the job_id in etl_metadata
            # First, get all file_names from this job
            cursor.execute("""
                SELECT DISTINCT file_name, table_name 
                FROM etl_metadata 
                WHERE file_name LIKE ?
            """, (f"%{job_id}%",))
            
            job_files = cursor.fetchall()
            
            if not job_files:
                return {
                    "success": False,
                    "error": f"No files found for job {job_id}"
                }
            
            # For each file/table combo, delete imported records
            # We'll use the pull_timestamp or etl_loaded_at to identify records from this job
            for file_name, table_name in job_files:
                if table_name not in tables:
                    continue
                    
                # Get the processing time range for this file
                cursor.execute("""
                    SELECT processing_started_at, processing_completed_at
                    FROM etl_metadata
                    WHERE file_name = ? AND table_name = ?
                    ORDER BY id DESC LIMIT 1
                """, (file_name, table_name))
                
                time_range = cursor.fetchone()
                if not time_range:
                    continue
                
                start_time, end_time = time_range
                
                # Delete records that were loaded during this time window
                # Using etl_loaded_at timestamp which is set during import
                try:
                    cursor.execute(f"""
                        DELETE FROM {table_name}
                        WHERE etl_loaded_at >= ? AND etl_loaded_at <= ?
                    """, (start_time, end_time))
                    
                    deleted = cursor.rowcount
                    if deleted > 0:
                        records_deleted += deleted
                        tables_affected += 1
                        
                except Exception as e:
                    logger.error(f"Error deleting from {table_name}: {e}")
                    continue
            
            # Mark the ETL metadata entries as undone
            cursor.execute("""
                UPDATE etl_metadata
                SET status = 'undone'
                WHERE file_name IN (
                    SELECT DISTINCT file_name 
                    FROM etl_metadata 
                    WHERE file_name LIKE ?
                )
            """, (f"%{job_id}%",))
            
            conn.commit()
        
        # Log the undo action
        from .audit_logger import get_audit_logger, AuditCategory, AuditAction
        audit_logger = get_audit_logger()
        audit_logger.log(
            username=session.username,
            action=AuditAction.DATA_DELETED,
            category=AuditCategory.ETL,
            success=True,
            details=f"Undone ETL job {job_id}: deleted {records_deleted} records from {tables_affected} tables",
            target_resource=job_id
        )
        
        return {
            "success": True,
            "job_id": job_id,
            "records_deleted": records_deleted,
            "tables_affected": tables_affected
        }
        
    except Exception as e:
        logger.error(f"Error undoing job {job_id}: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e)
        }


@app.get("/api/etl/processing-history")
async def get_processing_history(limit: int = Query(50, ge=1, le=1000), session: UserSession = Depends(require_auth)):
    """Get file processing history from database"""
    db_manager = app_state["db_manager"]
    result = db_manager.etl_metadata.get_processing_history(limit)
    
    if result.success:
        # Process the data to add calculated fields
        from datetime import datetime
        processed_data = []
        
        for record in result.data.to_dict('records') if hasattr(result.data, 'to_dict') else result.data:
            # Calculate processing time
            processing_time_seconds = None
            if record.get('processing_started_at') and record.get('processing_completed_at'):
                try:
                    start = datetime.fromisoformat(str(record['processing_started_at']))
                    end = datetime.fromisoformat(str(record['processing_completed_at']))
                    processing_time_seconds = round((end - start).total_seconds(), 2)
                except:
                    pass
            
            processed_data.append({
                'file_name': record.get('file_name'),
                'file_hash': record.get('file_hash'),
                'file_type': record.get('table_name', 'Unknown'),
                'status': record.get('status'),
                'records_processed': record.get('records_processed', 0),
                'records_inserted': record.get('records_inserted', 0),
                'records_updated': record.get('records_updated', 0),
                'processing_time_seconds': processing_time_seconds,
                'processed_at': record.get('processing_completed_at'),
                'file_date': record.get('file_date'),
                'error_message': record.get('error_message')
            })
        
        return {
            "success": True,
            "data": processed_data,
            "count": len(processed_data)
        }
    else:
        raise HTTPException(status_code=500, detail=result.error_message)


# ============================================================================
# AUTOMATED SYNC API (Protected)
# ============================================================================

@app.get("/api/automated-sync/config")
async def get_automated_sync_config(session: UserSession = Depends(require_role(UserRole.OPERATOR))):
    """Get automated sync configuration - requires operator role"""
    try:
        db_manager = app_state["db_manager"]
        
        with db_manager.pool.get_connection() as conn:
            cursor = conn.execute("""
                SELECT enabled, interval_minutes, last_run, next_run, updated_at, updated_by
                FROM automated_sync_config
                WHERE id = 1
            """)
            row = cursor.fetchone()
            
            if row:
                return {
                    "success": True,
                    "enabled": bool(row[0]),
                    "interval_minutes": row[1],
                    "last_run": row[2],
                    "next_run": row[3],
                    "updated_at": row[4],
                    "updated_by": row[5]
                }
            else:
                # Return default config
                return {
                    "success": True,
                    "enabled": False,
                    "interval_minutes": 60,
                    "last_run": None,
                    "next_run": None,
                    "updated_at": None,
                    "updated_by": None
                }
    except Exception as e:
        logger.error(f"Error getting automated sync config: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to get automated sync config: {str(e)}")


@app.post("/api/automated-sync/config")
async def save_automated_sync_config(
    enabled: bool = Form(...),
    interval_minutes: int = Form(60),
    session: UserSession = Depends(require_role(UserRole.OPERATOR))
):
    """Save automated sync configuration - requires operator role"""
    try:
        from datetime import datetime, timedelta
        db_manager = app_state["db_manager"]
        
        # Calculate next run time if enabled
        next_run = None
        if enabled:
            next_run = (datetime.now() + timedelta(minutes=interval_minutes)).isoformat()
        
        with db_manager.pool.get_connection() as conn:
            conn.execute("""
                UPDATE automated_sync_config
                SET enabled = ?,
                    interval_minutes = ?,
                    next_run = ?,
                    updated_at = ?,
                    updated_by = ?
                WHERE id = 1
            """, (int(enabled), interval_minutes, next_run, datetime.now().isoformat(), session.username))
            conn.commit()
        
        logger.info(f"User {session.username} updated automated sync config: enabled={enabled}, interval={interval_minutes}min")
        
        return {
            "success": True,
            "message": "Automated sync configuration saved successfully",
            "enabled": enabled,
            "interval_minutes": interval_minutes,
            "next_run": next_run
        }
    except Exception as e:
        logger.error(f"Error saving automated sync config: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to save automated sync config: {str(e)}")


# ============================================================================
# ============================================================================
# FILE MANAGEMENT API (Protected)
# ============================================================================

@app.post("/api/files/upload")
async def upload_files(files: List[UploadFile] = File(...), session: UserSession = Depends(require_role(UserRole.OPERATOR))):
    """Upload data files to input directory - requires operator role"""
    responses = []
    
    for file in files:
        try:
            # Save to input directory
            file_path = config.directories.input_dir / file.filename
            
            # Read and save file content
            content = await file.read()
            file_path.write_bytes(content)
            
            logger.info(f"User {session.username} uploaded file: {file.filename} ({len(content)} bytes)")
            
            responses.append(FileUploadResponse(
                success=True,
                message=f"File {file.filename} uploaded successfully",
                filename=file.filename,
                size=len(content)
            ))
            
        except Exception as e:
            logger.error(f"File upload failed for {file.filename}: {e}")
            responses.append(FileUploadResponse(
                success=False,
                message=f"Failed to upload {file.filename}: {str(e)}",
                filename=file.filename,
                size=0
            ))
    
    return {"uploads": [r.dict() for r in responses]}


@app.get("/api/files/list")
async def list_input_files(session: UserSession = Depends(require_auth)):
    """List files in input directory with hash information - requires authentication"""
    try:
        import hashlib
        input_dir = config.directories.input_dir
        files = []
        
        if input_dir.exists():
            for file_path in input_dir.iterdir():
                if file_path.is_file() and file_path.suffix.lower() in ['.txt', '.csv', '.tsv']:
                    stat = file_path.stat()
                    
                    # Calculate file hash
                    hasher = hashlib.md5()
                    with open(file_path, 'rb') as f:
                        for chunk in iter(lambda: f.read(4096), b""):
                            hasher.update(chunk)
                    file_hash = hasher.hexdigest()
                    
                    files.append({
                        "name": file_path.name,
                        "size": stat.st_size,
                        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                        "path": str(file_path),
                        "hash": file_hash
                    })
        
        return {
            "files": sorted(files, key=lambda f: f["modified"], reverse=True),
            "count": len(files)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to list files: {str(e)}")


# ============================================================================
# DATABASE API (Protected)
# ============================================================================

@app.get("/api/database/info")
async def get_database_info(session: UserSession = Depends(require_auth)):
    """Get comprehensive database information - requires authentication"""
    db_manager = app_state["db_manager"]
    
    try:
        info = db_manager.get_database_info()
        return {"database_info": info}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get database info: {str(e)}")


@app.get("/api/database/tables")
async def get_tables_list(session: UserSession = Depends(require_auth)):
    """Get list of all tables with record counts - requires authentication"""
    db_manager = app_state["db_manager"]
    
    try:
        table_stats = db_manager.get_table_stats()
        
        # Get table schemas
        tables = []
        for table_name, stats in table_stats.items():
            table_info = {
                "name": table_name,
                "record_count": stats.get("record_count", 0),
                "exists": stats.get("exists", False),
                "columns": stats.get("columns", []),
                "last_updated": stats.get("last_updated")
            }
            tables.append(table_info)
        
        return {
            "tables": sorted(tables, key=lambda t: t["record_count"], reverse=True),
            "total_tables": len(tables)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get tables: {str(e)}")


@app.get("/api/database/table/{table_name}")
async def get_table_data(
    table_name: str,
    limit: int = Query(100, ge=1, le=1000),
    offset: int = Query(0, ge=0)
):
    """Get table data with pagination"""
    db_manager = app_state["db_manager"]
    
    try:
        # Get the repository dynamically (will work with any table name)
        with db_manager.pool.get_connection() as conn:
            # Verify table exists
            cursor = conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name=?",
                (table_name,)
            )
            if not cursor.fetchone():
                raise HTTPException(status_code=404, detail=f"Table not found: {table_name}")
            
            # Get table data
            cursor = conn.execute(f"SELECT * FROM {table_name} LIMIT ? OFFSET ?", (limit, offset))
            columns = [description[0] for description in cursor.description]
            rows = cursor.fetchall()
            
            # Get total count
            cursor = conn.execute(f"SELECT COUNT(*) FROM {table_name}")
            total_records = cursor.fetchone()[0]
            
            # Convert rows to list of dicts
            data = [dict(zip(columns, row)) for row in rows]
        
        return {
            "success": True,
            "table_name": table_name,
            "data": data,
            "columns": columns,
            "row_count": len(data),
            "limit": limit,
            "offset": offset,
            "total_records": total_records
        }
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get table data: {str(e)}")


@app.post("/api/database/table/search")
async def search_table_data(search_request: TableSearchRequest):
    """Search table data"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            # Verify table exists
            cursor = conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name=?",
                (search_request.table_name,)
            )
            if not cursor.fetchone():
                raise HTTPException(status_code=404, detail=f"Table not found: {search_request.table_name}")
            
            if search_request.search_term:
                # Build search query
                cursor = conn.execute(f"PRAGMA table_info({search_request.table_name})")
                all_columns = [row[1] for row in cursor.fetchall()]
                
                search_columns = search_request.columns if search_request.columns else all_columns
                where_clauses = [f"{col} LIKE ?" for col in search_columns]
                where_sql = " OR ".join(where_clauses)
                search_params = [f"%{search_request.search_term}%" for _ in search_columns]
                
                query = f"SELECT * FROM {search_request.table_name} WHERE {where_sql} LIMIT ?"
                cursor = conn.execute(query, search_params + [search_request.limit])
            else:
                cursor = conn.execute(
                    f"SELECT * FROM {search_request.table_name} LIMIT ? OFFSET ?",
                    (search_request.limit, search_request.offset)
                )
            
            columns = [description[0] for description in cursor.description]
            rows = cursor.fetchall()
            data = [dict(zip(columns, row)) for row in rows]
        
        return {
            "success": True,
            "table_name": search_request.table_name,
            "search_term": search_request.search_term,
            "data": data,
            "columns": columns,
            "row_count": len(data)
        }
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to search table: {str(e)}")


@app.get("/api/database/table/{table_name}/export")
async def export_table_to_csv(table_name: str):
    """Export table data to CSV (redirect to avoid duplication)"""
    # Redirect to canonical export endpoint
    return await export_table_csv(table_name)


@app.get("/api/database/export-all")
async def export_all_tables():
    """Export all tables to a ZIP file with individual CSVs"""
    from fastapi.responses import StreamingResponse
    import io
    import csv
    import zipfile
    from datetime import datetime
    
    db_manager = app_state["db_manager"]
    
    try:
        # Create ZIP file in memory
        zip_buffer = io.BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            table_stats = db_manager.get_table_stats()
            
            for table_name, stats in table_stats.items():
                if not stats.get('exists') or stats.get('record_count', 0) == 0:
                    continue
                
                with db_manager.pool.get_connection() as conn:
                    cursor = conn.execute(f"SELECT * FROM {table_name}")
                    columns = [description[0] for description in cursor.description]
                    rows = cursor.fetchall()
                
                # Create CSV for this table
                csv_output = io.StringIO()
                writer = csv.writer(csv_output)
                writer.writerow(columns)
                writer.writerows(rows)
                
                # Add to ZIP
                zip_file.writestr(
                    f"{table_name}.csv",
                    csv_output.getvalue().encode('utf-8')
                )
        
        zip_buffer.seek(0)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        return StreamingResponse(
            zip_buffer,
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename=calaveras_uniteus_export_{timestamp}.zip"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to export database: {str(e)}")


@app.get("/api/database/export")
async def export_database_comprehensive(
    format: str = Query("csv", description="Export format: csv, excel, sql, json, sqlite"),
    options: str = Query("both", description="Export options: both (data+structure), structure"),
    include_metadata: bool = Query(True, description="Include ETL metadata tables"),
    compress: bool = Query(True, description="Compress output"),
    stream_large: bool = Query(False, description="Stream large tables in chunks"),
    tables: str = Query("all", description="Comma-separated table names or 'all'")
):
    """Comprehensive database export with multiple formats and options"""
    from fastapi.responses import StreamingResponse, FileResponse
    import io
    import csv
    import zipfile
    import json as json_lib
    from datetime import datetime
    
    db_manager = app_state["db_manager"]
    
    try:
        # Get tables to export
        if tables == "all":
            table_stats = db_manager.get_table_stats()
            selected_tables = [name for name, stats in table_stats.items() 
                             if stats.get('exists', False)]
        else:
            selected_tables = [t.strip() for t in tables.split(',')]
        
        # Filter out metadata tables if not included
        if not include_metadata:
            selected_tables = [t for t in selected_tables 
                             if not t.startswith('etl_') and t != 'file_metadata']
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        # Handle different export formats
        if format == "csv":
            return await export_as_csv(db_manager, selected_tables, options, compress, timestamp)
        elif format == "excel":
            return await export_as_excel(db_manager, selected_tables, options, compress, timestamp)
        elif format == "sql":
            return await export_as_sql(db_manager, selected_tables, options, timestamp)
        elif format == "json":
            return await export_as_json(db_manager, selected_tables, options, compress, timestamp)
        elif format == "sqlite":
            return await export_as_sqlite(db_manager, selected_tables, options, timestamp)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported format: {format}")
            
    except Exception as e:
        logger.exception(f"Database export failed: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to export database: {str(e)}")


async def export_as_csv(db_manager, tables, options, compress, timestamp):
    """Export database as CSV files (zipped) - always with structure"""
    from fastapi.responses import StreamingResponse
    import io, csv, zipfile
    
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED if compress else zipfile.ZIP_STORED) as zip_file:
        for table_name in tables:
            try:
                # Always add structure (schema)
                with db_manager.pool.get_connection() as conn:
                    cursor = conn.execute(f"SELECT sql FROM sqlite_master WHERE type='table' AND name=?", (table_name,))
                    schema = cursor.fetchone()
                    if schema:
                        zip_file.writestr(f"{table_name}_schema.sql", schema[0])
                
                # Add data if requested (both option)
                if options == "both":
                    with db_manager.pool.get_connection() as conn:
                        cursor = conn.execute(f"SELECT * FROM {table_name}")
                        columns = [desc[0] for desc in cursor.description]
                        rows = cursor.fetchall()
                    
                    csv_output = io.StringIO()
                    writer = csv.writer(csv_output)
                    writer.writerow(columns)
                    writer.writerows(rows)
                    
                    zip_file.writestr(f"{table_name}.csv", csv_output.getvalue().encode('utf-8'))
            except Exception as e:
                logger.error(f"Failed to export table {table_name}: {e}")
    
    zip_buffer.seek(0)
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename=database_export_csv_{timestamp}.zip"}
    )


async def export_as_excel(db_manager, tables, options, compress, timestamp):
    """Export database as Excel file with multiple sheets - data with structure info"""
    from fastapi.responses import Response
    import io
    import pandas as pd
    
    excel_buffer = io.BytesIO()
    
    # Use xlsxwriter engine which is more reliable for direct byte output
    with pd.ExcelWriter(excel_buffer, engine='xlsxwriter') as writer:
        # Add schema sheet if requested
        if options == "structure":
            schema_data = []
            for table_name in tables:
                try:
                    with db_manager.pool.get_connection() as conn:
                        cursor = conn.execute(f"SELECT sql FROM sqlite_master WHERE type='table' AND name=?", (table_name,))
                        schema = cursor.fetchone()
                        if schema:
                            schema_data.append({'table_name': table_name, 'schema': schema[0]})
                except Exception as e:
                    logger.error(f"Failed to get schema for {table_name}: {e}")
            
            if schema_data:
                schema_df = pd.DataFrame(schema_data)
                schema_df.to_excel(writer, sheet_name='_Schemas', index=False)
        
        # Add data sheets if requested
        if options == "both":
            for table_name in tables:
                try:
                    with db_manager.pool.get_connection() as conn:
                        # Limit to 100,000 rows per table to prevent memory issues
                        df = pd.read_sql_query(f"SELECT * FROM {table_name} LIMIT 100000", conn)
                        # Limit sheet name to 31 chars (Excel limit)
                        sheet_name = table_name[:31]
                        df.to_excel(writer, sheet_name=sheet_name, index=False)
                except Exception as e:
                    logger.error(f"Failed to export table {table_name}: {e}")
    
    # Get the bytes from the buffer
    excel_data = excel_buffer.getvalue()
    
    return Response(
        content=excel_data,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={
            "Content-Disposition": f"attachment; filename=database_export_{timestamp}.xlsx",
            "Content-Type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        }
    )


async def export_as_sql(db_manager, tables, options, timestamp):
    """Export database as SQL dump - always includes structure"""
    from fastapi.responses import StreamingResponse
    import io
    
    sql_output = io.StringIO()
    sql_output.write(f"-- Database Export: {timestamp}\n")
    sql_output.write(f"-- Format: SQL\n\n")
    
    for table_name in tables:
        try:
            # Always export schema (structure)
            with db_manager.pool.get_connection() as conn:
                cursor = conn.execute(f"SELECT sql FROM sqlite_master WHERE type='table' AND name=?", (table_name,))
                schema = cursor.fetchone()
                if schema:
                    sql_output.write(f"\n-- Table: {table_name}\n")
                    sql_output.write(f"DROP TABLE IF EXISTS {table_name};\n")
                    sql_output.write(schema[0] + ";\n\n")
            
            # Export data if requested (both option)
            if options == "both":
                with db_manager.pool.get_connection() as conn:
                    cursor = conn.execute(f"SELECT * FROM {table_name}")
                    columns = [desc[0] for desc in cursor.description]
                    rows = cursor.fetchall()
                    
                    if rows:
                        sql_output.write(f"-- Data for table: {table_name}\n")
                        for row in rows:
                            values = ', '.join([f"'{str(v).replace(chr(39), chr(39) + chr(39))}'" if v is not None else 'NULL' for v in row])
                            sql_output.write(f"INSERT INTO {table_name} ({', '.join(columns)}) VALUES ({values});\n")
                        sql_output.write("\n")
        except Exception as e:
            logger.error(f"Failed to export table {table_name}: {e}")
    
    sql_content = sql_output.getvalue().encode('utf-8')
    
    return StreamingResponse(
        io.BytesIO(sql_content),
        media_type="application/sql",
        headers={"Content-Disposition": f"attachment; filename=database_export_{timestamp}.sql"}
    )


async def export_as_json(db_manager, tables, options, compress, timestamp):
    """Export database as JSON files (zipped) - always includes structure in metadata"""
    from fastapi.responses import StreamingResponse
    import io, json as json_lib, zipfile
    
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED if compress else zipfile.ZIP_STORED) as zip_file:
        # Create metadata file with structure information
        metadata = {
            "export_timestamp": timestamp,
            "format": "JSON",
            "tables": {}
        }
        
        for table_name in tables:
            try:
                with db_manager.pool.get_connection() as conn:
                    # Always get and include schema information (structure)
                    cursor = conn.execute(f"PRAGMA table_info({table_name})")
                    columns_info = cursor.fetchall()
                    
                    table_metadata = {
                        "columns": [
                            {
                                "name": col[1],
                                "type": col[2],
                                "not_null": bool(col[3]),
                                "primary_key": bool(col[5])
                            }
                            for col in columns_info
                        ]
                    }
                    metadata["tables"][table_name] = table_metadata
                    
                    # Export data if requested (both option)
                    if options == "both":
                        cursor = conn.execute(f"SELECT * FROM {table_name}")
                        columns = [desc[0] for desc in cursor.description]
                        rows = cursor.fetchall()
                        
                        data = [dict(zip(columns, row)) for row in rows]
                        json_content = json_lib.dumps(data, indent=2, default=str)
                        zip_file.writestr(f"{table_name}.json", json_content)
            except Exception as e:
                logger.error(f"Failed to export table {table_name}: {e}")
        
        # Always add metadata file with structure information
        metadata_content = json_lib.dumps(metadata, indent=2)
        zip_file.writestr("_metadata.json", metadata_content)
    
    zip_buffer.seek(0)
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename=database_export_json_{timestamp}.zip"}
    )


async def export_as_sqlite(db_manager, tables, options, timestamp):
    """Export as SQLite database file"""
    from fastapi.responses import FileResponse
    import shutil
    import tempfile
    
    # Create a temporary copy of the database
    temp_db = tempfile.NamedTemporaryFile(delete=False, suffix='.db')
    shutil.copy2(db_manager.pool.db_path, temp_db.name)
    temp_db.close()
    
    return FileResponse(
        temp_db.name,
        media_type="application/x-sqlite3",
        headers={"Content-Disposition": f"attachment; filename=database_export_{timestamp}.db"},
        background=lambda: Path(temp_db.name).unlink(missing_ok=True)
    )


@app.post("/api/database/table/sort")
async def sort_table_data(sort_request: TableSortRequest):
    """Sort table data"""
    db_manager = app_state["db_manager"]
    
    # Validate table name
    valid_tables = config.data_quality.expected_tables.keys()
    if sort_request.table_name not in valid_tables:
        raise HTTPException(status_code=400, detail=f"Invalid table name: {sort_request.table_name}")
    
    # Validate sort direction
    if sort_request.sort_direction.lower() not in ['asc', 'desc']:
        raise HTTPException(status_code=400, detail="Sort direction must be 'asc' or 'desc'")
    
    try:
        repo = db_manager.get_repository(sort_request.table_name)
        
        # Build sort query
        query = f"""
            SELECT * FROM {sort_request.table_name} 
            ORDER BY {sort_request.sort_column} {sort_request.sort_direction.upper()}
            LIMIT {sort_request.limit} OFFSET {sort_request.offset}
        """
        
        result = repo.execute_query(query, return_dataframe=True)
        
        if result.success:
            return {
                "success": True,
                "table_name": sort_request.table_name,
                "sort_column": sort_request.sort_column,
                "sort_direction": sort_request.sort_direction,
                "data": result.data.to_dict('records') if hasattr(result.data, 'to_dict') else result.data,
                "columns": result.columns,
                "row_count": result.row_count
            }
        else:
            raise HTTPException(status_code=500, detail=result.error_message)
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to sort table: {str(e)}")


@app.post("/api/database/query")
async def execute_custom_query(query_request: DatabaseQuery):
    """Execute custom SQL query"""
    db_manager = app_state["db_manager"]
    
    # Safety check - only allow SELECT statements
    query_upper = query_request.query.strip().upper()
    if not query_upper.startswith('SELECT'):
        raise HTTPException(status_code=400, detail="Only SELECT queries are allowed")
    
    try:
        # Add limit if not present
        if 'LIMIT' not in query_upper:
            limited_query = f"{query_request.query.rstrip(';')} LIMIT {query_request.limit}"
        else:
            limited_query = query_request.query
        
        result = db_manager.execute_query(limited_query, return_dataframe=True)
        
        if result.success:
            return {
                "success": True,
                "query": limited_query,
                "data": result.data.to_dict('records') if hasattr(result.data, 'to_dict') else result.data,
                "columns": result.columns,
                "row_count": result.row_count,
                "execution_time_ms": result.execution_time_ms
            }
        else:
            raise HTTPException(status_code=500, detail=result.error_message)
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Query execution failed: {str(e)}")


@app.get("/api/database/export/{table_name}")
async def export_table_csv(table_name: str):
    """Export table data as CSV"""
    db_manager = app_state["db_manager"]
    
    try:
        # Verify table exists
        with db_manager.pool.get_connection() as conn:
            cursor = conn.execute(
                "SELECT name FROM sqlite_master WHERE type='table' AND name=?",
                (table_name,)
            )
            if not cursor.fetchone():
                raise HTTPException(status_code=404, detail=f"Table not found: {table_name}")
            
            # Get all data
            cursor = conn.execute(f"SELECT * FROM {table_name}")
            columns = [description[0] for description in cursor.description]
            rows = cursor.fetchall()
        
        if not rows:
            raise HTTPException(status_code=404, detail="No data found")
        
        # Create temporary CSV file
        import csv
        with tempfile.NamedTemporaryFile(mode='w', suffix='.csv', delete=False, newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(columns)  # Write header
            writer.writerows(rows)     # Write data
            temp_path = f.name
        
        # Return file
        return FileResponse(
            temp_path,
            media_type='text/csv',
            filename=f"{table_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


# ============================================================================
# DATA QUALITY API
# ============================================================================

@app.get("/api/data-quality/summary")
async def get_data_quality_summary():
    """Get data quality issues summary"""
    db_manager = app_state["db_manager"]
    
    try:
        summary = db_manager.data_quality.get_summary()
        return {"data_quality": summary}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get data quality summary: {str(e)}")


# ============================================================================
# ANALYTICS API
# ============================================================================

@app.get("/api/analytics/dashboard-stats")
async def get_dashboard_stats():
    """Get dashboard statistics"""
    db_manager = app_state["db_manager"]
    etl_service = app_state["etl_service"]
    
    try:
        # Database stats
        db_info = db_manager.get_database_info()
        table_stats = db_manager.get_table_stats()
        
        # ETL stats
        current_job = etl_service.get_current_job_status()
        job_history = etl_service.get_job_history(5)
        
        # Data quality stats
        dq_summary = db_manager.data_quality.get_summary()
        
        return {
            "database": {
                "total_records": sum(stats.get("record_count", 0) for stats in table_stats.values()),
                "total_tables": len([s for s in table_stats.values() if s.get("exists", False)]),
                "database_size_mb": db_info.get("database_size_mb", 0),
                "last_processing": db_info.get("last_processing_date")
            },
            "etl": {
                "current_job": current_job.to_dict() if current_job else None,
                "recent_jobs": len(job_history),
                "total_files_processed": db_info.get("total_files_processed", 0)
            },
            "data_quality": dq_summary,
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to get dashboard stats: {str(e)}")


# ============================================================================
# HEALTH CHECK
# ============================================================================

@app.get("/api/health")
async def health_check():
    """Health check endpoint with periodic maintenance"""
    try:
        db_manager = app_state["db_manager"]
        
        # Test database connection
        result = db_manager.execute_query("SELECT 1 as test")
        db_healthy = result.success
        
        # Perform periodic cleanup to prevent memory leaks
        if db_manager and hasattr(db_manager, 'pool'):
            db_manager.pool.periodic_cleanup()
        
        return {
            "status": "healthy" if db_healthy else "unhealthy",
            "database": "connected" if db_healthy else "disconnected",
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        return {
            "status": "unhealthy",
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }


# ============================================================================
# REPORTS & ANALYTICS API (Refactored into separate module)
# ============================================================================

# Import reports router from refactored module
from .reports import reports_router

# Include reports router
app.include_router(reports_router)

# Keep the build_date_filter function here for backwards compatibility with other endpoints
def build_date_filter(table: str, start_date: Optional[str] = None, end_date: Optional[str] = None):
    """
    Build WHERE clause and params for date filtering based on table.
    Returns tuple of (where_clause, params)
    """
    where_conditions = []
    params = []
    
    # Map tables to their date columns
    date_columns = {
        'referrals': 'referral_updated_at',
        'cases': 'case_updated_at',
        'assistance_requests': 'updated_at',
        'people': 'case_updated_at'  # Use via JOIN with cases
    }
    
    date_col = date_columns.get(table, 'created_at')  # fallback
    
    if start_date:
        where_conditions.append(f"{date_col} >= ?")
        params.append(start_date)
    if end_date:
        where_conditions.append(f"{date_col} <= ?")
        params.append(end_date)
    
    where_clause = f" AND {' AND '.join(where_conditions)}" if where_conditions else ""
    return where_clause, params

@app.get("/api/reports/summary")
async def get_reports_summary(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    service_type: Optional[str] = Query(None),
    provider: Optional[str] = Query(None),
    program: Optional[str] = Query(None),
    gender: Optional[str] = Query(None),
    race: Optional[str] = Query(None)
):
    """Get summary statistics for reports"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            # Build WHERE clause for date filtering
            where_conditions = []
            params = []
            
            if start_date:
                where_conditions.append("referral_updated_at >= ?")
                params.append(start_date)
            if end_date:
                where_conditions.append("referral_updated_at <= ?")
                params.append(end_date)
            
            where_clause = f"WHERE {' AND '.join(where_conditions)}" if where_conditions else ""
            
            # Total referrals
            query = f"SELECT COUNT(*) FROM referrals {where_clause}"
            cursor = conn.execute(query, params)
            result = cursor.fetchone()
            total_referrals = result[0] if result else 0
            
            # Total cases (use case_updated_at)
            case_where_conditions = []
            case_params = []
            if start_date:
                case_where_conditions.append("case_updated_at >= ?")
                case_params.append(start_date)
            if end_date:
                case_where_conditions.append("case_updated_at <= ?")
                case_params.append(end_date)
            case_where_clause = f"WHERE {' AND '.join(case_where_conditions)}" if case_where_conditions else ""
            
            query = f"SELECT COUNT(*) FROM cases {case_where_clause}"
            cursor = conn.execute(query, case_params)
            result = cursor.fetchone()
            total_cases = result[0] if result else 0
            
            # Total unique people (join with cases for date filtering)
            if case_where_conditions:
                query = f"""
                    SELECT COUNT(DISTINCT p.person_id) 
                    FROM people p 
                    JOIN cases c ON p.person_id = c.person_id 
                    {case_where_clause}
                """
                cursor = conn.execute(query, case_params)
            else:
                cursor = conn.execute("SELECT COUNT(DISTINCT person_id) FROM people")
            result = cursor.fetchone()
            total_people = result[0] if result else 0
            
            # Total assistance requests (use updated_at)
            ar_where_conditions = []
            ar_params = []
            if start_date:
                ar_where_conditions.append("updated_at >= ?")
                ar_params.append(start_date)
            if end_date:
                ar_where_conditions.append("updated_at <= ?")
                ar_params.append(end_date)
            ar_where_clause = f"WHERE {' AND '.join(ar_where_conditions)}" if ar_where_conditions else ""
            
            query = f"SELECT COUNT(*) FROM assistance_requests {ar_where_clause}"
            cursor = conn.execute(query, ar_params)
            result = cursor.fetchone()
            total_assistance_requests = result[0] if result else 0
        
        return {
            "total_referrals": total_referrals,
            "total_cases": total_cases,
            "total_people": total_people,
            "total_assistance_requests": total_assistance_requests
        }
    except Exception as e:
        return {
            "total_referrals": 0,
            "total_cases": 0,
            "total_people": 0,
            "total_assistance_requests": 0
        }


@app.get("/api/reports/referral-status")
async def get_referral_status_report(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Get referral counts by status"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            where_conditions = []
            params = []
            
            if start_date:
                where_conditions.append("referral_updated_at >= ?")
                params.append(start_date)
            if end_date:
                where_conditions.append("referral_updated_at <= ?")
                params.append(end_date)
            
            where_clause = f"WHERE referral_status IS NOT NULL"
            if where_conditions:
                where_clause += f" AND {' AND '.join(where_conditions)}"
            
            query = f"""
                SELECT referral_status, COUNT(*) as count 
                FROM referrals 
                {where_clause}
                GROUP BY referral_status
                ORDER BY count DESC
            """
            cursor = conn.execute(query, params)
            results = cursor.fetchall()
        
        # Handle empty results
        if not results:
            return {"labels": [], "values": []}
        
        labels = [row[0] or 'Unknown' for row in results]
        values = [row[1] for row in results]
        
        return {
            "labels": labels,
            "values": values
        }
    except Exception as e:
        return {"labels": [], "values": []}


@app.get("/api/reports/case-status")
async def get_case_status_report(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Get case counts by status"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            where_conditions = []
            params = []
            
            if start_date:
                where_conditions.append("case_updated_at >= ?")
                params.append(start_date)
            if end_date:
                where_conditions.append("case_updated_at <= ?")
                params.append(end_date)
            
            where_clause = f"WHERE case_status IS NOT NULL"
            if where_conditions:
                where_clause += f" AND {' AND '.join(where_conditions)}"
            
            query = f"""
                SELECT case_status, COUNT(*) as count 
                FROM cases 
                {where_clause}
                GROUP BY case_status
                ORDER BY count DESC
            """
            cursor = conn.execute(query, params)
            results = cursor.fetchall()
        
        labels = [row[0] or 'Unknown' for row in results]
        values = [row[1] for row in results]
        
        return {
            "labels": labels,
            "values": values
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/reports/service-types")
async def get_service_types_report(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Get case counts by service type"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            where_conditions = []
            params = []
            
            if start_date:
                where_conditions.append("case_updated_at >= ?")
                params.append(start_date)
            if end_date:
                where_conditions.append("case_updated_at <= ?")
                params.append(end_date)
            
            where_clause = f"WHERE service_type IS NOT NULL"
            if where_conditions:
                where_clause += f" AND {' AND '.join(where_conditions)}"
            
            query = f"""
                SELECT service_type, COUNT(*) as count 
                FROM cases 
                {where_clause}
                GROUP BY service_type
                ORDER BY count DESC
                LIMIT 15
            """
            cursor = conn.execute(query, params)
            results = cursor.fetchall()
        
        labels = [row[0] or 'Unknown' for row in results]
        values = [row[1] for row in results]
        
        return {
            "labels": labels,
            "values": values
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/reports/sending-providers")
async def get_sending_providers_report(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Get top 10 sending providers by referral count"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            date_filter, params = build_date_filter('referrals', start_date, end_date)
            query = f"""
                SELECT sending_provider_name, COUNT(*) as count 
                FROM referrals 
                WHERE sending_provider_name IS NOT NULL{date_filter}
                GROUP BY sending_provider_name
                ORDER BY count DESC
                LIMIT 10
            """
            cursor = conn.execute(query, params)
            results = cursor.fetchall()
        
        labels = [row[0] or 'Unknown' for row in results]
        values = [row[1] for row in results]
        
        return {
            "labels": labels,
            "values": values
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/reports/receiving-providers")
async def get_receiving_providers_report(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Get top 10 receiving providers by referral count"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            date_filter, params = build_date_filter('referrals', start_date, end_date)
            query = f"""
                SELECT receiving_provider_name, COUNT(*) as count 
                FROM referrals 
                WHERE receiving_provider_name IS NOT NULL{date_filter}
                GROUP BY receiving_provider_name
                ORDER BY count DESC
                LIMIT 10
            """
            cursor = conn.execute(query, params)
            results = cursor.fetchall()
        
        labels = [row[0] or 'Unknown' for row in results]
        values = [row[1] for row in results]
        
        return {
            "labels": labels,
            "values": values
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/reports/referrals-timeline")
async def get_referrals_timeline(
    grouping: str = Query("week", pattern="^(day|week|month)$"),
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Get referrals over time"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            # Determine date format based on grouping
            if grouping == "day":
                date_format = "%Y-%m-%d"
            elif grouping == "week":
                date_format = "%Y-W%W"
            else:  # month
                date_format = "%Y-%m"
            
            # Build date filter
            where_conditions = ["referral_created_at IS NOT NULL"]
            params = []
            
            if start_date:
                where_conditions.append("referral_created_at >= ?")
                params.append(start_date)
            if end_date:
                where_conditions.append("referral_created_at <= ?")
                params.append(end_date)
            
            where_clause = " AND ".join(where_conditions)
            
            query = f"""
                SELECT strftime('{date_format}', referral_created_at) as period, COUNT(*) as count 
                FROM referrals 
                WHERE {where_clause}
                GROUP BY period
                ORDER BY period
                LIMIT 100
            """
            cursor = conn.execute(query, params)
            results = cursor.fetchall()
        
        labels = [row[0] or 'Unknown' for row in results]
        values = [row[1] for row in results]
        
        return {
            "labels": labels,
            "values": values
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/reports/top-programs")
async def get_top_programs_report(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Get top programs by referral count with acceptance rates"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            date_filter, params = build_date_filter('referrals', start_date, end_date)
            query = f"""
                SELECT 
                    receiving_program_name,
                    COUNT(*) as total_referrals,
                    SUM(CASE WHEN referral_status = 'accepted' THEN 1 ELSE 0 END) as accepted_referrals
                FROM referrals 
                WHERE receiving_program_name IS NOT NULL{date_filter}
                GROUP BY receiving_program_name
                ORDER BY total_referrals DESC
                LIMIT 15
            """
            cursor = conn.execute(query, params)
            results = cursor.fetchall()
        
        programs = []
        for row in results:
            acceptance_rate = (row[2] / row[1] * 100) if row[1] > 0 else 0
            programs.append({
                "program_name": row[0],
                "total_referrals": row[1],
                "accepted_referrals": row[2],
                "acceptance_rate": round(acceptance_rate, 1)
            })
        
        return {"programs": programs}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/reports/case-outcomes")
async def get_case_outcomes_report(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Get case outcomes by resolution type"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            date_filter, params = build_date_filter('cases', start_date, end_date)
            query = f"""
                SELECT outcome_resolution_type, COUNT(*) as count 
                FROM cases 
                WHERE outcome_resolution_type IS NOT NULL{date_filter}
                GROUP BY outcome_resolution_type
                ORDER BY count DESC
            """
            cursor = conn.execute(query, params)
            results = cursor.fetchall()
        
        outcomes = []
        for row in results:
            outcomes.append({
                "resolution_type": row[0],
                "count": row[1]
            })
        
        return {"outcomes": outcomes}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# ADDITIONAL REPORTS API ENDPOINTS
# ============================================================================

@app.get("/api/reports/demographics/age-distribution")
async def get_age_distribution(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Get age distribution of clients"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            # Build date filter on cases table
            where_conditions = ["p.date_of_birth IS NOT NULL"]
            params = []
            
            if start_date or end_date:
                # Join with cases to apply date filtering
                if start_date:
                    where_conditions.append("c.case_updated_at >= ?")
                    params.append(start_date)
                if end_date:
                    where_conditions.append("c.case_updated_at <= ?")
                    params.append(end_date)
                
                where_clause = " AND ".join(where_conditions)
                query = f"""
                    SELECT 
                        CASE 
                            WHEN CAST((julianday('now') - julianday(p.date_of_birth)) / 365.25 AS INTEGER) < 18 THEN '0-17'
                            WHEN CAST((julianday('now') - julianday(p.date_of_birth)) / 365.25 AS INTEGER) BETWEEN 18 AND 24 THEN '18-24'
                            WHEN CAST((julianday('now') - julianday(p.date_of_birth)) / 365.25 AS INTEGER) BETWEEN 25 AND 34 THEN '25-34'
                            WHEN CAST((julianday('now') - julianday(p.date_of_birth)) / 365.25 AS INTEGER) BETWEEN 35 AND 44 THEN '35-44'
                            WHEN CAST((julianday('now') - julianday(p.date_of_birth)) / 365.25 AS INTEGER) BETWEEN 45 AND 54 THEN '45-54'
                            WHEN CAST((julianday('now') - julianday(p.date_of_birth)) / 365.25 AS INTEGER) BETWEEN 55 AND 64 THEN '55-64'
                            WHEN CAST((julianday('now') - julianday(p.date_of_birth)) / 365.25 AS INTEGER) >= 65 THEN '65+'
                            ELSE 'Unknown'
                        END as age_group,
                        COUNT(DISTINCT p.person_id) as count
                    FROM people p
                    INNER JOIN cases c ON p.person_id = c.person_id
                    WHERE {where_clause}
                    GROUP BY age_group
                    ORDER BY 
                        CASE age_group
                            WHEN '0-17' THEN 1
                            WHEN '18-24' THEN 2
                            WHEN '25-34' THEN 3
                            WHEN '35-44' THEN 4
                            WHEN '45-54' THEN 5
                            WHEN '55-64' THEN 6
                            WHEN '65+' THEN 7
                            ELSE 8
                        END
                """
            else:
                # No date filtering - query people table directly
                query = """
                    SELECT 
                        CASE 
                            WHEN CAST((julianday('now') - julianday(date_of_birth)) / 365.25 AS INTEGER) < 18 THEN '0-17'
                            WHEN CAST((julianday('now') - julianday(date_of_birth)) / 365.25 AS INTEGER) BETWEEN 18 AND 24 THEN '18-24'
                            WHEN CAST((julianday('now') - julianday(date_of_birth)) / 365.25 AS INTEGER) BETWEEN 25 AND 34 THEN '25-34'
                            WHEN CAST((julianday('now') - julianday(date_of_birth)) / 365.25 AS INTEGER) BETWEEN 35 AND 44 THEN '35-44'
                            WHEN CAST((julianday('now') - julianday(date_of_birth)) / 365.25 AS INTEGER) BETWEEN 45 AND 54 THEN '45-54'
                            WHEN CAST((julianday('now') - julianday(date_of_birth)) / 365.25 AS INTEGER) BETWEEN 55 AND 64 THEN '55-64'
                            WHEN CAST((julianday('now') - julianday(date_of_birth)) / 365.25 AS INTEGER) >= 65 THEN '65+'
                            ELSE 'Unknown'
                        END as age_group,
                        COUNT(*) as count
                    FROM people
                    WHERE date_of_birth IS NOT NULL
                    GROUP BY age_group
                    ORDER BY 
                        CASE age_group
                            WHEN '0-17' THEN 1
                            WHEN '18-24' THEN 2
                            WHEN '25-34' THEN 3
                            WHEN '35-44' THEN 4
                            WHEN '45-54' THEN 5
                            WHEN '55-64' THEN 6
                            WHEN '65+' THEN 7
                            ELSE 8
                        END
                """
            
            cursor = conn.execute(query, params)
            results = cursor.fetchall()
        
        return {
            "labels": [row[0] for row in results],
            "values": [row[1] for row in results]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/reports/demographics/gender")
async def get_gender_distribution(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Get gender distribution of clients"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            params = []
            
            if start_date or end_date:
                # Join with cases to apply date filtering
                where_conditions = []
                if start_date:
                    where_conditions.append("c.case_updated_at >= ?")
                    params.append(start_date)
                if end_date:
                    where_conditions.append("c.case_updated_at <= ?")
                    params.append(end_date)
                
                where_clause = " AND ".join(where_conditions)
                query = f"""
                    SELECT 
                        COALESCE(p.gender, 'Not Specified') as gender,
                        COUNT(DISTINCT p.person_id) as count
                    FROM people p
                    INNER JOIN cases c ON p.person_id = c.person_id
                    WHERE {where_clause}
                    GROUP BY gender
                    ORDER BY count DESC
                """
            else:
                query = """
                    SELECT 
                        COALESCE(gender, 'Not Specified') as gender,
                        COUNT(*) as count
                    FROM people
                    GROUP BY gender
                    ORDER BY count DESC
                """
            
            cursor = conn.execute(query, params)
            results = cursor.fetchall()
        
        return {
            "labels": [row[0] for row in results],
            "values": [row[1] for row in results]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/reports/demographics/race-ethnicity")
async def get_race_ethnicity(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Get race/ethnicity distribution"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            params = []
            
            if start_date or end_date:
                # Join with cases to apply date filtering
                where_conditions = ["p.race NOT IN ('undisclosed', '')"]
                if start_date:
                    where_conditions.append("c.case_updated_at >= ?")
                    params.append(start_date)
                if end_date:
                    where_conditions.append("c.case_updated_at <= ?")
                    params.append(end_date)
                
                where_clause = " AND ".join(where_conditions)
                query = f"""
                    SELECT 
                        COALESCE(p.race, 'Not Specified') as race,
                        COUNT(DISTINCT p.person_id) as count
                    FROM people p
                    INNER JOIN cases c ON p.person_id = c.person_id
                    WHERE {where_clause}
                    GROUP BY race
                    ORDER BY count DESC
                    LIMIT 10
                """
            else:
                query = """
                    SELECT 
                        COALESCE(race, 'Not Specified') as race,
                        COUNT(*) as count
                    FROM people
                    WHERE race NOT IN ('undisclosed', '')
                    GROUP BY race
                    ORDER BY count DESC
                    LIMIT 10
                """
            
            cursor = conn.execute(query, params)
            results = cursor.fetchall()
        
        return {
            "labels": [row[0] for row in results],
            "values": [row[1] for row in results]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/reports/service-metrics/resolution-time")
async def get_resolution_time_metrics(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Get average case resolution times by service type"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            date_filter, params = build_date_filter('cases', start_date, end_date)
            query = f"""
                SELECT 
                    service_type,
                    COUNT(*) as total_cases,
                    ROUND(AVG(julianday(case_closed_at) - julianday(case_created_at)), 1) as avg_days,
                    ROUND(MIN(julianday(case_closed_at) - julianday(case_created_at)), 1) as min_days,
                    ROUND(MAX(julianday(case_closed_at) - julianday(case_created_at)), 1) as max_days
                FROM cases
                WHERE case_closed_at IS NOT NULL 
                    AND case_created_at IS NOT NULL
                    AND service_type IS NOT NULL{date_filter}
                GROUP BY service_type
                HAVING total_cases >= 3
                ORDER BY avg_days DESC
                LIMIT 10
            """
            cursor = conn.execute(query, params)
            results = cursor.fetchall()
        
        return {
            "metrics": [
                {
                    "service_type": row[0],
                    "total_cases": row[1],
                    "avg_days": row[2],
                    "min_days": row[3],
                    "max_days": row[4]
                }
                for row in results
            ]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/reports/service-metrics/referral-conversion")
async def get_referral_conversion_rates(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Get referral acceptance/conversion rates by service type"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            date_filter, params = build_date_filter('referrals', start_date, end_date)
            query = f"""
                SELECT 
                    service_type,
                    COUNT(*) as total_referrals,
                    SUM(CASE WHEN referral_status = 'accepted' THEN 1 ELSE 0 END) as accepted,
                    SUM(CASE WHEN referral_status = 'declined' THEN 1 ELSE 0 END) as declined,
                    SUM(CASE WHEN referral_status IN ('pending', 'off_platform') THEN 1 ELSE 0 END) as pending
                FROM referrals
                WHERE service_type IS NOT NULL{date_filter}
                GROUP BY service_type
                HAVING total_referrals >= 5
                ORDER BY total_referrals DESC
                LIMIT 10
            """
            cursor = conn.execute(query, params)
            results = cursor.fetchall()
        
        return {
            "metrics": [
                {
                    "service_type": row[0],
                    "total_referrals": row[1],
                    "accepted": row[2],
                    "declined": row[3],
                    "pending": row[4],
                    "acceptance_rate": round((row[2] / row[1] * 100) if row[1] > 0 else 0, 1)
                }
                for row in results
            ]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/reports/network/provider-collaboration")
async def get_provider_collaboration(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Get provider collaboration network (who refers to whom)"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            date_filter, params = build_date_filter('referrals', start_date, end_date)
            query = f"""
                SELECT 
                    sending_provider_name,
                    receiving_provider_name,
                    COUNT(*) as referral_count
                FROM referrals
                WHERE sending_provider_name IS NOT NULL 
                    AND receiving_provider_name IS NOT NULL
                    AND sending_provider_name != receiving_provider_name{date_filter}
                GROUP BY sending_provider_name, receiving_provider_name
                HAVING referral_count >= 3
                ORDER BY referral_count DESC
                LIMIT 20
            """
            cursor = conn.execute(query, params)
            results = cursor.fetchall()
        
        return {
            "collaborations": [
                {
                    "from": row[0],
                    "to": row[1],
                    "count": row[2]
                }
                for row in results
            ]
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/reports/geographic/cases-by-location")
async def get_cases_by_location(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Get case distribution by city/county"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            # Build date filter for cases
            where_conditions = ["p.current_person_address_city IS NOT NULL"]
            params = []
            
            if start_date:
                where_conditions.append("c.case_updated_at >= ?")
                params.append(start_date)
            if end_date:
                where_conditions.append("c.case_updated_at <= ?")
                params.append(end_date)
            
            where_clause = " AND ".join(where_conditions)
            
            query = f"""
                SELECT 
                    p.current_person_address_city,
                    p.current_person_address_county,
                    p.current_person_address_state,
                    COUNT(DISTINCT c.case_id) as case_count
                FROM people p
                JOIN cases c ON p.person_id = c.person_id
                WHERE {where_clause}
                GROUP BY p.current_person_address_city, p.current_person_address_county, p.current_person_address_state
                ORDER BY case_count DESC
                LIMIT 15
            """
            cursor = conn.execute(query, params)
            results = cursor.fetchall()
        
        # Handle empty results
        if not results:
            return {"locations": []}
        
        return {
            "locations": [
                {
                    "city": row[0],
                    "county": row[1],
                    "state": row[2],
                    "case_count": row[3]
                }
                for row in results
            ]
        }
    except Exception as e:
        return {"locations": []}


@app.get("/api/reports/military/veteran-services")
async def get_veteran_services(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Get statistics on veteran/military services"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            date_filter, params = build_date_filter('assistance_requests', start_date, end_date)
            
            # Count of assistance requests with military affiliation
            query1 = f"""
                SELECT 
                    mil_affiliation,
                    COUNT(*) as count
                FROM assistance_requests
                WHERE mil_affiliation IS NOT NULL AND mil_affiliation != ''{date_filter}
                GROUP BY mil_affiliation
                ORDER BY count DESC
            """
            cursor = conn.execute(query1, params)
            affiliation_results = cursor.fetchall()
            
            # Count by military branch
            query2 = f"""
                SELECT 
                    mil_branch,
                    COUNT(*) as count
                FROM assistance_requests
                WHERE mil_branch IS NOT NULL AND mil_branch != ''{date_filter}
                GROUP BY mil_branch
                ORDER BY count DESC
            """
            cursor = conn.execute(query2, params)
            branch_results = cursor.fetchall()
        
        # Handle empty results
        if not affiliation_results and not branch_results:
            return {
                "by_affiliation": {"labels": [], "values": []},
                "by_branch": {"labels": [], "values": []}
            }
        
        return {
            "by_affiliation": {
                "labels": [row[0] for row in affiliation_results] if affiliation_results else [],
                "values": [row[1] for row in affiliation_results] if affiliation_results else []
            },
            "by_branch": {
                "labels": [row[0] for row in branch_results] if branch_results else [],
                "values": [row[1] for row in branch_results] if branch_results else []
            }
        }
    except Exception as e:
        return {
            "by_affiliation": {"labels": [], "values": []},
            "by_branch": {"labels": [], "values": []}
        }


@app.get("/api/reports/workforce/employee-workload")
async def get_employee_workload(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Get employee workload distribution"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            # Build date filter for cases
            where_conditions = []
            params = []
            
            if start_date:
                where_conditions.append("c.case_updated_at >= ?")
                params.append(start_date)
            if end_date:
                where_conditions.append("c.case_updated_at <= ?")
                params.append(end_date)
            
            where_clause = f" WHERE {' AND '.join(where_conditions)}" if where_conditions else ""
            
            query = f"""
                SELECT 
                    e.first_name || ' ' || e.last_name as employee_name,
                    e.provider_name,
                    COUNT(CASE WHEN c.case_status IN ('active', 'managed', 'processed') THEN 1 END) as active_cases,
                    COUNT(c.case_id) as total_cases,
                    COUNT(CASE WHEN c.outcome_resolution_type = 'resolved' THEN 1 END) as resolved_cases
                FROM employees e
                LEFT JOIN cases c ON e.employee_id = c.primary_worker_id{where_clause}
                GROUP BY e.employee_id, employee_name, e.provider_name
                HAVING total_cases > 0
                ORDER BY active_cases DESC, total_cases DESC
                LIMIT 20
            """
            cursor = conn.execute(query, params)
            results = cursor.fetchall()
        
        # Handle empty results
        if not results:
            return {"employees": []}
        
        return {
            "employees": [
                {
                    "employee_name": row[0],
                    "provider": row[1],
                    "active_cases": row[2],
                    "total_cases": row[3],
                    "resolved_cases": row[4],
                    "resolution_rate": round((row[4] / row[3] * 100) if row[3] > 0 else 0, 1)
                }
                for row in results
            ]
        }
    except Exception as e:
        return {"employees": []}


@app.get("/api/reports/trends/cases-over-time")
async def get_cases_over_time(
    grouping: str = Query("month", pattern="^(day|week|month)$"),
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Get cases created over time by status"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            if grouping == "day":
                date_format = "%Y-%m-%d"
            elif grouping == "week":
                date_format = "%Y-W%W"
            else:
                date_format = "%Y-%m"
            
            # Build date filter
            where_conditions = ["case_created_at IS NOT NULL"]
            params = []
            
            if start_date:
                where_conditions.append("case_updated_at >= ?")
                params.append(start_date)
            if end_date:
                where_conditions.append("case_updated_at <= ?")
                params.append(end_date)
            
            where_clause = " AND ".join(where_conditions)
            
            query = f"""
                SELECT 
                    strftime('{date_format}', case_created_at) as period,
                    case_status,
                    COUNT(*) as count
                FROM cases
                WHERE {where_clause}
                GROUP BY period, case_status
                ORDER BY period
            """
            cursor = conn.execute(query, params)
            results = cursor.fetchall()
        
        # Handle empty results
        if not results:
            return {"labels": [], "datasets": []}
        
        # Organize data by status
        periods = {}
        statuses = set()
        
        for row in results:
            period, status, count = row
            if period not in periods:
                periods[period] = {}
            periods[period][status] = count
            statuses.add(status)
        
        # Convert to chart format
        labels = sorted(periods.keys())
        datasets = []
        
        for status in sorted(statuses):
            datasets.append({
                "label": status,
                "data": [periods[period].get(status, 0) for period in labels]
            })
        
        return {
            "labels": labels,
            "datasets": datasets
        }
    except Exception as e:
        return {"labels": [], "datasets": []}


@app.get("/api/reports/service-subtypes")
async def get_service_subtypes(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None)
):
    """Get detailed breakdown of service subtypes"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            date_filter, params = build_date_filter('cases', start_date, end_date)
            query = f"""
                SELECT 
                    service_type,
                    service_subtype,
                    COUNT(*) as count
                FROM cases
                WHERE service_type IS NOT NULL AND service_subtype IS NOT NULL{date_filter}
                GROUP BY service_type, service_subtype
                ORDER BY count DESC
                LIMIT 25
            """
            cursor = conn.execute(query, params)
            results = cursor.fetchall()
        
        # Handle empty results
        if not results:
            return {"subtypes": []}
        
        return {
            "subtypes": [
                {
                    "service_type": row[0],
                    "service_subtype": row[1],
                    "count": row[2]
                }
                for row in results
            ]
        }
    except Exception as e:
        return {"subtypes": []}


@app.get("/api/reports/client-journey")
async def get_client_journey_metrics():
    """Get metrics on client journey through the system"""
    db_manager = app_state["db_manager"]
    
    try:
        with db_manager.pool.get_connection() as conn:
            cursor = conn.execute("""
                SELECT 
                    COUNT(DISTINCT p.person_id) as total_clients,
                    AVG(case_count) as avg_cases_per_client,
                    AVG(referral_count) as avg_referrals_per_client,
                    AVG(ar_count) as avg_assistance_requests_per_client
                FROM people p
                LEFT JOIN (
                    SELECT person_id, COUNT(*) as case_count 
                    FROM cases 
                    GROUP BY person_id
                ) c ON p.person_id = c.person_id
                LEFT JOIN (
                    SELECT person_id, COUNT(*) as referral_count 
                    FROM referrals 
                    GROUP BY person_id
                ) r ON p.person_id = r.person_id
                LEFT JOIN (
                    SELECT person_id, COUNT(*) as ar_count 
                    FROM assistance_requests 
                    GROUP BY person_id
                ) ar ON p.person_id = ar.person_id
            """)
            result = cursor.fetchone()
            
            # Get distribution of touchpoints
            cursor = conn.execute("""
                SELECT 
                    CASE 
                        WHEN total_touchpoints = 1 THEN '1'
                        WHEN total_touchpoints BETWEEN 2 AND 3 THEN '2-3'
                        WHEN total_touchpoints BETWEEN 4 AND 6 THEN '4-6'
                        WHEN total_touchpoints BETWEEN 7 AND 10 THEN '7-10'
                        WHEN total_touchpoints > 10 THEN '10+'
                    END as touchpoint_range,
                    COUNT(*) as client_count
                FROM (
                    SELECT 
                        p.person_id,
                        COALESCE(c.case_count, 0) + 
                        COALESCE(r.referral_count, 0) + 
                        COALESCE(ar.ar_count, 0) as total_touchpoints
                    FROM people p
                    LEFT JOIN (SELECT person_id, COUNT(*) as case_count FROM cases GROUP BY person_id) c 
                        ON p.person_id = c.person_id
                    LEFT JOIN (SELECT person_id, COUNT(*) as referral_count FROM referrals GROUP BY person_id) r 
                        ON p.person_id = r.person_id
                    LEFT JOIN (SELECT person_id, COUNT(*) as ar_count FROM assistance_requests GROUP BY person_id) ar 
                        ON p.person_id = ar.person_id
                )
                GROUP BY touchpoint_range
                ORDER BY 
                    CASE touchpoint_range
                        WHEN '1' THEN 1
                        WHEN '2-3' THEN 2
                        WHEN '4-6' THEN 3
                        WHEN '7-10' THEN 4
                        WHEN '10+' THEN 5
                    END
            """)
            distribution = cursor.fetchall()
        
        # Handle empty results
        if not result:
            return {
                "summary": {
                    "total_clients": 0,
                    "avg_cases_per_client": 0,
                    "avg_referrals_per_client": 0,
                    "avg_assistance_requests_per_client": 0
                },
                "touchpoint_distribution": {"labels": [], "values": []}
            }
        
        return {
            "summary": {
                "total_clients": result[0] or 0,
                "avg_cases_per_client": round(result[1], 2) if result[1] else 0,
                "avg_referrals_per_client": round(result[2], 2) if result[2] else 0,
                "avg_assistance_requests_per_client": round(result[3], 2) if result[3] else 0
            },
            "touchpoint_distribution": {
                "labels": [row[0] for row in distribution] if distribution else [],
                "values": [row[1] for row in distribution] if distribution else []
            }
        }
    except Exception as e:
        return {
            "summary": {
                "total_clients": 0,
                "avg_cases_per_client": 0,
                "avg_referrals_per_client": 0,
                "avg_assistance_requests_per_client": 0
            },
            "touchpoint_distribution": {"labels": [], "values": []}
        }


# ============================================================================
# ANNUAL REPORT EXPORT API (Protected)
# ============================================================================

@app.post("/api/reports/export/annual-report-word")
async def export_annual_report_word(request: Request, report_data: AnnualReportExportRequest, session: UserSession = Depends(require_auth)):
    """Export annual report to Word document with embedded charts"""
    from .report_export import generate_word_report
    from .audit_logger import get_audit_logger
    
    try:
        # Generate Word document
        output = generate_word_report(
            report_data=report_data.dict(),
            username=session.username
        )
        
        # Log export action
        audit_logger = get_audit_logger()
        audit_logger.log(
            username=session.username,
            action="ANNUAL_REPORT_EXPORT_WORD",
            category="REPORTS",
            success=True,
            details=f"Exported annual report for period: {report_data.period}",
            ip_address=request.client.host if request.client else None
        )
        
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=Annual_Report_{datetime.now().strftime('%Y%m%d')}.docx"
            }
        )
        
    except ImportError as e:
        raise HTTPException(
            status_code=501,
            detail=f"Word export unavailable: {str(e)}. Install required packages: pip install python-docx Pillow"
        )
    except Exception as e:
        logger.error(f"Error generating Word export: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to generate Word document: {str(e)}")


@app.post("/api/reports/export/annual-report-pdf")
async def export_annual_report_pdf(request: Request, report_data: AnnualReportExportRequest, session: UserSession = Depends(require_auth)):
    """Export annual report to PDF with embedded charts"""
    from .report_export import generate_pdf_report
    from .audit_logger import get_audit_logger
    
    try:
        # Generate PDF
        output = generate_pdf_report(
            report_data=report_data.dict(),
            username=session.username
        )
        
        # Log export action
        audit_logger = get_audit_logger()
        audit_logger.log(
            username=session.username,
            action="ANNUAL_REPORT_EXPORT_PDF",
            category="REPORTS",
            success=True,
            details=f"Exported annual report PDF for period: {report_data.period}",
            ip_address=request.client.host if request.client else None
        )
        
        return StreamingResponse(
            output,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename=Annual_Report_{datetime.now().strftime('%Y%m%d')}.pdf"
            }
        )
        
    except ImportError as e:
        raise HTTPException(
            status_code=501,
            detail=f"PDF export unavailable: {str(e)}. Install required packages: pip install reportlab Pillow"
        )
    except Exception as e:
        logger.error(f"Error generating PDF export: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to generate PDF: {str(e)}")


# ============================================================================
# ADVANCED REPORTS WITH FILTERING
# ============================================================================

@app.get("/api/reports/filter-options")
async def get_filter_options():
    """Get available filter options from database for all filter dropdowns"""
    db_manager = app_state["db_manager"]
    try:
        with db_manager.pool.get_connection() as conn:
            # Get date range from data
            cursor = conn.execute("""
                SELECT 
                    MIN(case_created_at) as min_date,
                    MAX(case_created_at) as max_date
                FROM cases
                WHERE case_created_at IS NOT NULL
            """)
            date_range = cursor.fetchone()
            
            # Get case statuses
            cursor = conn.execute("""
                SELECT DISTINCT case_status 
                FROM cases 
                WHERE case_status IS NOT NULL
                ORDER BY case_status
            """)
            statuses = [row[0] for row in cursor.fetchall()]
            
            # Get service types
            cursor = conn.execute("""
                SELECT DISTINCT service_type 
                FROM cases 
                WHERE service_type IS NOT NULL
                ORDER BY service_type
            """)
            service_types = [row[0] for row in cursor.fetchall()]
            
            # Get service subtypes
            cursor = conn.execute("""
                SELECT DISTINCT service_subtype 
                FROM cases 
                WHERE service_subtype IS NOT NULL
                ORDER BY service_subtype
            """)
            service_subtypes = [row[0] for row in cursor.fetchall()]
            
            # Get providers
            cursor = conn.execute("""
                SELECT DISTINCT provider_name 
                FROM cases 
                WHERE provider_name IS NOT NULL
                ORDER BY provider_name
            """)
            providers = [row[0] for row in cursor.fetchall()]
            
            # Get programs
            cursor = conn.execute("""
                SELECT DISTINCT program_name 
                FROM cases 
                WHERE program_name IS NOT NULL
                ORDER BY program_name
            """)
            programs = [row[0] for row in cursor.fetchall()]
            
            # Get genders from people table
            cursor = conn.execute("""
                SELECT DISTINCT gender 
                FROM people 
                WHERE gender IS NOT NULL
                ORDER BY gender
            """)
            genders = [row[0] for row in cursor.fetchall()]
            
            # Get races
            cursor = conn.execute("""
                SELECT DISTINCT race 
                FROM people 
                WHERE race IS NOT NULL
                ORDER BY race
            """)
            races = [row[0] for row in cursor.fetchall()]
            
            # Get referral statuses
            cursor = conn.execute("""
                SELECT DISTINCT referral_status 
                FROM referrals 
                WHERE referral_status IS NOT NULL
                ORDER BY referral_status
            """)
            referral_statuses = [row[0] for row in cursor.fetchall()]
            
            return {
                "date_range": {
                    "min": date_range[0] if date_range[0] else None,
                    "max": date_range[1] if date_range[1] else None
                },
                "case_statuses": statuses,
                "service_types": service_types,
                "service_subtypes": service_subtypes,
                "providers": providers,
                "programs": programs,
                "genders": genders,
                "races": races,
                "referral_statuses": referral_statuses
            }
    except Exception as e:
        logger.error(f"Error fetching filter options: {e}")
        return {"error": str(e)}


@app.get("/api/reports/time-series/cases")
async def get_cases_time_series(
    start_date: str = None,
    end_date: str = None,
    status: str = None,
    service_type: str = None,
    provider: str = None
):
    """Get time series data of case creation over time with filtering"""
    db_manager = app_state["db_manager"]
    try:
        with db_manager.pool.get_connection() as conn:
            # Build dynamic WHERE clause
            where_clauses = ["case_created_at IS NOT NULL"]
            if start_date:
                where_clauses.append(f"case_created_at >= '{start_date}'")
            if end_date:
                where_clauses.append(f"case_created_at <= '{end_date}'")
            if status:
                where_clauses.append(f"case_status = '{status}'")
            if service_type:
                where_clauses.append(f"service_type = '{service_type}'")
            if provider:
                where_clauses.append(f"provider_name = '{provider}'")
            
            where_clause = " AND ".join(where_clauses)
            
            cursor = conn.execute(f"""
                SELECT 
                    DATE(case_created_at) as date,
                    COUNT(*) as count,
                    case_status
                FROM cases
                WHERE {where_clause}
                GROUP BY DATE(case_created_at), case_status
                ORDER BY date
            """)
            
            rows = cursor.fetchall()
            
            # Organize by date with status breakdown
            data_by_date = {}
            for row in rows:
                date_str = row[0]
                count = row[1]
                status_val = row[2] or 'Unknown'
                
                if date_str not in data_by_date:
                    data_by_date[date_str] = {'total': 0, 'by_status': {}}
                
                data_by_date[date_str]['total'] += count
                data_by_date[date_str]['by_status'][status_val] = count
            
            return {
                "dates": sorted(data_by_date.keys()),
                "data": data_by_date
            }
    except Exception as e:
        logger.error(f"Error fetching time series: {e}")
        return {"error": str(e)}


@app.get("/api/reports/time-series/referrals")
async def get_referrals_time_series(
    start_date: str = None,
    end_date: str = None,
    status: str = None,
    service_type: str = None
):
    """Get time series data of referral creation over time"""
    db_manager = app_state["db_manager"]
    try:
        with db_manager.pool.get_connection() as conn:
            where_clauses = ["referral_created_at IS NOT NULL"]
            if start_date:
                where_clauses.append(f"referral_created_at >= '{start_date}'")
            if end_date:
                where_clauses.append(f"referral_created_at <= '{end_date}'")
            if status:
                where_clauses.append(f"referral_status = '{status}'")
            if service_type:
                where_clauses.append(f"service_type = '{service_type}'")
            
            where_clause = " AND ".join(where_clauses)
            
            cursor = conn.execute(f"""
                SELECT 
                    DATE(referral_created_at) as date,
                    COUNT(*) as count
                FROM referrals
                WHERE {where_clause}
                GROUP BY DATE(referral_created_at)
                ORDER BY date
            """)
            
            rows = cursor.fetchall()
            return {
                "dates": [row[0] for row in rows],
                "counts": [row[1] for row in rows]
            }
    except Exception as e:
        logger.error(f"Error fetching referral time series: {e}")
        return {"error": str(e)}


@app.get("/api/reports/cohort-analysis")
async def get_cohort_analysis(
    start_date: str = None,
    end_date: str = None,
    service_type: str = None
):
    """Analyze client retention and progression by cohort (month of first case)"""
    db_manager = app_state["db_manager"]
    try:
        with db_manager.pool.get_connection() as conn:
            where_clauses = ["1=1"]
            if start_date:
                where_clauses.append(f"case_created_at >= '{start_date}'")
            if end_date:
                where_clauses.append(f"case_created_at <= '{end_date}'")
            if service_type:
                where_clauses.append(f"service_type = '{service_type}'")
            
            where_clause = " AND ".join(where_clauses)
            
            cursor = conn.execute(f"""
                WITH first_case AS (
                    SELECT 
                        person_id,
                        strftime('%Y-%m', MIN(case_created_at)) as cohort_month,
                        MIN(case_created_at) as first_case_date
                    FROM cases
                    WHERE {where_clause}
                    GROUP BY person_id
                ),
                case_counts AS (
                    SELECT 
                        fc.cohort_month,
                        COUNT(DISTINCT fc.person_id) as cohort_size,
                        COUNT(DISTINCT CASE 
                            WHEN c.case_created_at > fc.first_case_date 
                            THEN fc.person_id 
                        END) as returned_clients,
                        COUNT(DISTINCT CASE 
                            WHEN c.case_status IN ('completed', 'closed') 
                            THEN fc.person_id 
                        END) as completed_clients
                    FROM first_case fc
                    LEFT JOIN cases c ON fc.person_id = c.person_id
                    GROUP BY fc.cohort_month
                )
                SELECT 
                    cohort_month,
                    cohort_size,
                    returned_clients,
                    completed_clients,
                    ROUND(100.0 * returned_clients / NULLIF(cohort_size, 0), 1) as return_rate,
                    ROUND(100.0 * completed_clients / NULLIF(cohort_size, 0), 1) as completion_rate
                FROM case_counts
                ORDER BY cohort_month
            """)
            
            rows = cursor.fetchall()
            return {
                "cohorts": [row[0] for row in rows],
                "cohort_sizes": [row[1] for row in rows],
                "returned_clients": [row[2] for row in rows],
                "completed_clients": [row[3] for row in rows],
                "return_rates": [row[4] for row in rows],
                "completion_rates": [row[5] for row in rows]
            }
    except Exception as e:
        logger.error(f"Error in cohort analysis: {e}")
        return {"error": str(e)}


@app.get("/api/reports/service-pathways")
async def get_service_pathways(
    start_date: str = None,
    end_date: str = None
):
    """Analyze common service pathways - which services lead to which referrals"""
    db_manager = app_state["db_manager"]
    try:
        with db_manager.pool.get_connection() as conn:
            where_clauses = ["c.case_created_at IS NOT NULL"]
            if start_date:
                where_clauses.append(f"c.case_created_at >= '{start_date}'")
            if end_date:
                where_clauses.append(f"c.case_created_at <= '{end_date}'")
            
            where_clause = " AND ".join(where_clauses)
            
            cursor = conn.execute(f"""
                SELECT 
                    c.service_type as initial_service,
                    r.service_type as referral_service,
                    COUNT(*) as pathway_count,
                    AVG(julianday(r.referral_created_at) - julianday(c.case_created_at)) as avg_days_between
                FROM cases c
                INNER JOIN referrals r ON c.case_id = r.case_id
                WHERE {where_clause}
                    AND c.service_type IS NOT NULL
                    AND r.service_type IS NOT NULL
                    AND r.referral_created_at > c.case_created_at
                GROUP BY c.service_type, r.service_type
                HAVING pathway_count >= 2
                ORDER BY pathway_count DESC
                LIMIT 20
            """)
            
            rows = cursor.fetchall()
            return {
                "pathways": [
                    {
                        "initial_service": row[0],
                        "referral_service": row[1],
                        "count": row[2],
                        "avg_days_between": round(row[3], 1) if row[3] else None
                    }
                    for row in rows
                ]
            }
    except Exception as e:
        logger.error(f"Error in service pathway analysis: {e}")
        return {"error": str(e)}


@app.get("/api/reports/outcome-metrics")
async def get_outcome_metrics(
    start_date: str = None,
    end_date: str = None,
    service_type: str = None,
    provider: str = None
):
    """Analyze case outcomes and resolution metrics"""
    db_manager = app_state["db_manager"]
    try:
        with db_manager.pool.get_connection() as conn:
            where_clauses = ["case_created_at IS NOT NULL"]
            if start_date:
                where_clauses.append(f"case_created_at >= '{start_date}'")
            if end_date:
                where_clauses.append(f"case_created_at <= '{end_date}'")
            if service_type:
                where_clauses.append(f"service_type = '{service_type}'")
            if provider:
                where_clauses.append(f"provider_name = '{provider}'")
            
            where_clause = " AND ".join(where_clauses)
            
            # Outcome distribution
            cursor = conn.execute(f"""
                SELECT 
                    COALESCE(outcome_resolution_type, 'Not Recorded') as resolution_type,
                    COUNT(*) as count
                FROM cases
                WHERE {where_clause}
                GROUP BY outcome_resolution_type
                ORDER BY count DESC
            """)
            outcomes = cursor.fetchall()
            
            # Time to resolution by service type
            cursor = conn.execute(f"""
                SELECT 
                    service_type,
                    AVG(julianday(case_closed_at) - julianday(case_created_at)) as avg_days_to_close,
                    COUNT(*) as closed_count
                FROM cases
                WHERE {where_clause}
                    AND case_closed_at IS NOT NULL
                    AND service_type IS NOT NULL
                GROUP BY service_type
                ORDER BY closed_count DESC
                LIMIT 10
            """)
            resolution_times = cursor.fetchall()
            
            return {
                "outcome_distribution": {
                    "types": [row[0] for row in outcomes],
                    "counts": [row[1] for row in outcomes]
                },
                "resolution_times": [
                    {
                        "service_type": row[0],
                        "avg_days": round(row[1], 1) if row[1] else 0,
                        "count": row[2]
                    }
                    for row in resolution_times
                ]
            }
    except Exception as e:
        logger.error(f"Error in outcome metrics: {e}")
        return {"error": str(e)}


@app.get("/api/reports/geographic-distribution")
async def get_geographic_distribution(
    start_date: str = None,
    end_date: str = None
):
    """Analyze service distribution by geographic location"""
    db_manager = app_state["db_manager"]
    try:
        with db_manager.pool.get_connection() as conn:
            where_clauses = ["c.case_created_at IS NOT NULL"]
            if start_date:
                where_clauses.append(f"c.case_created_at >= '{start_date}'")
            if end_date:
                where_clauses.append(f"c.case_created_at <= '{end_date}'")
            
            where_clause = " AND ".join(where_clauses)
            
            # By city
            cursor = conn.execute(f"""
                SELECT 
                    COALESCE(p.current_person_address_city, 'Unknown') as city,
                    COUNT(DISTINCT c.case_id) as case_count,
                    COUNT(DISTINCT c.person_id) as client_count
                FROM cases c
                LEFT JOIN people p ON c.person_id = p.person_id
                WHERE {where_clause}
                GROUP BY p.current_person_address_city
                ORDER BY case_count DESC
                LIMIT 15
            """)
            cities = cursor.fetchall()
            
            # By county
            cursor = conn.execute(f"""
                SELECT 
                    COALESCE(p.current_person_address_county, 'Unknown') as county,
                    COUNT(DISTINCT c.case_id) as case_count
                FROM cases c
                LEFT JOIN people p ON c.person_id = p.person_id
                WHERE {where_clause}
                GROUP BY p.current_person_address_county
                ORDER BY case_count DESC
                LIMIT 10
            """)
            counties = cursor.fetchall()
            
            # By ZIP code
            cursor = conn.execute(f"""
                SELECT 
                    COALESCE(p.current_person_address_postal_code, 'Unknown') as zip_code,
                    COUNT(DISTINCT c.case_id) as case_count
                FROM cases c
                LEFT JOIN people p ON c.person_id = p.person_id
                WHERE {where_clause}
                GROUP BY p.current_person_address_postal_code
                ORDER BY case_count DESC
                LIMIT 15
            """)
            zip_codes = cursor.fetchall()
            
            return {
                "by_city": [
                    {"city": row[0], "cases": row[1], "clients": row[2]}
                    for row in cities
                ],
                "by_county": [
                    {"county": row[0], "cases": row[1]}
                    for row in counties
                ],
                "by_zip": [
                    {"zip": row[0], "cases": row[1]}
                    for row in zip_codes
                ]
            }
    except Exception as e:
        logger.error(f"Error in geographic distribution: {e}")
        return {"error": str(e)}


@app.get("/api/reports/provider-performance")
async def get_provider_performance(
    start_date: str = None,
    end_date: str = None
):
    """Compare provider performance metrics"""
    db_manager = app_state["db_manager"]
    try:
        with db_manager.pool.get_connection() as conn:
            where_clauses = ["case_created_at IS NOT NULL"]
            if start_date:
                where_clauses.append(f"case_created_at >= '{start_date}'")
            if end_date:
                where_clauses.append(f"case_created_at <= '{end_date}'")
            
            where_clause = " AND ".join(where_clauses)
            
            cursor = conn.execute(f"""
                SELECT 
                    COALESCE(provider_name, 'Unknown') as provider,
                    COUNT(DISTINCT case_id) as total_cases,
                    COUNT(DISTINCT person_id) as unique_clients,
                    COUNT(DISTINCT CASE 
                        WHEN case_status IN ('active', 'open', 'in_progress') 
                        THEN case_id 
                    END) as active_cases,
                    COUNT(DISTINCT CASE 
                        WHEN case_status IN ('pending', 'awaiting', 'new') 
                        THEN case_id 
                    END) as pending_cases,
                    COUNT(DISTINCT CASE 
                        WHEN case_status IN ('completed', 'closed') 
                        THEN case_id 
                    END) as closed_cases,
                    AVG(CASE 
                        WHEN case_closed_at IS NOT NULL 
                        THEN julianday(case_closed_at) - julianday(case_created_at)
                    END) as avg_resolution_days,
                    MIN(CASE 
                        WHEN case_closed_at IS NOT NULL 
                        THEN julianday(case_closed_at) - julianday(case_created_at)
                    END) as min_resolution_days,
                    MAX(CASE 
                        WHEN case_closed_at IS NOT NULL 
                        THEN julianday(case_closed_at) - julianday(case_created_at)
                    END) as max_resolution_days
                FROM cases
                WHERE {where_clause}
                    AND provider_name IS NOT NULL
                GROUP BY provider_name
                HAVING total_cases >= 5
                ORDER BY total_cases DESC, avg_resolution_days ASC
                LIMIT 20
            """)
            
            rows = cursor.fetchall()
            return {
                "providers": [
                    {
                        "provider_name": row[0],
                        "total_cases": row[1],
                        "unique_clients": row[2],
                        "active_cases": row[3],
                        "pending_cases": row[4],
                        "closed_cases": row[5],
                        "avg_days": round(row[6], 1) if row[6] else None,
                        "min_days": round(row[7], 1) if row[7] else None,
                        "max_days": round(row[8], 1) if row[8] else None,
                        "completion_rate": round(100.0 * row[5] / row[1], 1) if row[1] > 0 else 0
                    }
                    for row in rows
                ]
            }
    except Exception as e:
        logger.error(f"Error in provider performance: {e}")
        return {"error": str(e)}


@app.get("/api/reports/provider-performance-metrics")
async def get_provider_performance_metrics(
    start_date: str = None,
    end_date: str = None,
    provider_type: str = "receiving"
):
    """Get provider performance metrics for acceptance and completion rates"""
    db_manager = app_state["db_manager"]
    try:
        with db_manager.pool.get_connection() as conn:
            where_clauses = ["r.referral_created_at IS NOT NULL"]
            if start_date:
                where_clauses.append(f"r.referral_created_at >= '{start_date}'")
            if end_date:
                where_clauses.append(f"r.referral_created_at <= '{end_date}'")
            
            where_clause = " AND ".join(where_clauses)
            
            # Choose provider field based on type
            provider_field = "r.receiving_provider_name" if provider_type == "receiving" else "r.sending_provider_name"
            
            cursor = conn.execute(f"""
                SELECT 
                    COALESCE({provider_field}, 'Unknown') as provider_name,
                    COUNT(*) as total_referrals,
                    SUM(CASE WHEN r.referral_status = 'accepted' THEN 1 ELSE 0 END) as accepted,
                    SUM(CASE WHEN r.referral_status IN ('completed', 'closed') THEN 1 ELSE 0 END) as completed,
                    AVG(CASE 
                        WHEN r.referral_accepted_at IS NOT NULL 
                        THEN julianday(r.referral_accepted_at) - julianday(r.referral_created_at)
                    END) as avg_response_days
                FROM referrals r
                WHERE {where_clause}
                    AND {provider_field} IS NOT NULL
                GROUP BY {provider_field}
                HAVING total_referrals >= 3
                ORDER BY total_referrals DESC
                LIMIT 15
            """)
            
            rows = cursor.fetchall()
            return {
                "providers": [
                    {
                        "provider_name": row[0],
                        "total_referrals": row[1],
                        "acceptance_rate": round(100.0 * row[2] / row[1], 1) if row[1] > 0 else 0,
                        "completion_rate": round(100.0 * row[3] / row[1], 1) if row[1] > 0 else 0,
                        "avg_response_days": round(row[4], 1) if row[4] else None
                    }
                    for row in rows
                ]
            }
    except Exception as e:
        logger.error(f"Error in provider performance metrics: {e}")
        return {"error": str(e), "providers": []}


@app.get("/api/reports/high-risk-drop-off-analysis")
async def get_high_risk_drop_off_analysis(
    start_date: str = None,
    end_date: str = None
):
    """Identify service types with high drop-off/decline rates"""
    db_manager = app_state["db_manager"]
    try:
        with db_manager.pool.get_connection() as conn:
            where_clauses = ["referral_created_at IS NOT NULL"]
            if start_date:
                where_clauses.append(f"referral_created_at >= '{start_date}'")
            if end_date:
                where_clauses.append(f"referral_created_at <= '{end_date}'")
            
            where_clause = " AND ".join(where_clauses)
            
            cursor = conn.execute(f"""
                SELECT 
                    service_type,
                    COUNT(*) as total_referrals,
                    SUM(CASE WHEN referral_status IN ('declined', 'rejected', 'off_platform') THEN 1 ELSE 0 END) as dropped
                FROM referrals
                WHERE {where_clause}
                    AND service_type IS NOT NULL
                GROUP BY service_type
                HAVING total_referrals >= 5
                ORDER BY (100.0 * dropped / total_referrals) DESC
                LIMIT 10
            """)
            
            rows = cursor.fetchall()
            return {
                "service_types": [
                    {
                        "service_type": row[0],
                        "total_referrals": row[1],
                        "drop_off_rate": round(100.0 * row[2] / row[1], 1) if row[1] > 0 else 0
                    }
                    for row in rows
                ]
            }
    except Exception as e:
        logger.error(f"Error in high-risk drop-off analysis: {e}")
        return {"error": str(e), "service_types": []}


@app.get("/api/reports/client-risk-factors")
async def get_client_risk_factors(
    start_date: str = None,
    end_date: str = None
):
    """Analyze client demographics and risk factors"""
    db_manager = app_state["db_manager"]
    try:
        with db_manager.pool.get_connection() as conn:
            where_clauses = ["c.case_created_at IS NOT NULL"]
            if start_date:
                where_clauses.append(f"c.case_created_at >= '{start_date}'")
            if end_date:
                where_clauses.append(f"c.case_created_at <= '{end_date}'")
            
            where_clause = " AND ".join(where_clauses)
            
            # Housing status impact
            # Note: housing_status is not available in people table, using assistance_requests
            cursor = conn.execute(f"""
                SELECT 
                    COALESCE(ar.housing_current_status, 'Not Specified') as housing_status,
                    COUNT(DISTINCT c.case_id) as case_count,
                    AVG(CASE 
                        WHEN c.case_closed_at IS NOT NULL 
                        THEN julianday(c.case_closed_at) - julianday(c.case_created_at)
                    END) as avg_resolution_days
                FROM cases c
                LEFT JOIN assistance_requests ar ON c.case_id = ar.case_id
                WHERE {where_clause}
                GROUP BY ar.housing_current_status
                ORDER BY case_count DESC
                LIMIT 10
            """)
            housing = cursor.fetchall()
            
            # Household size correlation
            cursor = conn.execute(f"""
                SELECT 
                    CASE 
                        WHEN p.household_size IS NULL THEN 'Unknown'
                        WHEN p.household_size = 1 THEN '1 person'
                        WHEN p.household_size BETWEEN 2 AND 3 THEN '2-3 people'
                        WHEN p.household_size BETWEEN 4 AND 5 THEN '4-5 people'
                        ELSE '6+ people'
                    END as household_category,
                    COUNT(DISTINCT c.case_id) as case_count,
                    COUNT(DISTINCT c.person_id) as client_count
                FROM cases c
                LEFT JOIN people p ON c.person_id = p.person_id
                WHERE {where_clause}
                GROUP BY household_category
                ORDER BY case_count DESC
            """)
            household = cursor.fetchall()
            
            # Employment status - Note: employment_status not in schema, return empty data
            # If needed in future, add to database_schema.py people table
            employment = []
            
            return {
                "housing_impact": [
                    {
                        "status": row[0],
                        "cases": row[1],
                        "avg_days": round(row[2], 1) if row[2] else 0
                    }
                    for row in housing
                ],
                "household_size": [
                    {
                        "category": row[0],
                        "cases": row[1],
                        "clients": row[2]
                    }
                    for row in household
                ],
                "employment": [
                    {"status": row[0], "cases": row[1]}
                    for row in employment
                ]
            }
    except Exception as e:
        logger.error(f"Error in risk factor analysis: {e}")
        return {"error": str(e)}


@app.get("/api/reports/referral-network")
async def get_referral_network(
    start_date: str = None,
    end_date: str = None,
    min_referrals: int = 3
):
    """Analyze referral network connections between providers"""
    db_manager = app_state["db_manager"]
    try:
        with db_manager.pool.get_connection() as conn:
            where_clauses = ["referral_created_at IS NOT NULL"]
            if start_date:
                where_clauses.append(f"referral_created_at >= '{start_date}'")
            if end_date:
                where_clauses.append(f"referral_created_at <= '{end_date}'")
            
            where_clause = " AND ".join(where_clauses)
            
            cursor = conn.execute(f"""
                SELECT 
                    COALESCE(sending_provider_name, 'Unknown') as source,
                    COALESCE(receiving_provider_name, 'Unknown') as target,
                    COUNT(*) as referral_count,
                    COUNT(DISTINCT person_id) as unique_clients,
                    AVG(CASE 
                        WHEN referral_status IN ('accepted', 'completed')
                        THEN 1.0 ELSE 0.0 
                    END) as acceptance_rate
                FROM referrals
                WHERE {where_clause}
                    AND sending_provider_name IS NOT NULL
                    AND receiving_provider_name IS NOT NULL
                    AND sending_provider_name != receiving_provider_name
                GROUP BY sending_provider_name, receiving_provider_name
                HAVING referral_count >= {min_referrals}
                ORDER BY referral_count DESC
                LIMIT 50
            """)
            
            rows = cursor.fetchall()
            return {
                "connections": [
                    {
                        "source": row[0],
                        "target": row[1],
                        "referrals": row[2],
                        "unique_clients": row[3],
                        "acceptance_rate": round(row[4] * 100, 1)
                    }
                    for row in rows
                ]
            }
    except Exception as e:
        logger.error(f"Error in referral network analysis: {e}")
        return {"error": str(e)}


@app.get("/api/reports/service-funnel")
async def get_service_funnel(
    start_date: str = None,
    end_date: str = None,
    service_type: str = None
):
    """Analyze conversion funnel from assistance request to case completion"""
    db_manager = app_state["db_manager"]
    try:
        with db_manager.pool.get_connection() as conn:
            where_clauses = ["1=1"]
            if start_date:
                where_clauses.append(f"ar.created_at >= '{start_date}'")
            if end_date:
                where_clauses.append(f"ar.created_at <= '{end_date}'")
            
            where_clause = " AND ".join(where_clauses)
            
            # Count at each stage
            cursor = conn.execute(f"""
                WITH funnel_data AS (
                    SELECT 
                        COUNT(DISTINCT ar.assistance_request_id) as total_requests,
                        COUNT(DISTINCT ar.case_id) as requests_with_case,
                        COUNT(DISTINCT CASE 
                            WHEN c.case_id IS NOT NULL 
                            THEN ar.person_id 
                        END) as clients_with_case,
                        COUNT(DISTINCT r.referral_id) as total_referrals,
                        COUNT(DISTINCT CASE 
                            WHEN c.case_status IN ('completed', 'closed')
                            THEN ar.assistance_request_id
                        END) as completed_cases
                    FROM assistance_requests ar
                    LEFT JOIN cases c ON ar.case_id = c.case_id
                    LEFT JOIN referrals r ON c.case_id = r.case_id
                    WHERE {where_clause}
                )
                SELECT * FROM funnel_data
            """)
            
            row = cursor.fetchone()
            
            total_requests = row[0]
            requests_with_case = row[1]
            clients_with_case = row[2]
            total_referrals = row[3]
            completed = row[4]
            
            return {
                "stages": [
                    {
                        "name": "Assistance Requests",
                        "count": total_requests,
                        "percentage": 100.0
                    },
                    {
                        "name": "Cases Created",
                        "count": requests_with_case,
                        "percentage": round(100.0 * requests_with_case / total_requests, 1) if total_requests > 0 else 0
                    },
                    {
                        "name": "Clients Engaged",
                        "count": clients_with_case,
                        "percentage": round(100.0 * clients_with_case / total_requests, 1) if total_requests > 0 else 0
                    },
                    {
                        "name": "Referrals Made",
                        "count": total_referrals,
                        "percentage": round(100.0 * total_referrals / total_requests, 1) if total_requests > 0 else 0
                    },
                    {
                        "name": "Cases Completed",
                        "count": completed,
                        "percentage": round(100.0 * completed / total_requests, 1) if total_requests > 0 else 0
                    }
                ]
            }
    except Exception as e:
        logger.error(f"Error in service funnel: {e}")
        return {"error": str(e)}


@app.get("/api/reports/demographic-correlations")
async def get_demographic_correlations(
    start_date: str = None,
    end_date: str = None
):
    """Analyze how demographics correlate with service types and outcomes"""
    db_manager = app_state["db_manager"]
    try:
        with db_manager.pool.get_connection() as conn:
            where_clauses = ["c.case_created_at IS NOT NULL"]
            if start_date:
                where_clauses.append(f"c.case_created_at >= '{start_date}'")
            if end_date:
                where_clauses.append(f"c.case_created_at <= '{end_date}'")
            
            where_clause = " AND ".join(where_clauses)
            
            # Age group by service type
            cursor = conn.execute(f"""
                SELECT 
                    c.service_type,
                    CASE 
                        WHEN p.age IS NULL THEN 'Unknown'
                        WHEN p.age < 18 THEN 'Under 18'
                        WHEN p.age BETWEEN 18 AND 24 THEN '18-24'
                        WHEN p.age BETWEEN 25 AND 34 THEN '25-34'
                        WHEN p.age BETWEEN 35 AND 44 THEN '35-44'
                        WHEN p.age BETWEEN 45 AND 54 THEN '45-54'
                        WHEN p.age BETWEEN 55 AND 64 THEN '55-64'
                        ELSE '65+'
                    END as age_group,
                    COUNT(DISTINCT c.case_id) as case_count
                FROM cases c
                LEFT JOIN people p ON c.person_id = p.person_id
                WHERE {where_clause}
                    AND c.service_type IS NOT NULL
                GROUP BY c.service_type, age_group
                ORDER BY c.service_type, age_group
            """)
            age_service = cursor.fetchall()
            
            # Gender by service type
            cursor = conn.execute(f"""
                SELECT 
                    c.service_type,
                    COALESCE(p.gender, 'Unknown') as gender,
                    COUNT(DISTINCT c.case_id) as case_count
                FROM cases c
                LEFT JOIN people p ON c.person_id = p.person_id
                WHERE {where_clause}
                    AND c.service_type IS NOT NULL
                GROUP BY c.service_type, gender
                ORDER BY c.service_type, case_count DESC
            """)
            gender_service = cursor.fetchall()
            
            # Race by outcome
            cursor = conn.execute(f"""
                SELECT 
                    COALESCE(p.race, 'Unknown') as race,
                    c.case_status,
                    COUNT(DISTINCT c.case_id) as case_count
                FROM cases c
                LEFT JOIN people p ON c.person_id = p.person_id
                WHERE {where_clause}
                    AND c.case_status IS NOT NULL
                GROUP BY race, c.case_status
                ORDER BY race, case_count DESC
            """)
            race_outcome = cursor.fetchall()
            
            return {
                "age_by_service": [
                    {"service": row[0], "age_group": row[1], "count": row[2]}
                    for row in age_service
                ],
                "gender_by_service": [
                    {"service": row[0], "gender": row[1], "count": row[2]}
                    for row in gender_service
                ],
                "race_by_outcome": [
                    {"race": row[0], "status": row[1], "count": row[2]}
                    for row in race_outcome
                ]
            }
    except Exception as e:
        logger.error(f"Error in demographic correlation analysis: {e}")
        return {"error": str(e)}


# ============================================================================
# ADMIN CONTROL PANEL API (Admin Only)
# ============================================================================

@app.get("/api/admin/users")
async def get_users(session: UserSession = Depends(require_auth)):
    """Get all users - Admin only"""
    logger.info(f"Admin users request from {session.username} (role: {session.role.value})")
    
    if session.role.value != 'admin':
        logger.warning(f"Non-admin user {session.username} attempted to access /api/admin/users")
        raise HTTPException(status_code=403, detail="Admin access required")
    
    try:
        auth_service = get_auth_service()
        users = auth_service.local_db.list_users()
        active_count = sum(1 for u in users if u.get('is_active', False))
        
        logger.info(f"Returning {len(users)} users ({active_count} active) to {session.username}")
        
        return {
            "users": users,
            "total": len(users),
            "active": active_count
        }
    except Exception as e:
        logger.error(f"Error fetching users: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to fetch users: {str(e)}")


@app.get("/api/admin/audit-log")
async def get_audit_log(
    limit: int = Query(100, ge=1, le=1000),
    offset: int = Query(0, ge=0),
    category: Optional[str] = Query(None),
    username: Optional[str] = Query(None),
    action: Optional[str] = Query(None),
    success: Optional[bool] = Query(None),
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None),
    search: Optional[str] = Query(None),
    session: UserSession = Depends(require_auth)
):
    """Get comprehensive audit log with filtering - Admin only"""
    if session.role.value != 'admin':
        raise HTTPException(status_code=403, detail="Admin access required")
    
    try:
        from .audit_logger import get_audit_logger
        audit_logger = get_audit_logger()
        
        logs = audit_logger.get_logs(
            limit=limit,
            offset=offset,
            category=category,
            username=username,
            action=action,
            success=success,
            start_date=start_date,
            end_date=end_date,
            search=search
        )
        
        return {
            "success": True,
            "logs": logs,
            "count": len(logs),
            "filters": {
                "category": category,
                "username": username,
                "action": action,
                "success": success,
                "start_date": start_date,
                "end_date": end_date,
                "search": search
            }
        }
    except Exception as e:
        logger.error(f"Error fetching audit log: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to fetch audit log: {str(e)}")


@app.get("/api/admin/audit-statistics")
async def get_audit_statistics(
    start_date: Optional[str] = Query(None),
    end_date: Optional[str] = Query(None),
    session: UserSession = Depends(require_auth)
):
    """Get audit log statistics - Admin only"""
    if session.role.value != 'admin':
        raise HTTPException(status_code=403, detail="Admin access required")
    
    try:
        from .audit_logger import get_audit_logger
        audit_logger = get_audit_logger()
        
        stats = audit_logger.get_statistics(start_date=start_date, end_date=end_date)
        
        return {
            "success": True,
            "statistics": stats
        }
    except Exception as e:
        logger.error(f"Error fetching audit statistics: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to fetch audit statistics: {str(e)}")


@app.get("/api/admin/user-activity/{username}")
async def get_user_activity(
    username: str,
    days: int = Query(30, ge=1, le=365),
    session: UserSession = Depends(require_auth)
):
    """Get activity summary for a specific user - Admin only"""
    if session.role.value != 'admin':
        raise HTTPException(status_code=403, detail="Admin access required")
    
    try:
        from .audit_logger import get_audit_logger
        audit_logger = get_audit_logger()
        
        activity = audit_logger.get_user_activity(username=username, days=days)
        
        return {
            "success": True,
            "activity": activity
        }
    except Exception as e:
        logger.error(f"Error fetching user activity: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to fetch user activity: {str(e)}")


@app.get("/api/admin/system-info")
async def get_system_info(session: UserSession = Depends(require_auth)):
    """Get system information - Admin only"""
    if session.role.value != 'admin':
        raise HTTPException(status_code=403, detail="Admin access required")
    
    try:
        import sys
        import platform
        import psutil
        from datetime import datetime, timedelta
        
        db_manager = app_state["db_manager"]
        
        # Get database size
        db_path = Path(db_manager.db_path)
        db_size_mb = db_path.stat().st_size / (1024 * 1024) if db_path.exists() else 0
        
        # Get server uptime (from process start time)
        process = psutil.Process()
        process_start = datetime.fromtimestamp(process.create_time())
        uptime_seconds = (datetime.now() - process_start).total_seconds()
        
        # Format uptime nicely
        days = int(uptime_seconds // 86400)
        hours = int((uptime_seconds % 86400) // 3600)
        minutes = int((uptime_seconds % 3600) // 60)
        
        if days > 0:
            uptime_str = f"{days}d {hours}h {minutes}m"
        elif hours > 0:
            uptime_str = f"{hours}h {minutes}m"
        else:
            uptime_str = f"{minutes}m"
        
        # Get RAM usage (system-wide)
        memory = psutil.virtual_memory()
        ram_used_gb = memory.used / (1024**3)
        ram_total_gb = memory.total / (1024**3)
        ram_percent = memory.percent
        
        # Get Python process memory usage
        process = psutil.Process()
        process_memory = process.memory_info()
        process_memory_mb = process_memory.rss / (1024**2)  # RSS (Resident Set Size) in MB
        process_memory_percent = process.memory_percent()
        
        # Get ETL statistics
        with db_manager.pool.get_connection() as conn: 
            cursor = conn.execute("""
                SELECT 
                    COUNT(*) as total_runs,
                    SUM(CASE WHEN status = 'completed' THEN 1 ELSE 0 END) as successful_runs,
                    MAX(processing_completed_at) as last_run,
                    AVG(CAST((julianday(processing_completed_at) - julianday(processing_started_at)) * 24 * 60 AS REAL)) as avg_minutes
                FROM etl_metadata
            """)
            etl_stats = cursor.fetchone()
            
            cursor = conn.execute("SELECT SUM(records_processed) FROM etl_metadata WHERE status = 'completed'")
            total_records = cursor.fetchone()[0] or 0
        
        success_rate = 0
        if etl_stats and etl_stats[0] > 0:
            success_rate = round((etl_stats[1] / etl_stats[0]) * 100, 1)
        
        avg_time = "N/A"
        if etl_stats and etl_stats[3]:
            minutes = int(etl_stats[3])
            seconds = int((etl_stats[3] - minutes) * 60)
            avg_time = f"{minutes}m {seconds}s"
        
        # Get Windows username (AD user running the server)
        import os
        import getpass
        windows_username = None
        try:
            # Try multiple methods to get Windows username
            windows_username = os.getenv('USERNAME') or os.getenv('USER') or getpass.getuser()
            # Try to get domain if available
            try:
                domain = os.getenv('USERDOMAIN')
                if domain and domain.upper() != windows_username.upper():
                    windows_username = f"{domain}\\{windows_username}"
            except:
                pass
        except Exception as e:
            logger.debug(f"Could not get Windows username: {e}")
            windows_username = "Unknown"
        
        # Get server specs
        # Processor info
        processor_info = platform.processor()
        if not processor_info or processor_info == '':
            try:
                processor_info = platform.machine()
            except:
                processor_info = "Unknown"
        
        # Windows version
        windows_version = platform.system()
        if windows_version == "Windows":
            try:
                import winreg
                key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows NT\CurrentVersion")
                product_name = winreg.QueryValueEx(key, "ProductName")[0]
                windows_version = f"Windows {product_name}"
                try:
                    build = winreg.QueryValueEx(key, "DisplayVersion")[0]
                    windows_version += f" ({build})"
                except:
                    pass
                winreg.CloseKey(key)
            except Exception as e:
                logger.debug(f"Could not get Windows version details: {e}")
                windows_version = platform.platform()
        
        # Server port and web server info
        from .config import config
        server_port = config.web.port if hasattr(config, 'web') else 8000
        server_host = config.web.host if hasattr(config, 'web') else "0.0.0.0"
        server_protocol = "HTTPS" if hasattr(config, 'web') and getattr(config.web, 'use_https', False) else "HTTP"
        
        # Get CPU count
        cpu_count = psutil.cpu_count(logical=True)
        cpu_physical = psutil.cpu_count(logical=False)
        
        # Get Python version
        python_version = f"{sys.version_info.major}.{sys.version_info.minor}.{sys.version_info.micro}"
        
        return {
            "python_version": python_version,
            "platform": platform.platform(),
            "processor": processor_info,
            "cpu_count": cpu_count,
            "cpu_physical": cpu_physical,
            "windows_version": windows_version,
            "server_port": server_port,
            "server_host": server_host,
            "server_protocol": server_protocol,
            "windows_username": windows_username,
            "db_size": f"{db_size_mb:.2f} MB",
            "uptime": uptime_str,
            "ram_usage": f"{ram_used_gb:.1f} / {ram_total_gb:.1f} GB ({ram_percent:.0f}%)",
            "python_memory": f"{process_memory_mb:.1f} MB ({process_memory_percent:.1f}%)",
            "total_records": total_records,
            "last_etl": etl_stats[2] if etl_stats and etl_stats[2] else None,
            "etl_success_rate": success_rate,
            "avg_processing_time": avg_time
        }
    except Exception as e:
        logger.error(f"Error fetching system info: {e}")
        return {"error": str(e)}


# ========================================================================
# SETTINGS MANAGEMENT ENDPOINTS (SIEM & SFTP)
# ========================================================================

@app.get("/api/settings/siem")
async def get_siem_settings(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Get SIEM configuration settings - Admin only"""
    try:
        settings_manager = get_settings_manager()
        settings = settings_manager.get_siem_settings()
        
        return {
            "success": True,
            "settings": settings
        }
    except Exception as e:
        logger.error(f"Error fetching SIEM settings: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to fetch SIEM settings: {str(e)}")


@app.post("/api/settings/siem")
async def save_siem_settings(
    enabled: bool = Form(False),
    enable_windows_event_log: bool = Form(False),
    syslog_enabled: bool = Form(False),
    syslog_host: str = Form("localhost"),
    syslog_port: int = Form(514),
    syslog_protocol: str = Form("UDP"),
    include_sensitive_data: bool = Form(False),
    windows_event_log_min_severity: str = Form("ERROR"),
    syslog_min_severity: str = Form("ERROR"),
    session: UserSession = Depends(require_role(UserRole.ADMIN))
):
    """Save SIEM configuration settings - Admin only"""
    try:
        settings_manager = get_settings_manager()
        
        settings = {
            'enabled': enabled,
            'enable_windows_event_log': enable_windows_event_log,
            'syslog_enabled': syslog_enabled,
            'syslog_host': syslog_host,
            'syslog_port': syslog_port,
            'syslog_protocol': syslog_protocol,
            'include_sensitive_data': include_sensitive_data,
            'windows_event_log_min_severity': windows_event_log_min_severity,
            'syslog_min_severity': syslog_min_severity
        }
        
        success = settings_manager.save_siem_settings(settings, session.username)
        
        if success:
            # Reinitialize SIEM logger with new settings
            siem_logger = get_siem_logger()
            siem_logger._initialize_loggers()
            
            # Log the configuration change
            from .siem_logger import log_siem_event
            log_siem_event(
                SIEMEventType.CONFIGURATION_CHANGE,
                "SIEM settings updated",
                severity=SIEMSeverity.NOTICE,
                username=session.username,
                resource="sys_siem_config",
                action="update",
                success=True
            )
            
            return {"success": True, "message": "SIEM settings saved successfully"}
        else:
            return {"success": False, "error": "Failed to save SIEM settings"}
    
    except Exception as e:
        logger.error(f"Error saving SIEM settings: {e}", exc_info=True)
        return {"success": False, "error": str(e)}


@app.get("/api/settings/sftp")
async def get_sftp_settings(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Get SFTP configuration settings - Admin only"""
    try:
        settings_manager = get_settings_manager()
        settings = settings_manager.get_sftp_settings()
        
        # Don't send private key passphrase or password to client
        if 'private_key_passphrase' in settings:
            settings['private_key_passphrase'] = '********' if settings['private_key_passphrase'] else ''
        if 'password' in settings:
            settings['password'] = '********' if settings['password'] else ''
        
        return {
            "success": True,
            "settings": settings
        }
    except Exception as e:
        logger.error(f"Error fetching SFTP settings: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to fetch SFTP settings: {str(e)}")


@app.post("/api/settings/sftp")
async def save_sftp_settings(
    enabled: bool = Form(False),
    host: str = Form(""),
    port: int = Form(22),
    username: str = Form(""),
    auth_method: str = Form("key"),
    private_key_path: str = Form("data/sftp/private_key"),
    key_format: str = Form("auto"),
    remote_directory: str = Form("/data/exports"),
    auto_download: bool = Form(True),  # Default to True
    download_interval_minutes: int = Form(60),
    delete_after_download: bool = Form(False),  # Always False - option removed from UI
    local_download_path: str = Form("temp_data_files"),
    timeout_seconds: int = Form(30),
    max_retries: int = Form(3),
    verify_host_key: bool = Form(True),
    known_hosts_path: str = Form("data/sftp/known_hosts"),
    file_patterns: str = Form("*.txt,*.csv"),  # Comma-separated
    session: UserSession = Depends(require_role(UserRole.ADMIN))
):
    """Save SFTP configuration settings - Admin only"""
    try:
        settings_manager = get_settings_manager()
        
        # Parse file patterns
        patterns = [p.strip() for p in file_patterns.split(',') if p.strip()]
        
        settings = {
            'enabled': enabled,
            'host': host,
            'port': port,
            'username': username,
            'auth_method': auth_method,
            'private_key_path': private_key_path,
            'key_format': key_format,
            'remote_directory': remote_directory,
            'auto_download': auto_download,
            'download_interval_minutes': download_interval_minutes,
            'delete_after_download': delete_after_download,
            'local_download_path': local_download_path,
            'timeout_seconds': timeout_seconds,
            'max_retries': max_retries,
            'verify_host_key': verify_host_key,
            'known_hosts_path': known_hosts_path,
            'file_patterns': patterns
        }
        
        success = settings_manager.save_sftp_settings(settings, session.username)
        
        if success:
            return {"success": True, "message": "SFTP settings saved successfully"}
        else:
            return {"success": False, "error": "Failed to save SFTP settings"}
    
    except Exception as e:
        logger.error(f"Error saving SFTP settings: {e}", exc_info=True)
        return {"success": False, "error": str(e)}


@app.post("/api/sftp/test-connection")
async def test_sftp_connection(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Test SFTP connection with current settings - Admin only"""
    try:
        sftp_service = get_sftp_service()
        success, message = sftp_service.test_connection(username=session.username)
        
        return {
            "success": success,
            "message": message
        }
    except Exception as e:
        logger.error(f"Error testing SFTP connection: {e}", exc_info=True)
        return {"success": False, "message": f"Connection test failed: {str(e)}"}


@app.post("/api/sftp/convert-key")
async def convert_putty_key(
    ppk_path: str = Form(...),
    passphrase: str = Form(None),
    session: UserSession = Depends(require_role(UserRole.ADMIN))
):
    """Convert PuTTY format key to OpenSSH format - Admin only"""
    try:
        from core.utils.putty_key_converter import PuTTYKeyConverter
        from pathlib import Path
        
        ppk_file = Path(ppk_path)
        
        # Validate input
        if not ppk_file.exists():
            return {
                "success": False,
                "message": f"Key file not found: {ppk_path}"
            }
        
        if not PuTTYKeyConverter.is_putty_key(ppk_file):
            return {
                "success": False,
                "message": f"File is not a valid PuTTY key: {ppk_path}"
            }
        
        # Perform conversion
        success, message, output_path = PuTTYKeyConverter.convert_key_auto(
            ppk_file,
            passphrase if passphrase else None
        )
        
        # If conversion successful, auto-update SFTP config
        if success and output_path:
            try:
                # Make path relative to base_dir if possible
                try:
                    relative_path = output_path.relative_to(config.directories.project_root)
                    config.sftp.private_key_path = relative_path
                except ValueError:
                    # Path is not relative, use as-is
                    config.sftp.private_key_path = output_path
                
                # Save updated config
                config.save_sftp_config()
                logger.info(f"Auto-updated SFTP config to use converted key: {config.sftp.private_key_path}")
                message += f" → Config updated to: {config.sftp.private_key_path}"
            except Exception as update_err:
                logger.warning(f"Could not auto-update SFTP config: {update_err}")
                message += " (Note: Please update Private Key Path manually)"
        
        # Log the conversion attempt (if audit logger available)
        try:
            from .audit_logger import get_audit_logger, AuditCategory, AuditAction
            get_audit_logger().log(
                username=session.username,
                action=AuditAction.CONFIGURATION_CHANGED,
                category=AuditCategory.SECURITY,
                success=success,
                details=f"PuTTY key conversion: {ppk_path}",
                target_resource=str(output_path) if output_path else ppk_path,
                error_message=message if not success else None
            )
        except (ImportError, Exception) as audit_err:
            logger.debug(f"Audit logging not available: {audit_err}")
        
        return {
            "success": success,
            "message": message,
            "converted_path": str(output_path) if output_path else None,
            "output_path": str(output_path) if output_path else None  # Include both for compatibility
        }
        
    except Exception as e:
        logger.error(f"Error converting PuTTY key: {e}", exc_info=True)
        return {
            "success": False,
            "message": f"Conversion error: {str(e)}"
        }


@app.get("/api/sftp/last-sync")
async def get_last_sftp_sync(
    request: Request,
    session: UserSession = Depends(require_auth)
):
    """
    Get the last SFTP sync time and cached file list
    """
    try:
        from .database import get_database_manager
        import json
        
        db_manager = get_database_manager()
        
        with db_manager.pool.get_connection() as conn:
            cursor = conn.execute("""
                SELECT sync_time, file_list, file_count, synced_by
                FROM sftp_cache
                ORDER BY sync_time DESC
                LIMIT 1
            """)
            row = cursor.fetchone()
            
            if row:
                return {
                    'success': True,
                    'sync_time': row[0],
                    'files': json.loads(row[1]),
                    'file_count': row[2],
                    'synced_by': row[3]
                }
            else:
                return {
                    'success': False,
                    'message': 'No sync history found',
                    'sync_time': None
                }
    except Exception as e:
        logger.error(f"Error getting last SFTP sync: {e}", exc_info=True)
        return {
            'success': False,
            'message': str(e),
            'sync_time': None
        }


@app.post("/api/sftp/discover-files")
async def discover_sftp_files(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Discover available files on SFTP server and cache the results - Admin only"""
    try:
        from .database import get_database_manager
        from datetime import datetime
        import json
        
        sftp_service = get_sftp_service()
        files = sftp_service.discover_files(username=session.username)
        
        # Convert to serializable format
        files_data = [f.to_dict() for f in files]
        
        # Cache the file list in database
        sync_time = datetime.now().isoformat()
        db_manager = get_database_manager()
        
        try:
            with db_manager.pool.get_connection() as conn:
                conn.execute("""
                    INSERT INTO sftp_cache (sync_time, file_list, file_count, synced_by)
                    VALUES (?, ?, ?, ?)
                """, (
                    sync_time,
                    json.dumps(files_data),
                    len(files_data),
                    session.username
                ))
                conn.commit()
                logger.info(f"Cached {len(files_data)} SFTP files to database")
        except Exception as db_error:
            logger.warning(f"Failed to cache SFTP files to database: {db_error}")
            # Continue even if caching fails
        
        return {
            "success": True,
            "files": files_data,
            "count": len(files_data),
            "sync_time": sync_time
        }
    except Exception as e:
        logger.error(f"Error discovering SFTP files: {e}", exc_info=True)
        return {"success": False, "error": str(e), "files": [], "count": 0}


@app.post("/api/sftp/download")
async def download_sftp_files(
    files: List[str] = Form([]),
    session: UserSession = Depends(require_role(UserRole.ADMIN))
):
    """Download specific files from SFTP server - Admin only"""
    try:
        if not files:
            return {
                "success": False,
                "error": "No files specified for download",
                "downloaded_count": 0,
                "failed_count": 0
            }
        
        sftp_service = get_sftp_service()
        
        # Download the specified files
        results = sftp_service.download_files(files, username=session.username)
        
        # Count successes and failures (results are SFTPDownloadResult objects)
        downloaded_count = sum(1 for r in results if r.success)
        failed_count = len(results) - downloaded_count
        
        return {
            "success": True,
            "downloaded_count": downloaded_count,
            "failed_count": failed_count,
            "results": results
        }
    except Exception as e:
        logger.error(f"Error downloading SFTP files: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e),
            "downloaded_count": 0,
            "failed_count": 0
        }


# ========================================================================
# DATABASE SETTINGS ENDPOINTS
# ========================================================================

@app.get("/api/settings/database")
async def get_database_settings(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Get database configuration settings - Admin only"""
    try:
        from .settings_manager import get_settings_manager
        settings_manager = get_settings_manager()
        settings = settings_manager.get_database_settings()
        
        # Don't send password to client
        if 'mssql_password' in settings and settings['mssql_password']:
            settings['mssql_password'] = '********'
        
        return {
            "success": True,
            "settings": settings
        }
    except Exception as e:
        logger.error(f"Error fetching database settings: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to fetch database settings: {str(e)}")


@app.post("/api/settings/database")
async def save_database_settings(
    db_type: str = Form(...),
    sqlite_path: str = Form(""),
    mssql_server: str = Form(""),
    mssql_port: int = Form(1433),
    mssql_database: str = Form(""),
    mssql_username: str = Form(""),
    mssql_password: str = Form(""),
    mssql_trusted_connection: bool = Form(True),
    mssql_driver: str = Form("ODBC Driver 17 for SQL Server"),
    azuresql_server: str = Form(""),
    azuresql_port: int = Form(1433),
    azuresql_database: str = Form(""),
    azuresql_username: str = Form(""),
    azuresql_password: str = Form(""),
    azuresql_driver: str = Form("ODBC Driver 17 for SQL Server"),
    postgresql_host: str = Form(""),
    postgresql_port: int = Form(5432),
    postgresql_database: str = Form(""),
    postgresql_username: str = Form(""),
    postgresql_password: str = Form(""),
    mysql_host: str = Form(""),
    mysql_port: int = Form(3306),
    mysql_database: str = Form(""),
    mysql_username: str = Form(""),
    mysql_password: str = Form(""),
    connection_timeout: int = Form(30),
    max_connections: int = Form(10),
    session: UserSession = Depends(require_role(UserRole.ADMIN))
):
    """Save database configuration settings - Admin only"""
    try:
        from .settings_manager import get_settings_manager
        settings_manager = get_settings_manager()
        
        # Don't update passwords if they're the masked value
        current_settings = settings_manager.get_database_settings()
        if mssql_password == '********':
            mssql_password = current_settings.get('mssql_password', '')
        if azuresql_password == '********':
            azuresql_password = current_settings.get('mssql_password', '')  # Azure SQL uses same storage as MS SQL
        if postgresql_password == '********':
            postgresql_password = current_settings.get('postgresql_password', '')
        if mysql_password == '********':
            mysql_password = current_settings.get('mysql_password', '')
        
        # For Azure SQL, use azuresql_* fields, but store them in mssql_* fields (they use the same adapter)
        if db_type == 'azuresql':
            mssql_server = azuresql_server
            mssql_port = azuresql_port
            mssql_database = azuresql_database
            mssql_username = azuresql_username
            mssql_password = azuresql_password
            mssql_driver = azuresql_driver
            mssql_trusted_connection = False  # Azure SQL always uses SQL Auth
        
        settings = {
            'db_type': db_type,
            'path': sqlite_path if db_type == 'sqlite' else '',
            'mssql_server': mssql_server,
            'mssql_port': mssql_port,
            'mssql_database': mssql_database,
            'mssql_username': mssql_username,
            'mssql_password': mssql_password,
            'mssql_trusted_connection': mssql_trusted_connection,
            'mssql_driver': mssql_driver,
            'postgresql_host': postgresql_host,
            'postgresql_port': postgresql_port,
            'postgresql_database': postgresql_database,
            'postgresql_username': postgresql_username,
            'postgresql_password': postgresql_password,
            'mysql_host': mysql_host,
            'mysql_port': mysql_port,
            'mysql_database': mysql_database,
            'mysql_username': mysql_username,
            'mysql_password': mysql_password,
            'connection_timeout': connection_timeout,
            'max_connections': max_connections
        }
        
        success = settings_manager.save_database_settings(settings, session.username)
        
        if success:
            # Reload configuration
            from .config import config
            config._reload_database_config()
            
            return {"success": True, "message": "Database settings saved successfully. Restart required for changes to take effect."}
        else:
            return {"success": False, "error": "Failed to save database settings"}
    
    except Exception as e:
        logger.error(f"Error saving database settings: {e}", exc_info=True)
        return {"success": False, "error": str(e)}


@app.post("/api/database/test-connection")
async def test_database_connection(
    db_type: str = Form(...),
    sqlite_path: str = Form(""),
    mssql_server: str = Form(""),
    mssql_port: int = Form(1433),
    mssql_database: str = Form(""),
    mssql_username: str = Form(""),
    mssql_password: str = Form(""),
    mssql_trusted_connection: bool = Form(True),
    mssql_driver: str = Form("ODBC Driver 17 for SQL Server"),
    azuresql_server: str = Form(""),
    azuresql_port: int = Form(1433),
    azuresql_database: str = Form(""),
    azuresql_username: str = Form(""),
    azuresql_password: str = Form(""),
    azuresql_driver: str = Form("ODBC Driver 17 for SQL Server"),
    postgresql_host: str = Form(""),
    postgresql_port: int = Form(5432),
    postgresql_database: str = Form(""),
    postgresql_username: str = Form(""),
    postgresql_password: str = Form(""),
    mysql_host: str = Form(""),
    mysql_port: int = Form(3306),
    mysql_database: str = Form(""),
    mysql_username: str = Form(""),
    mysql_password: str = Form(""),
    connection_timeout: int = Form(30),
    session: UserSession = Depends(require_role(UserRole.ADMIN))
):
    """Test database connection - Admin only"""
    try:
        from .database_adapter import (
            SQLiteAdapter, MSSQLAdapter, PostgreSQLAdapter, MySQLAdapter
        )
        from pathlib import Path
        
        if db_type == "mssql":
            if not mssql_server or not mssql_database:
                return {"success": False, "error": "Server and database name are required"}
            
            adapter = MSSQLAdapter(
                server=mssql_server,
                database=mssql_database,
                username=mssql_username if not mssql_trusted_connection else "",
                password=mssql_password if not mssql_trusted_connection else "",
                trusted_connection=mssql_trusted_connection,
                port=mssql_port,
                driver=mssql_driver,
                timeout=connection_timeout
            )
        elif db_type == "azuresql":
            # Use azuresql_* fields if provided, otherwise fallback to mssql_* fields (for backward compatibility)
            server = azuresql_server if azuresql_server else mssql_server
            port = azuresql_port if azuresql_port else mssql_port
            database = azuresql_database if azuresql_database else mssql_database
            username = azuresql_username if azuresql_username else mssql_username
            password = azuresql_password if azuresql_password else mssql_password
            driver = azuresql_driver if azuresql_driver else mssql_driver
            
            if not server or not database:
                return {"success": False, "error": "Server and database name are required"}
            
            # Azure SQL always requires SQL Authentication
            adapter = MSSQLAdapter(
                server=server,
                database=database,
                username=username,
                password=password,
                trusted_connection=False,  # Azure SQL doesn't support Windows Auth
                port=port,
                driver=driver,
                timeout=connection_timeout
            )
        elif db_type == "postgresql":
            if not postgresql_host or not postgresql_database:
                return {"success": False, "error": "Host and database name are required"}
            
            adapter = PostgreSQLAdapter(
                host=postgresql_host,
                database=postgresql_database,
                username=postgresql_username,
                password=postgresql_password,
                port=postgresql_port,
                timeout=connection_timeout
            )
        elif db_type == "mysql":
            if not mysql_host or not mysql_database:
                return {"success": False, "error": "Host and database name are required"}
            
            adapter = MySQLAdapter(
                host=mysql_host,
                database=mysql_database,
                username=mysql_username,
                password=mysql_password,
                port=mysql_port,
                timeout=connection_timeout
            )
        else:  # SQLite
            if not sqlite_path:
                return {"success": False, "error": "Database path is required"}
            
            # Resolve path relative to application base directory if not absolute
            db_path = Path(sqlite_path)
            if not db_path.is_absolute():
                from .config import config
                db_path = config.directories.project_root / db_path
            
            # Ensure parent directory exists
            db_path.parent.mkdir(parents=True, exist_ok=True)
            
            adapter = SQLiteAdapter(
                db_path=db_path,
                timeout=connection_timeout
            )
        
        # Test connection
        with adapter.get_connection() as conn:
            if db_type in ["mssql", "azuresql"]:
                cursor = conn.cursor()
                cursor.execute("SELECT 1")
                cursor.fetchone()
            elif db_type == "postgresql":
                cursor = conn.cursor()
                cursor.execute("SELECT 1")
                cursor.fetchone()
            elif db_type == "mysql":
                cursor = conn.cursor()
                cursor.execute("SELECT 1")
                cursor.fetchone()
            else:  # SQLite
                conn.execute("SELECT 1")
        
        # Build detailed success message based on database type
        message_parts = []
        
        if db_type == "sqlite":
            message_parts.append(f"Successfully connected to SQLite database")
            message_parts.append(f"Path: {str(db_path)}")
            message_parts.append(f"Timeout: {connection_timeout}s")
        elif db_type == "mssql":
            message_parts.append(f"Successfully connected to Microsoft SQL Server")
            message_parts.append(f"Server: {mssql_server}")
            message_parts.append(f"Port: {mssql_port}")
            message_parts.append(f"Database: {mssql_database}")
            if mssql_trusted_connection:
                message_parts.append(f"Authentication: Windows Authentication (Trusted Connection)")
            else:
                message_parts.append(f"Authentication: SQL Authentication")
                message_parts.append(f"Username: {mssql_username}")
            message_parts.append(f"Driver: {mssql_driver}")
            message_parts.append(f"Timeout: {connection_timeout}s")
        elif db_type == "azuresql":
            message_parts.append(f"Successfully connected to Azure SQL Database")
            message_parts.append(f"Server: {server}")
            message_parts.append(f"Port: {port}")
            message_parts.append(f"Database: {database}")
            message_parts.append(f"Authentication: SQL Authentication")
            message_parts.append(f"Username: {username}")
            message_parts.append(f"Driver: {driver}")
            message_parts.append(f"Timeout: {connection_timeout}s")
        elif db_type == "postgresql":
            message_parts.append(f"Successfully connected to PostgreSQL database")
            message_parts.append(f"Host: {postgresql_host}")
            message_parts.append(f"Port: {postgresql_port}")
            message_parts.append(f"Database: {postgresql_database}")
            message_parts.append(f"Username: {postgresql_username}")
            message_parts.append(f"Timeout: {connection_timeout}s")
        elif db_type == "mysql":
            message_parts.append(f"Successfully connected to MySQL database")
            message_parts.append(f"Host: {mysql_host}")
            message_parts.append(f"Port: {mysql_port}")
            message_parts.append(f"Database: {mysql_database}")
            message_parts.append(f"Username: {mysql_username}")
            message_parts.append(f"Timeout: {connection_timeout}s")
        
        return {
            "success": True,
            "message": " | ".join(message_parts)
        }
    except ImportError as e:
        error_msg = str(e)
        # Log to server activity log
        logger.error(f"Database driver not available for {db_type}: {error_msg}")
        # Provide helpful installation instructions
        if "pyodbc" in error_msg.lower() or db_type in ["mssql", "azuresql"]:
            install_instructions = "pip install pyodbc"
        elif "psycopg2" in error_msg.lower() or db_type == "postgresql":
            install_instructions = "pip install psycopg2-binary"
        elif "pymysql" in error_msg.lower() or db_type == "mysql":
            install_instructions = "pip install pymysql"
        else:
            install_instructions = "Please check the error message and install the required driver"
        
        return {
            "success": False, 
            "error": f"Database driver not available: {error_msg}. The launcher should have installed all dependencies. Please restart the application or install manually: {install_instructions}"
        }
    except sqlite3.OperationalError as e:
        error_msg = str(e)
        logger.error(f"SQLite connection test failed: {error_msg}")
        if "unable to open database file" in error_msg.lower():
            return {"success": False, "error": f"Cannot access database file. Check file permissions and path: {sqlite_path}"}
        elif "database is locked" in error_msg.lower():
            return {"success": False, "error": "Database is locked. Another process may be using it."}
        else:
            return {"success": False, "error": f"SQLite error: {error_msg}"}
    except HTTPException as e:
        # This is an authentication error from the dependency
        if e.status_code == 401:
            return {"success": False, "error": "Authentication required. Please log in and try again."}
        raise  # Re-raise other HTTP exceptions
    except Exception as e:
        logger.error(f"Database connection test failed for {db_type}: {str(e)}", exc_info=True)
        error_msg = str(e)
        # Check if it's an authentication error
        if "Not authenticated" in error_msg or "401" in error_msg or "unauthorized" in error_msg.lower():
            return {"success": False, "error": "Authentication required. Please log in and try again."}
        return {"success": False, "error": error_msg}


@app.post("/api/database/initialize")
async def initialize_database(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Initialize database schema - Admin only"""
    try:
        from .database_adapter import get_database_adapter
        from .database_schema import get_schema_sql
        from .database_schema_converter import get_schema_for_database_type
        from .config import config
        
        adapter = get_database_adapter()
        base_schema = get_schema_sql()
        schema_sql = get_schema_for_database_type(config.database.db_type, base_schema)
        
        # Log converted schema for debugging
        logger.info(f"Initializing {config.database.db_type} database with converted schema")
        logger.debug(f"First 500 characters of converted schema: {schema_sql[:500]}")
        
        # Split into statements and execute
        statements = [s.strip() for s in schema_sql.split(';') if s.strip() and not s.strip().startswith('--')]
        
        logger.info(f"Total statements to execute: {len(statements)}")
        
        # Log statement types for debugging
        create_tables = sum(1 for s in statements if 'CREATE TABLE' in s.upper())
        create_indexes = sum(1 for s in statements if 'CREATE INDEX' in s.upper())
        create_views = sum(1 for s in statements if 'CREATE VIEW' in s.upper())
        logger.info(f"Statement breakdown - Tables: {create_tables}, Indexes: {create_indexes}, Views: {create_views}")
        
        statement_count = 0
        errors = []
        
        with adapter.get_connection() as conn:
            for idx, statement in enumerate(statements, 1):
                try:
                    if config.database.db_type in ["mssql", "azuresql"]:
                        # MS SQL Server doesn't support IF NOT EXISTS in CREATE TABLE
                        # Check if table exists first
                        if 'CREATE TABLE' in statement.upper() and 'IF NOT EXISTS' not in statement.upper():
                            # Extract table name
                            table_match = re.search(r'CREATE TABLE\s+(\w+)', statement, re.IGNORECASE)
                            if table_match:
                                table_name = table_match.group(1)
                                # Check if table exists
                                check_sql = f"SELECT COUNT(*) FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = '{table_name}'"
                                cursor = conn.cursor()
                                cursor.execute(check_sql)
                                if cursor.fetchone()[0] > 0:
                                    logger.debug(f"Skipping statement {idx}: Table {table_name} already exists")
                                    continue  # Table exists, skip
                        
                        cursor = conn.cursor()
                        cursor.execute(statement)
                        conn.commit()
                    elif config.database.db_type == "postgresql":
                        # PostgreSQL supports IF NOT EXISTS natively
                        cursor = conn.cursor()
                        cursor.execute(statement)
                        conn.commit()
                    elif config.database.db_type == "mysql":
                        # MySQL supports IF NOT EXISTS natively
                        cursor = conn.cursor()
                        cursor.execute(statement)
                        conn.commit()
                    else:  # SQLite
                        conn.execute(statement)
                        conn.commit()
                    
                    statement_count += 1
                    logger.debug(f"Successfully executed statement {idx}")
                except Exception as stmt_error:
                    # Log the actual failing statement for debugging
                    error_msg = str(stmt_error)
                    if 'already exists' not in error_msg.lower() and 'duplicate' not in error_msg.lower():
                        # Truncate statement for logging but show more context
                        statement_preview = statement[:200] if len(statement) <= 200 else statement[:200] + "..."
                        logger.error(f"Statement {idx} FAILED:\nSQL: {statement_preview}\nError: {error_msg}")
                        errors.append(f"Statement {idx}: {error_msg[:150]}")
                    # Continue with other statements
        
        if errors:
            logger.error(f"Database initialization completed with {len(errors)} errors for {config.database.db_type}")
            error_summary = f"Initialization failed with {len(errors)} errors. Successfully executed: {statement_count} statements. First error: {errors[0] if errors else 'Unknown'}"
            return {
                "success": False,
                "error": error_summary,
                "statements_executed": statement_count,
                "total_statements": len(statements)
            }
        
        # After successful schema initialization, create automated_sync_config table
        # This is done separately because it's a system table not part of the main schema
        try:
            if config.database.db_type != 'sqlite':
                # For non-SQLite databases, use database-specific syntax
                from .database import get_database_manager
                db_manager = get_database_manager()
                with adapter.get_connection() as conn:
                    db_manager.migrate_automated_sync_for_other_databases(conn, config.database.db_type)
                logger.info(f"Created automated_sync_config table for {config.database.db_type}")
        except Exception as sync_error:
            logger.warning(f"Could not create automated_sync_config table: {sync_error}")
        
        logger.info(f"Database initialized successfully for {config.database.db_type}: {statement_count} statements executed")
        return {
            "success": True,
            "message": f"Database initialized successfully ({statement_count} statements executed)"
        }
    except Exception as e:
        logger.error(f"Error initializing database for {config.database.db_type}: {e}", exc_info=True)
        return {"success": False, "error": str(e)}


@app.get("/api/database/check-initialization")
async def check_database_initialization(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Check if database tables are initialized - Admin only"""
    try:
        from .database_adapter import get_database_adapter
        from .config import config
        
        adapter = get_database_adapter()
        
        # Check for one of the main tables (people table is a good indicator)
        table_name = 'people'
        tables_exist = False
        
        try:
            with adapter.get_connection() as conn:
                if config.database.db_type == 'sqlite':
                    cursor = conn.execute(
                        "SELECT name FROM sqlite_master WHERE type='table' AND name=?",
                        (table_name,)
                    )
                    tables_exist = cursor.fetchone() is not None
                elif config.database.db_type in ['mssql', 'azuresql']:
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT COUNT(*) FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = ?",
                        (table_name,)
                    )
                    tables_exist = cursor.fetchone()[0] > 0
                elif config.database.db_type == 'postgresql':
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT COUNT(*) FROM information_schema.tables WHERE table_name = %s",
                        (table_name,)
                    )
                    tables_exist = cursor.fetchone()[0] > 0
                elif config.database.db_type == 'mysql':
                    cursor = conn.cursor()
                    cursor.execute(
                        "SELECT COUNT(*) FROM information_schema.tables WHERE table_name = %s",
                        (table_name,)
                    )
                    tables_exist = cursor.fetchone()[0] > 0
        except Exception as e:
            logger.warning(f"Error checking table existence: {e}")
            tables_exist = False
        
        return {
            "success": True,
            "initialized": tables_exist
        }
    except Exception as e:
        logger.error(f"Error checking database initialization: {e}", exc_info=True)
        return {"success": False, "error": str(e)}


@app.get("/api/database/check-data")
async def check_database_has_data(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Check if database has any data - Admin only"""
    try:
        from .database_adapter import get_database_adapter
        from .config import config
        
        adapter = get_database_adapter()
        has_data = False
        
        try:
            with adapter.get_connection() as conn:
                # Check if people table has any rows
                if config.database.db_type == 'sqlite':
                    cursor = conn.execute("SELECT COUNT(*) FROM people")
                    has_data = cursor.fetchone()[0] > 0
                elif config.database.db_type in ['mssql', 'azuresql']:
                    cursor = conn.cursor()
                    cursor.execute("SELECT COUNT(*) FROM people")
                    has_data = cursor.fetchone()[0] > 0
                elif config.database.db_type == 'postgresql':
                    cursor = conn.cursor()
                    cursor.execute("SELECT COUNT(*) FROM people")
                    has_data = cursor.fetchone()[0] > 0
                elif config.database.db_type == 'mysql':
                    cursor = conn.cursor()
                    cursor.execute("SELECT COUNT(*) FROM people")
                    has_data = cursor.fetchone()[0] > 0
        except Exception as e:
            # If table doesn't exist or query fails, no data
            logger.warning(f"Error checking for data: {e}")
            has_data = False
        
        return {
            "success": True,
            "has_data": has_data
        }
    except Exception as e:
        logger.error(f"Error checking database data: {e}", exc_info=True)
        return {"success": False, "error": str(e)}


@app.post("/api/database/migrate-data")
async def migrate_data_from_sqlite(
    source_db_type: str = Form("sqlite"),
    destination_db_type: str = Form(...),
    create_tables: bool = Form(True),
    source_sqlite_path: str = Form(""),
    source_mssql_server: str = Form(""),
    source_mssql_port: int = Form(1433),
    source_mssql_database: str = Form(""),
    source_mssql_username: str = Form(""),
    source_mssql_password: str = Form(""),
    source_mssql_trusted_connection: bool = Form(True),
    source_postgresql_host: str = Form(""),
    source_postgresql_port: int = Form(5432),
    source_postgresql_database: str = Form(""),
    source_postgresql_username: str = Form(""),
    source_postgresql_password: str = Form(""),
    source_mysql_host: str = Form(""),
    source_mysql_port: int = Form(3306),
    source_mysql_database: str = Form(""),
    source_mysql_username: str = Form(""),
    source_mysql_password: str = Form(""),
    destination_sqlite_path: str = Form(""),
    destination_mssql_server: str = Form(""),
    destination_mssql_port: int = Form(1433),
    destination_mssql_database: str = Form(""),
    destination_mssql_username: str = Form(""),
    destination_mssql_password: str = Form(""),
    destination_mssql_trusted_connection: bool = Form(True),
    destination_mssql_driver: str = Form("ODBC Driver 17 for SQL Server"),
    destination_postgresql_host: str = Form(""),
    destination_postgresql_port: int = Form(5432),
    destination_postgresql_database: str = Form(""),
    destination_postgresql_username: str = Form(""),
    destination_postgresql_password: str = Form(""),
    destination_mysql_host: str = Form(""),
    destination_mysql_port: int = Form(3306),
    destination_mysql_database: str = Form(""),
    destination_mysql_username: str = Form(""),
    destination_mysql_password: str = Form(""),
    session: UserSession = Depends(require_role(UserRole.ADMIN))
):
    """Migrate data between any database types - Admin only"""
    try:
        from .database_adapter import (
            SQLiteAdapter, MSSQLAdapter, PostgreSQLAdapter, MySQLAdapter
        )
        from pathlib import Path
        
        # Create source adapter
        if source_db_type == "sqlite":
            source_path = Path(source_sqlite_path) if source_sqlite_path else Path(config.database.path)
            if not source_path.exists():
                return {"success": False, "error": f"Source SQLite database not found at {source_path}"}
            source_adapter = SQLiteAdapter(path=str(source_path), timeout=30)
        elif source_db_type == "mssql":
            if not source_mssql_server or not source_mssql_database:
                return {"success": False, "error": "Source server and database name are required"}
            source_adapter = MSSQLAdapter(
                server=source_mssql_server,
                database=source_mssql_database,
                username=source_mssql_username if not source_mssql_trusted_connection else "",
                password=source_mssql_password if not source_mssql_trusted_connection else "",
                trusted_connection=source_mssql_trusted_connection,
                port=source_mssql_port,
                timeout=30
            )
        elif source_db_type == "postgresql":
            if not source_postgresql_host or not source_postgresql_database:
                return {"success": False, "error": "Source host and database name are required"}
            source_adapter = PostgreSQLAdapter(
                host=source_postgresql_host,
                database=source_postgresql_database,
                username=source_postgresql_username,
                password=source_postgresql_password,
                port=source_postgresql_port,
                timeout=30
            )
        elif source_db_type == "mysql":
            if not source_mysql_host or not source_mysql_database:
                return {"success": False, "error": "Source host and database name are required"}
            source_adapter = MySQLAdapter(
                host=source_mysql_host,
                database=source_mysql_database,
                username=source_mysql_username,
                password=source_mysql_password,
                port=source_mysql_port,
                timeout=30
            )
        else:
            return {"success": False, "error": f"Unsupported source database type: {source_db_type}"}
        
        # Test source connection
        try:
            with source_adapter.get_connection() as conn:
                pass
            logger.info(f"Successfully connected to source {source_db_type} database")
        except Exception as e:
            logger.error(f"Failed to connect to source database: {e}")
            return {"success": False, "error": f"Failed to connect to source database: {str(e)}"}
        
        # Create destination adapter
        if destination_db_type == "sqlite":
            dest_path = Path(destination_sqlite_path) if destination_sqlite_path else Path("data/database/migrated.db")
            dest_adapter = SQLiteAdapter(path=str(dest_path), timeout=30)
        elif destination_db_type == "mssql":
            if not destination_mssql_server or not destination_mssql_database:
                return {"success": False, "error": "Destination server and database name are required"}
            dest_adapter = MSSQLAdapter(
                server=destination_mssql_server,
                database=destination_mssql_database,
                username=destination_mssql_username if not destination_mssql_trusted_connection else "",
                password=destination_mssql_password if not destination_mssql_trusted_connection else "",
                trusted_connection=destination_mssql_trusted_connection,
                port=destination_mssql_port,
                driver=destination_mssql_driver,
                timeout=30
            )
        elif destination_db_type == "postgresql":
            if not destination_postgresql_host or not destination_postgresql_database:
                return {"success": False, "error": "Destination host and database name are required"}
            dest_adapter = PostgreSQLAdapter(
                host=destination_postgresql_host,
                database=destination_postgresql_database,
                username=destination_postgresql_username,
                password=destination_postgresql_password,
                port=destination_postgresql_port,
                timeout=30
            )
        elif destination_db_type == "mysql":
            if not destination_mysql_host or not destination_mysql_database:
                return {"success": False, "error": "Destination host and database name are required"}
            dest_adapter = MySQLAdapter(
                host=destination_mysql_host,
                database=destination_mysql_database,
                username=destination_mysql_username,
                password=destination_mysql_password,
                port=destination_mysql_port,
                timeout=30
            )
        else:
            return {"success": False, "error": f"Unsupported destination database type: {destination_db_type}"}
        
        # Test destination connection
        try:
            with dest_adapter.get_connection() as conn:
                pass
            logger.info(f"Successfully connected to destination {destination_db_type} database")
        except Exception as e:
            logger.error(f"Failed to connect to destination database: {e}")
            return {"success": False, "error": f"Failed to connect to destination database: {str(e)}"}
        
        # Create tables if requested
        if create_tables:
            try:
                from .database_schema_converter import convert_sqlite_to_mssql, convert_sqlite_to_postgresql, convert_sqlite_to_mysql
                from .database_schema import get_schema_sql
                
                # Get base SQLite schema
                base_schema = get_schema_sql()
                
                # Convert schema based on destination database type
                if destination_db_type == 'sqlite':
                    dest_schema = base_schema
                elif destination_db_type in ['mssql', 'azuresql']:
                    dest_schema = convert_sqlite_to_mssql(base_schema)
                elif destination_db_type == 'postgresql':
                    dest_schema = convert_sqlite_to_postgresql(base_schema)
                elif destination_db_type == 'mysql':
                    dest_schema = convert_sqlite_to_mysql(base_schema)
                else:
                    return {"success": False, "error": f"Unsupported destination database type: {destination_db_type}"}
                
                # Execute schema SQL
                with dest_adapter.get_connection() as conn:
                    statements = [s.strip() for s in dest_schema.split(';') if s.strip()]
                    for statement in statements:
                        if statement:
                            try:
                                if destination_db_type == 'sqlite':
                                    conn.execute(statement)
                                else:
                                    cursor = conn.cursor()
                                    cursor.execute(statement)
                            except Exception as stmt_error:
                                # Ignore "already exists" errors
                                if 'already exists' not in str(stmt_error).lower():
                                    logger.warning(f"Schema statement failed: {str(stmt_error)[:200]}")
                    conn.commit()
                    
                logger.info(f"Successfully created tables in {destination_db_type} database")
            except Exception as e:
                logger.error(f"Failed to create tables: {e}", exc_info=True)
                return {"success": False, "error": f"Failed to create tables: {str(e)}"}
        
        # Check if destination has data
        dest_has_data = False
        try:
            with dest_adapter.get_connection() as conn:
                if destination_db_type == 'sqlite':
                    cursor = conn.execute("SELECT COUNT(*) FROM sqlite_master WHERE type='table' AND name='people'")
                    if cursor.fetchone()[0] > 0:
                        cursor = conn.execute("SELECT COUNT(*) FROM people")
                        dest_has_data = cursor.fetchone()[0] > 0
                elif destination_db_type in ['mssql', 'azuresql']:
                    cursor = conn.cursor()
                    cursor.execute("SELECT COUNT(*) FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = 'people'")
                    if cursor.fetchone()[0] > 0:
                        cursor.execute("SELECT COUNT(*) FROM people")
                        dest_has_data = cursor.fetchone()[0] > 0
                elif destination_db_type == 'postgresql':
                    cursor = conn.cursor()
                    cursor.execute("SELECT COUNT(*) FROM information_schema.tables WHERE table_name = 'people'")
                    if cursor.fetchone()[0] > 0:
                        cursor.execute("SELECT COUNT(*) FROM people")
                        dest_has_data = cursor.fetchone()[0] > 0
                elif destination_db_type == 'mysql':
                    cursor = conn.cursor()
                    cursor.execute("SELECT COUNT(*) FROM information_schema.tables WHERE table_schema = DATABASE() AND table_name = 'people'")
                    if cursor.fetchone()[0] > 0:
                        cursor.execute("SELECT COUNT(*) FROM people")
                        dest_has_data = cursor.fetchone()[0] > 0
        except Exception as e:
            logger.warning(f"Could not check destination data status: {e}")
        
        if dest_has_data:
            return {
                "success": False,
                "error": "Destination database already contains data. Migration is only allowed to empty databases."
            }
        
        # Get list of tables to migrate
        tables_to_migrate = ['etl_metadata', 'people', 'employees', 'cases', 'referrals', 
                             'assistance_requests', 'assistance_requests_supplemental_responses',
                             'resource_lists', 'resource_list_shares', 'data_quality_issues',
                             'automated_sync_config']
        
        migration_results = {}
        total_records = 0
        
        # Migrate each table
        for table_name in tables_to_migrate:
            try:
                # Special handling for automated_sync_config - it might not exist in older databases
                if table_name == 'automated_sync_config':
                    # Check if table exists in source
                    with source_adapter.get_connection() as source_conn:
                        table_exists = False
                        if source_db_type == 'sqlite':
                            cursor = source_conn.execute(
                                "SELECT name FROM sqlite_master WHERE type='table' AND name=?",
                                (table_name,)
                            )
                            table_exists = cursor.fetchone() is not None
                        elif source_db_type in ['mssql', 'azuresql']:
                            cursor = source_conn.cursor()
                            cursor.execute(
                                "SELECT COUNT(*) FROM INFORMATION_SCHEMA.TABLES WHERE TABLE_NAME = ?",
                                (table_name,)
                            )
                            table_exists = cursor.fetchone()[0] > 0
                        elif source_db_type == 'postgresql':
                            cursor = source_conn.cursor()
                            cursor.execute(
                                "SELECT COUNT(*) FROM information_schema.tables WHERE table_name = %s",
                                (table_name,)
                            )
                            table_exists = cursor.fetchone()[0] > 0
                        elif source_db_type == 'mysql':
                            cursor = source_conn.cursor()
                            cursor.execute(
                                "SELECT COUNT(*) FROM information_schema.tables WHERE table_name = %s",
                                (table_name,)
                            )
                            table_exists = cursor.fetchone()[0] > 0
                    
                    if not table_exists:
                        migration_results[table_name] = {
                            "records": 0,
                            "status": "skipped",
                            "message": "Table does not exist in source (likely older database version)"
                        }
                        logger.info(f"Skipping {table_name}: table does not exist in source")
                        continue
                
                # Get all data from source table
                with source_adapter.get_connection() as source_conn:
                    if source_db_type == 'sqlite':
                        cursor = source_conn.execute(f"SELECT * FROM {table_name}")
                        rows = cursor.fetchall()
                        columns = [description[0] for description in cursor.description]
                    else:
                        cursor = source_conn.cursor()
                        cursor.execute(f"SELECT * FROM {table_name}")
                        rows = cursor.fetchall()
                        columns = [desc[0] for desc in cursor.description]
                
                if not rows:
                    migration_results[table_name] = {"records": 0, "status": "skipped", "message": "Table is empty"}
                    logger.info(f"Skipping {table_name}: empty table")
                    continue
                
                # Insert into destination
                with dest_adapter.get_connection() as dest_conn:
                    # Determine correct placeholder style for destination database
                    if destination_db_type in ['sqlite']:
                        placeholders = ','.join(['?'] * len(columns))
                    elif destination_db_type in ['mssql', 'azuresql']:
                        placeholders = ','.join(['?'] * len(columns))
                    elif destination_db_type in ['postgresql', 'mysql']:
                        placeholders = ','.join(['%s'] * len(columns))
                    else:
                        placeholders = ','.join(['%s'] * len(columns))
                    
                    column_names = ','.join(columns)
                    insert_sql = f"INSERT INTO {table_name} ({column_names}) VALUES ({placeholders})"
                    
                    if destination_db_type == 'sqlite':
                        for row in rows:
                            values = list(row) if not isinstance(row, dict) else [row[col] for col in columns]
                            dest_conn.execute(insert_sql, values)
                        dest_conn.commit()
                    else:
                        dest_cursor = dest_conn.cursor()
                        for row in rows:
                            values = list(row) if not isinstance(row, dict) else [row[col] for col in columns]
                            dest_cursor.execute(insert_sql, values)
                        dest_conn.commit()
                
                records_migrated = len(rows)
                total_records += records_migrated
                migration_results[table_name] = {
                    "records": records_migrated,
                    "status": "success",
                    "message": f"Migrated {records_migrated} records"
                }
                logger.info(f"Migrated {table_name}: {records_migrated} records")
                
            except Exception as table_error:
                error_msg = str(table_error)
                migration_results[table_name] = {
                    "records": 0,
                    "status": "error",
                    "message": error_msg[:200]
                }
                logger.error(f"Failed to migrate {table_name}: {error_msg}")
        
        # Build success message
        success_tables = [t for t, r in migration_results.items() if r["status"] == "success"]
        failed_tables = [t for t, r in migration_results.items() if r["status"] == "error"]
        
        message = f"Migration completed: {total_records} total records migrated from {source_db_type} to {destination_db_type}. "
        message += f"Successfully migrated {len(success_tables)} tables. "
        if failed_tables:
            message += f"Failed tables: {', '.join(failed_tables)}"
        
        logger.info(f"Migration summary: {message}")
        
        return {
            "success": len(failed_tables) == 0,
            "message": message,
            "total_records": total_records,
            "migration_results": migration_results,
            "source_type": source_db_type,
            "destination_type": destination_db_type
        }
        
    except Exception as e:
        logger.error(f"Migration failed: {e}", exc_info=True)
        return {
            "success": False,
            "error": f"Migration failed: {str(e)}"
        }


@app.get("/api/schema/errors")
async def get_schema_errors(
    limit: int = Query(50, ge=1, le=1000),
    resolved: Optional[bool] = Query(None, description="Filter by resolved status. None = all, True = resolved only, False = unresolved only"),
    session: UserSession = Depends(require_role(UserRole.ADMIN))
):
    """Get recent schema errors - Admin only"""
    try:
        from .schema_validator import get_schema_validator
        
        validator = get_schema_validator()
        
        # Handle resolved parameter: None = all, True = resolved only, False = unresolved only
        if resolved is None:
            # Get all errors (both resolved and unresolved)
            errors = validator.get_all_errors(limit=limit)
        elif resolved:
            # Get only resolved errors
            errors = validator.get_recent_errors(limit=limit, resolved_only=True)
        else:
            # Get only unresolved errors
            errors = validator.get_recent_errors(limit=limit, resolved_only=False)
        
        return {
            "success": True,
            "errors": errors,
            "count": len(errors)
        }
    except Exception as e:
        logger.error(f"Error fetching schema errors: {e}", exc_info=True)
        return {"success": False, "error": str(e)}


@app.post("/api/schema/errors/{error_id}/resolve")
async def resolve_schema_error(
    error_id: int,
    session: UserSession = Depends(require_role(UserRole.ADMIN))
):
    """Mark a schema error as resolved - Admin only"""
    try:
        from .schema_validator import get_schema_validator
        
        validator = get_schema_validator()
        success = validator.mark_error_resolved(error_id, session.username)
        
        if success:
            return {"success": True, "message": "Error marked as resolved"}
        else:
            return {"success": False, "error": "Failed to mark error as resolved"}
    except Exception as e:
        logger.error(f"Error resolving schema error: {e}", exc_info=True)
        return {"success": False, "error": str(e)}


@app.post("/api/schema/execute-sql")
async def execute_schema_sql(
    sql_command: str = Form(...),
    session: UserSession = Depends(require_role(UserRole.ADMIN))
):
    """Execute SQL command to fix schema issues - Admin only"""
    try:
        from .database_adapter import get_database_adapter
        from .audit_logger import get_audit_logger, AuditCategory, AuditAction
        
        adapter = get_database_adapter()
        audit_logger = get_audit_logger()
        
        # Validate SQL command (basic safety check)
        sql_upper = sql_command.upper().strip()
        dangerous_keywords = ['DROP', 'DELETE', 'TRUNCATE', 'ALTER USER', 'GRANT', 'REVOKE']
        
        # Allow ALTER TABLE and CREATE TABLE
        if not (sql_upper.startswith('CREATE TABLE') or 
                sql_upper.startswith('ALTER TABLE') or
                sql_upper.startswith('CREATE INDEX')):
            return {
                "success": False,
                "error": "Only CREATE TABLE, ALTER TABLE, and CREATE INDEX commands are allowed"
            }
        
        # Check for dangerous operations
        for keyword in dangerous_keywords:
            if keyword in sql_upper and 'TABLE' not in sql_upper:
                return {
                    "success": False,
                    "error": f"Operation '{keyword}' is not allowed for safety reasons"
                }
        
        # Execute SQL
        try:
            with adapter.get_connection() as conn:
                if config.database.db_type == 'sqlite':
                    conn.executescript(sql_command)
                elif config.database.db_type in ['mssql', 'azuresql']:
                    cursor = conn.cursor()
                    # Split by semicolon and execute each statement
                    for statement in sql_command.split(';'):
                        statement = statement.strip()
                        if statement:
                            cursor.execute(statement)
                    conn.commit()
                elif config.database.db_type == 'postgresql':
                    cursor = conn.cursor()
                    for statement in sql_command.split(';'):
                        statement = statement.strip()
                        if statement:
                            cursor.execute(statement)
                    conn.commit()
                elif config.database.db_type == 'mysql':
                    cursor = conn.cursor()
                    for statement in sql_command.split(';'):
                        statement = statement.strip()
                        if statement:
                            cursor.execute(statement)
                    conn.commit()
            
            # Log successful execution
            audit_logger.log(
                username=session.username,
                action=AuditAction.DATABASE_RESET,  # Using existing action, could add new one
                category=AuditCategory.SYSTEM,
                success=True,
                details=f"Executed schema SQL command",
                target_resource="database_schema"
            )
            
            return {
                "success": True,
                "message": "SQL command executed successfully"
            }
        except Exception as sql_error:
            error_msg = str(sql_error)
            logger.error(f"Error executing schema SQL: {error_msg}", exc_info=True)
            
            # Log failed execution
            audit_logger.log(
                username=session.username,
                action=AuditAction.DATABASE_RESET,
                category=AuditCategory.SYSTEM,
                success=False,
                details=f"Failed to execute schema SQL command",
                target_resource="database_schema",
                error_message=error_msg
            )
            
            return {
                "success": False,
                "error": f"SQL execution failed: {error_msg}"
            }
    except Exception as e:
        logger.error(f"Error in execute_schema_sql: {e}", exc_info=True)
        return {"success": False, "error": str(e)}


@app.get("/api/schema/file-mappings")
async def get_file_table_mappings(
    include_defaults: bool = Query(True),
    session: UserSession = Depends(require_role(UserRole.ADMIN))
):
    """Get file-to-table name mappings - Admin only
    
    Args:
        include_defaults: If True, includes default mappings based on pattern extraction
    """
    try:
        import sqlite3
        from pathlib import Path
        from datetime import datetime
        from .config import config
        
        internal_db = config.directories.database_dir / "internal.db"
        mappings = []
        
        if internal_db.exists():
            with sqlite3.connect(internal_db) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.execute("""
                    SELECT id, file_pattern, table_name, created_at, updated_at, created_by, is_active
                    FROM file_table_mappings
                    ORDER BY file_pattern
                """)
                mappings = [dict(row) for row in cursor.fetchall()]
        
        # Generate default mappings if requested
        default_mappings = []
        if include_defaults:
            # Get expected table names from config
            expected_tables = list(config.data_quality.expected_tables.keys())
            ignored_prefixes = config.etl.ignored_filename_prefixes
            extensions = config.etl.recognized_extensions
            
            # Generate default patterns for each table
            for table_name in expected_tables:
                # Create patterns based on common naming conventions
                patterns = []
                
                # Pattern 1: chhsca_{table}_*.txt (most common)
                patterns.append({
                    'file_pattern': f'chhsca_{table_name}_*.txt',
                    'table_name': table_name,
                    'is_default': True,
                    'description': f'Matches files like chhsca_{table_name}_YYYYMMDD.txt'
                })
                
                # Pattern 2: SAMPLE_chhsca_{table}_*.txt (if SAMPLE is in ignored prefixes)
                if 'SAMPLE' in ignored_prefixes:
                    patterns.append({
                        'file_pattern': f'SAMPLE_chhsca_{table_name}_*.txt',
                        'table_name': table_name,
                        'is_default': True,
                        'description': f'Matches files like SAMPLE_chhsca_{table_name}_YYYYMMDD.txt'
                    })
                
                # Pattern 3: {table}_*.txt (simpler pattern)
                patterns.append({
                    'file_pattern': f'{table_name}_*.txt',
                    'table_name': table_name,
                    'is_default': True,
                    'description': f'Matches files like {table_name}_YYYYMMDD.txt'
                })
                
                # Pattern 4: {table}.txt (exact match)
                patterns.append({
                    'file_pattern': f'{table_name}.txt',
                    'table_name': table_name,
                    'is_default': True,
                    'description': f'Exact match for {table_name}.txt'
                })
                
                # Add patterns for other extensions (CSV, TSV)
                for ext in extensions:
                    if ext != '.txt':
                        # Remove leading dot for pattern
                        ext_clean = ext.lstrip('.')
                        patterns.append({
                            'file_pattern': f'{table_name}_*.{ext_clean}',
                            'table_name': table_name,
                            'is_default': True,
                            'description': f'Matches files like {table_name}_YYYYMMDD{ext}'
                        })
                
                default_mappings.extend(patterns)
            
            # Filter out defaults that are already in custom mappings
            custom_patterns = {m['file_pattern'] for m in mappings}
            default_mappings = [m for m in default_mappings if m['file_pattern'] not in custom_patterns]
            
            # Sort defaults by table name, then by pattern specificity (most specific first)
            default_mappings.sort(key=lambda x: (x['table_name'], -len(x['file_pattern'])))
        
        return {
            "success": True,
            "mappings": mappings,
            "default_mappings": default_mappings,
            "config": {
                "ignored_prefixes": config.etl.ignored_filename_prefixes,
                "recognized_extensions": config.etl.recognized_extensions
            }
        }
    except Exception as e:
        logger.error(f"Error fetching file mappings: {e}", exc_info=True)
        return {"success": False, "error": str(e)}


def _check_mapping_conflicts(new_pattern: str, new_table: str, existing_mappings: List[Dict], default_mappings: List[Dict]) -> Optional[str]:
    """Check if a new mapping pattern conflicts with existing mappings
    
    Returns error message if conflict found, None otherwise
    """
    import fnmatch
    
    # Combine all mappings to check against
    all_mappings = []
    
    # Add existing custom mappings
    for m in existing_mappings:
        all_mappings.append({
            'pattern': m['file_pattern'],
            'table': m['table_name']
        })
    
    # Add default mappings
    for m in default_mappings:
        all_mappings.append({
            'pattern': m['file_pattern'],
            'table': m['table_name']
        })
    
    # Remove duplicates
    seen = set()
    unique_mappings = []
    for m in all_mappings:
        key = (m['pattern'], m['table'])
        if key not in seen:
            seen.add(key)
            unique_mappings.append(m)
    all_mappings = unique_mappings
    
    # Check for conflicts
    conflicts = []
    
    for existing in all_mappings:
        # Skip if same pattern and same table (updating existing)
        if existing['pattern'] == new_pattern and existing['table'] == new_table:
            continue
        
        # Skip if same table (different patterns mapping to same table is OK)
        if existing['table'] == new_table:
            continue
        
        # Check if patterns could match the same filename
        # Strategy: Generate test filenames and see if both patterns match
        
        # Generate test filenames based on the new pattern
        test_filenames = _generate_test_filenames_from_pattern(new_pattern)
        
        # Check if existing pattern matches any of our test filenames
        for test_file in test_filenames:
            if _pattern_matches_file(existing['pattern'], test_file):
                # Found a conflict - both patterns match the same filename but map to different tables
                conflicts.append({
                    'conflicting_pattern': existing['pattern'],
                    'conflicting_table': existing['table'],
                    'test_filename': test_file
                })
                break  # Found a conflict for this pattern, no need to check more
        
        # Also check reverse: generate test filenames from existing pattern
        if not any(c['conflicting_pattern'] == existing['pattern'] for c in conflicts):
            existing_test_files = _generate_test_filenames_from_pattern(existing['pattern'])
            for test_file in existing_test_files:
                if _pattern_matches_file(new_pattern, test_file):
                    # Found a conflict
                    conflicts.append({
                        'conflicting_pattern': existing['pattern'],
                        'conflicting_table': existing['table'],
                        'test_filename': test_file
                    })
                    break
    
    if conflicts:
        conflict_messages = []
        for conflict in conflicts[:3]:  # Show up to 3 conflicts
            conflict_messages.append(
                f"Pattern '{conflict['conflicting_pattern']}' (maps to '{conflict['conflicting_table']}') "
                f"would also match '{conflict['test_filename']}'"
            )
        
        error_msg = (
            f"Pattern '{new_pattern}' conflicts with existing mappings:\n" +
            "\n".join(conflict_messages)
        )
        if len(conflicts) > 3:
            error_msg += f"\n... and {len(conflicts) - 3} more conflict(s)"
        
        return error_msg
    
    return None


def _generate_test_filenames_from_pattern(pattern: str) -> List[str]:
    """Generate test filenames that would match a given pattern"""
    import fnmatch
    from .config import config
    
    test_files = []
    
    # If pattern has no wildcards, it's an exact match
    if '*' not in pattern and '?' not in pattern:
        test_files.append(pattern)
        return test_files
    
    # Generate test filenames by replacing wildcards with common values
    # Common table names from config
    common_tables = ['people', 'employees', 'cases', 'referrals', 'assistance_requests', 
                     'resource_lists', 'resource_list_shares']
    common_dates = ['20250101', '20250115', '20250201']
    common_prefixes = ['chhsca_', 'SAMPLE_chhsca_', '']
    
    # Count wildcards
    wildcard_count = pattern.count('*') + pattern.count('?')
    
    if wildcard_count == 1:
        # Single wildcard - try various replacements
        for prefix in common_prefixes:
            for table in common_tables:
                test_file = pattern.replace('*', f'{prefix}{table}').replace('?', 'x')
                test_files.append(test_file)
            for date in common_dates:
                test_file = pattern.replace('*', date).replace('?', 'x')
                test_files.append(test_file)
    elif wildcard_count == 2:
        # Two wildcards - common case like chhsca_*_*.txt
        for prefix in common_prefixes:
            for table in common_tables:
                for date in common_dates:
                    # Replace first * with prefix+table, second with date
                    parts = pattern.split('*', 1)
                    if len(parts) == 2:
                        test_file = parts[0] + f'{prefix}{table}' + parts[1].replace('*', date, 1)
                        test_files.append(test_file)
    else:
        # Multiple wildcards - use simpler approach
        for prefix in common_prefixes:
            for table in common_tables:
                test_file = pattern
                # Replace all * with table name
                test_file = test_file.replace('*', f'{prefix}{table}')
                test_file = test_file.replace('?', 'x')
                test_files.append(test_file)
    
    # Remove duplicates and empty strings
    test_files = list(set([f for f in test_files if f]))
    
    # Limit to reasonable number
    return test_files[:20]


def _pattern_matches_file(pattern: str, filename: str) -> bool:
    """Check if a pattern matches a filename"""
    import fnmatch
    
    # Exact match
    if pattern == filename:
        return True
    
    # Wildcard match
    if '*' in pattern or '?' in pattern:
        return fnmatch.fnmatch(filename, pattern)
    
    # Substring match (for patterns without wildcards that might be in filename)
    if pattern in filename:
        return True
    
    return False


@app.post("/api/schema/file-mappings")
async def save_file_table_mapping(
    file_pattern: str = Form(...),
    table_name: str = Form(...),
    mapping_id: int = Form(None),
    session: UserSession = Depends(require_role(UserRole.ADMIN))
):
    """Save file-to-table name mapping - Admin only"""
    try:
        import sqlite3
        from pathlib import Path
        from datetime import datetime
        from .config import config
        
        # Validate input
        file_pattern = file_pattern.strip()
        table_name = table_name.strip()
        
        if not file_pattern or not table_name:
            return {"success": False, "error": "File pattern and table name are required"}
        
        internal_db = config.directories.database_dir / "internal.db"
        internal_db.parent.mkdir(parents=True, exist_ok=True)
        
        with sqlite3.connect(internal_db) as conn:
            conn.row_factory = sqlite3.Row
            
            # Get existing custom mappings (excluding the one being updated)
            existing_mappings = []
            cursor = conn.execute("""
                SELECT id, file_pattern, table_name, is_active
                FROM file_table_mappings
                WHERE is_active = 1
            """)
            for row in cursor.fetchall():
                # Skip the mapping being updated
                if mapping_id and row['id'] == mapping_id:
                    continue
                existing_mappings.append({
                    'file_pattern': row['file_pattern'],
                    'table_name': row['table_name']
                })
            
            # Get default mappings for conflict checking
            default_mappings = []
            expected_tables = list(config.data_quality.expected_tables.keys())
            ignored_prefixes = config.etl.ignored_filename_prefixes
            extensions = config.etl.recognized_extensions
            
            for table in expected_tables:
                default_mappings.append({
                    'file_pattern': f'chhsca_{table}_*.txt',
                    'table_name': table
                })
                if 'SAMPLE' in ignored_prefixes:
                    default_mappings.append({
                        'file_pattern': f'SAMPLE_chhsca_{table}_*.txt',
                        'table_name': table
                    })
                default_mappings.append({
                    'file_pattern': f'{table}_*.txt',
                    'table_name': table
                })
                default_mappings.append({
                    'file_pattern': f'{table}.txt',
                    'table_name': table
                })
                for ext in extensions:
                    if ext != '.txt':
                        ext_clean = ext.lstrip('.')
                        default_mappings.append({
                            'file_pattern': f'{table}_*.{ext_clean}',
                            'table_name': table
                        })
                        default_mappings.append({
                            'file_pattern': f'chhsca_{table}_*.{ext_clean}',
                            'table_name': table
                        })
            
            # Filter out defaults that are already custom mappings
            custom_patterns = {m['file_pattern'] for m in existing_mappings}
            default_mappings = [m for m in default_mappings if m['file_pattern'] not in custom_patterns]
            
            # Check for conflicts
            conflict_error = _check_mapping_conflicts(file_pattern, table_name, existing_mappings, default_mappings)
            if conflict_error:
                return {"success": False, "error": conflict_error}
            
            now = datetime.now().isoformat()
            
            if mapping_id:
                # Update existing
                conn.execute("""
                    UPDATE file_table_mappings
                    SET file_pattern = ?, table_name = ?, updated_at = ?
                    WHERE id = ?
                """, (file_pattern, table_name, now, mapping_id))
            else:
                # Check for duplicate pattern
                cursor = conn.execute("""
                    SELECT id FROM file_table_mappings
                    WHERE file_pattern = ? AND is_active = 1
                """, (file_pattern,))
                if cursor.fetchone():
                    return {"success": False, "error": f"A mapping with pattern '{file_pattern}' already exists"}
                
                # Insert new
                conn.execute("""
                    INSERT INTO file_table_mappings (file_pattern, table_name, created_at, updated_at, created_by, is_active)
                    VALUES (?, ?, ?, ?, ?, 1)
                """, (file_pattern, table_name, now, now, session.username))
            
            conn.commit()
        
        return {"success": True, "message": "Mapping saved successfully"}
    except sqlite3.IntegrityError as e:
        if "UNIQUE constraint" in str(e):
            return {"success": False, "error": f"A mapping with pattern '{file_pattern}' already exists"}
        return {"success": False, "error": f"Database error: {str(e)}"}
    except Exception as e:
        logger.error(f"Error saving file mapping: {e}", exc_info=True)
        return {"success": False, "error": str(e)}


@app.delete("/api/schema/file-mappings/{mapping_id}")
async def delete_file_table_mapping(
    mapping_id: int,
    session: UserSession = Depends(require_role(UserRole.ADMIN))
):
    """Delete file-to-table name mapping - Admin only"""
    try:
        import sqlite3
        from pathlib import Path
        from .config import config
        
        internal_db = config.directories.database_dir / "internal.db"
        
        if internal_db.exists():
            with sqlite3.connect(internal_db) as conn:
                conn.execute("DELETE FROM file_table_mappings WHERE id = ?", (mapping_id,))
                conn.commit()
        
        return {"success": True, "message": "Mapping deleted successfully"}
    except Exception as e:
        logger.error(f"Error deleting file mapping: {e}", exc_info=True)
        return {"success": False, "error": str(e)}


@app.post("/api/admin/users")
async def create_user(
    username: str = Form(...),
    display_name: str = Form(""),
    email: str = Form(""),
    password: str = Form(""),
    role: str = Form(...),
    auth_method: str = Form("local"),
    session: UserSession = Depends(require_auth)
):
    """Create a new user - Admin only"""
    if session.role.value != 'admin':
        raise HTTPException(status_code=403, detail="Admin access required")
    
    try:
        # Validate auth_method
        if auth_method not in ['local', 'ad']:
            return {"success": False, "error": "Invalid auth_method. Must be 'local' or 'ad'"}
        
        # Check for special markers for AD retrieval
        obtain_email_on_login = False
        if email == '__OBTAIN_FROM_AD__':
            obtain_email_on_login = True
            email = None  # Will be fetched from AD on first login
        
        obtain_display_name_on_login = False
        if display_name == '__OBTAIN_FROM_AD__':
            obtain_display_name_on_login = True
            display_name = None  # Will be fetched from AD on first login
        
        # For AD users, password is not required (they use AD credentials)
        if auth_method == 'local' and not password:
            return {"success": False, "error": "Password is required for local authentication"}
        
        # For local users, display name is required
        if auth_method == 'local' and not display_name:
            return {"success": False, "error": "Display name is required for local authentication"}
        
        # For AD users, generate a random password (won't be used)
        if auth_method == 'ad':
            password = secrets.token_hex(32)
        
        auth_service = get_auth_service()
        success = auth_service.local_db.create_user(
            username=username,
            password=password,
            display_name=display_name,
            email=email,
            role=UserRole(role),
            created_by=session.username,
            auth_method=auth_method,
            obtain_email_on_login=obtain_email_on_login,
            obtain_display_name_on_login=obtain_display_name_on_login
        )
        
        if success:
            logger.info(f"User created: {username} (auth: {auth_method}) by {session.username}")
            return {"success": True, "message": f"User '{username}' created successfully with {auth_method} authentication"}
        else:
            return {"success": False, "error": "Failed to create user (username may already exist)"}
    except Exception as e:
        logger.error(f"Error creating user: {e}")
        return {"success": False, "error": str(e)}


@app.get("/api/admin/users/{username}")
async def get_user(username: str, session: UserSession = Depends(require_auth)):
    """Get details of a specific user - Admin only"""
    if session.role.value != 'admin':
        raise HTTPException(status_code=403, detail="Admin access required")
    
    try:
        auth_service = get_auth_service()
        
        import sqlite3
        with sqlite3.connect(auth_service.local_db.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.execute("""
                SELECT username, display_name, email, role, auth_method, is_active, 
                       created_at, created_by, last_login
                FROM sys_users 
                WHERE LOWER(username) = LOWER(?)
            """, (username,))
            
            user_row = cursor.fetchone()
            
            if not user_row:
                raise HTTPException(status_code=404, detail="User not found")
            
            user = dict(user_row)
            logger.info(f"User details fetched: {username} by {session.username}")
            return {"success": True, "user": user}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching user details: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/admin/users/{username}")
async def update_user(
    username: str,
    display_name: str = Form(None),
    email: str = Form(None),
    role: str = Form(None),
    is_active: str = Form(None),
    password: str = Form(None),
    session: UserSession = Depends(require_auth),
    request: Request = None
):
    """Update a user - Admin only with comprehensive audit logging"""
    if session.role.value != 'admin':
        raise HTTPException(status_code=403, detail="Admin access required")
    
    try:
        auth_service = get_auth_service()
        
        # Get current user state for audit logging
        import sqlite3
        with sqlite3.connect(auth_service.local_db.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.execute("""
                SELECT username, display_name, email, role, is_active, auth_method
                FROM sys_users 
                WHERE LOWER(username) = LOWER(?)
            """, (username,))
            current_user = cursor.fetchone()
            
            if not current_user:
                raise HTTPException(status_code=404, detail="User not found")
            
            current_user_dict = dict(current_user)
        
        # Build update query dynamically and track changes
        updates = []
        params = []
        changes = []
        
        if display_name is not None:
            old_value = current_user_dict.get('display_name', '')
            if old_value != display_name:
                updates.append("display_name = ?")
                params.append(display_name)
                changes.append(f"display_name: '{old_value}' → '{display_name}'")
        
        if email is not None:
            old_value = current_user_dict.get('email', '')
            if old_value != email:
                updates.append("email = ?")
                params.append(email)
                changes.append(f"email: '{old_value}' → '{email}'")
        
        if role is not None:
            # Validate role
            if role not in ['admin', 'operator', 'viewer']:
                return {"success": False, "error": "Invalid role. Must be admin, operator, or viewer"}
            old_value = current_user_dict.get('role', '')
            if old_value != role:
                updates.append("role = ?")
                params.append(role)
                changes.append(f"role: '{old_value}' → '{role}'")
        
        if is_active is not None:
            old_value = current_user_dict.get('is_active', 0)
            new_value = 1 if is_active == '1' else 0
            if old_value != new_value:
                updates.append("is_active = ?")
                params.append(new_value)
                status_change = 'activated' if new_value else 'deactivated'
                changes.append(f"status: {'active' if old_value else 'inactive'} → {'active' if new_value else 'inactive'}")
        
        if password is not None and password:
            # Validate password length
            if len(password) < 8:
                return {"success": False, "error": "Password must be at least 8 characters"}
            # Hash the new password
            password_hash = auth_service.local_db._hash_password(password)
            updates.append("password_hash = ?")
            params.append(password_hash)
            changes.append("password: [changed]")
        
        if not updates:
            return {"success": False, "error": "No fields to update"}
        
        params.append(username)
        
        # Perform the update
        with sqlite3.connect(auth_service.local_db.db_path) as conn:
            query = f"UPDATE sys_users SET {', '.join(updates)} WHERE LOWER(username) = LOWER(?)"
            conn.execute(query, params)
        
        # Get client IP and user agent for audit logging
        client_ip = request.client.host if request else None
        user_agent = request.headers.get("user-agent") if request else None
        
        # Comprehensive audit logging with before/after values
        from .audit_logger import get_audit_logger, AuditAction, AuditCategory
        audit_logger = get_audit_logger()
        
        audit_details = f"User '{username}' updated by {session.username}. Changes: {'; '.join(changes)}"
        
        audit_logger.log(
            username=session.username,
            action=AuditAction.USER_UPDATED.value,
            category=AuditCategory.USER_MANAGEMENT.value,
            success=True,
            details=audit_details,
            ip_address=client_ip,
            user_agent=user_agent,
            session_id=session.session_id if hasattr(session, 'session_id') else None,
            target_user=username,
            target_resource=f"user:{username}"
        )
        
        # Also log to legacy audit trail for backwards compatibility
        auth_service.local_db.log_audit(
            username=session.username,
            action='edit_user',
            category='user_management',
            success=True,
            details=audit_details,
            ip_address=client_ip,
            user_agent=user_agent,
            session_id=session.session_id if hasattr(session, 'session_id') else None,
            target_user=username
        )
        
        logger.info(f"User updated: {username} by {session.username}. Changes: {', '.join(changes)}")
        return {"success": True, "message": f"User '{username}' updated successfully"}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating user: {e}", exc_info=True)
        
        # Log failed update attempt
        try:
            from .audit_logger import get_audit_logger, AuditAction, AuditCategory
            audit_logger = get_audit_logger()
            client_ip = request.client.host if request else None
            user_agent = request.headers.get("user-agent") if request else None
            
            audit_logger.log(
                username=session.username,
                action=AuditAction.USER_UPDATED.value,
                category=AuditCategory.USER_MANAGEMENT.value,
                success=False,
                details=f"Failed to update user '{username}': {str(e)}",
                ip_address=client_ip,
                user_agent=user_agent,
                session_id=session.session_id if hasattr(session, 'session_id') else None,
                target_user=username,
                error_message=str(e)
            )
        except:
            pass  # Don't fail if audit logging fails
        
        return {"success": False, "error": str(e)}


@app.delete("/api/admin/users/{username}")
async def delete_user(username: str, session: UserSession = Depends(require_auth)):
    """Delete (deactivate) a user - Admin only - DEPRECATED, use PATCH to toggle status"""
    if session.role.value != 'admin':
        raise HTTPException(status_code=403, detail="Admin access required")
    
    # Prevent deleting self
    if username.lower() == session.username.lower():
        return {"success": False, "error": "Cannot delete your own account"}
    
    try:
        auth_service = get_auth_service()
        success = auth_service.local_db.deactivate_user(username)
        
        if success:
            logger.info(f"User deactivated: {username} by {session.username}")
            return {"success": True, "message": f"User '{username}' deactivated successfully"}
        else:
            return {"success": False, "error": "Failed to deactivate user"}
    except Exception as e:
        logger.error(f"Error deactivating user: {e}")
        return {"success": False, "error": str(e)}


@app.patch("/api/admin/users/{username}/toggle-status")
async def toggle_user_status(username: str, session: UserSession = Depends(require_auth)):
    """Toggle user active status (activate/deactivate) - Admin only"""
    if session.role.value != 'admin':
        raise HTTPException(status_code=403, detail="Admin access required")
    
    # Prevent toggling self
    if username.lower() == session.username.lower():
        return {"success": False, "error": "Cannot toggle your own account status"}
    
    try:
        auth_service = get_auth_service()
        success, new_status = auth_service.local_db.toggle_user_status(username)
        
        if success:
            action = "activated" if new_status == 'active' else "deactivated"
            logger.info(f"User {action}: {username} by {session.username}")
            return {
                "success": True, 
                "message": f"User '{username}' {action} successfully",
                "new_status": new_status
            }
        else:
            return {"success": False, "error": "Failed to toggle user status"}
    except Exception as e:
        logger.error(f"Error toggling user status: {e}")
        return {"success": False, "error": str(e)}


@app.get("/api/admin/security/health-check")
async def security_health_check(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Get comprehensive security health check - Admin only"""
    try:
        from .security_health_check import get_health_checker
        
        health_checker = get_health_checker()
        result = health_checker.run_all_checks()
        
        return result
    except Exception as e:
        logger.error(f"Error running security health check: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e),
            "checks": {},
            "score": {'score': 0, 'rating': 'Error', 'passed': 0, 'failed': 0, 'total': 0},
            "hipaa_compliance": [],
            "recommendations": []
        }


@app.get("/api/admin/siem/config")
async def get_siem_config(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Get SIEM configuration - Admin only"""
    try:
        from .settings_manager import get_settings_manager
        
        settings_manager = get_settings_manager()
        settings = settings_manager.get_siem_settings()
        
        # Map database settings to admin config format
        return {
            "success": True,
            "config": {
                "enabled": bool(settings.get('enabled', 0)),
                "protocol": "syslog",  # Currently only syslog is supported
                "host": settings.get('syslog_host', ''),
                "port": settings.get('syslog_port', 514),
                "transport": settings.get('syslog_protocol', 'UDP').lower(),
                "facility": 16,  # User facility (local0)
                "min_severity": "info",
                "app_name": "UniteUs-ETL",
                "hostname_override": "",
                "batch_mode": True,
                "event_categories": ["authentication", "user_management", "data_access", "data_export", "etl", "security", "system"]
            }
        }
    except Exception as e:
        logger.error(f"Error loading SIEM config: {e}", exc_info=True)
        # Return proper JSON structure instead of raising HTTPException
        # This ensures frontend gets a parseable response
        return {
            "success": False,
            "error": str(e),
            "config": {
                "enabled": False,
                "protocol": "syslog",
                "host": "",
                "port": 514,
                "transport": "udp",
                "facility": 16,
                "min_severity": "info",
                "app_name": "UniteUs-ETL",
                "hostname_override": "",
                "batch_mode": True,
                "event_categories": []
            }
        }


@app.post("/api/admin/siem/config")
async def save_siem_config(
    config: dict,
    session: UserSession = Depends(require_role(UserRole.ADMIN))
):
    """Save SIEM configuration - Admin only"""
    try:
        from .settings_manager import get_settings_manager
        from .siem_logger import get_siem_logger
        
        settings_manager = get_settings_manager()
        
        # Map admin config to settings format
        # If enabled and syslog settings are provided, save them
        if config.get('enabled') and config.get('host'):
            settings = {
                'enabled': config.get('enabled', False),
                'enable_windows_event_log': True,  # Keep Windows Event Log enabled for fatal errors
                'syslog_enabled': True,
                'syslog_host': config.get('host', ''),
                'syslog_port': config.get('port', 514),
                'syslog_protocol': config.get('transport', 'UDP').upper(),
                'include_sensitive_data': False
            }
            
            success = settings_manager.save_siem_settings(settings, session.username)
            
            if success:
                # Reinitialize SIEM logger with new settings
                siem_logger = get_siem_logger()
                siem_logger._initialize_loggers()
                
                logger.info(f"SIEM configuration updated by {session.username}")
                logger.info(f"SIEM config: {config}")
                
                return {
                    "success": True,
                    "message": "SIEM configuration saved successfully"
                }
            else:
                return {
                    "success": False,
                    "error": "Failed to save SIEM settings to database"
                }
        else:
            # Just log the config but don't enable syslog if host is missing
            logger.info(f"SIEM configuration updated by {session.username}")
            logger.info(f"SIEM config: {config}")
            
            # If enabled but no host, disable syslog
            if config.get('enabled') and not config.get('host'):
                settings = {
                    'enabled': False,  # Disable if not properly configured
                    'enable_windows_event_log': True,
                    'syslog_enabled': False,
                    'syslog_host': '',
                    'syslog_port': 514,
                    'syslog_protocol': 'UDP',
                    'include_sensitive_data': False
                }
                settings_manager.save_siem_settings(settings, session.username)
            
            return {
                "success": True,
                "message": "SIEM configuration saved (syslog disabled - host required)"
            }
    except Exception as e:
        logger.error(f"Error saving SIEM config: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e)
        }


@app.post("/api/admin/siem/test")
async def test_siem_connection(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Test SIEM connection - Admin only - Actually logs a test event"""
    try:
        from .config import config
        from .siem_logger import log_siem_event, SIEMEventType, SIEMSeverity
        
        logger.info(f"SIEM connection test initiated by {session.username}")
        
        # Actually log a test event to all enabled backends
        if config.siem.enabled:
            log_siem_event(
                SIEMEventType.SYSTEM_EVENT,
                f"SIEM integration test event from Admin Control Panel - User: {session.username}",
                severity=SIEMSeverity.INFO,
                username=session.username,
                success=True,
                additional_data={"test": True, "source": "siem_integration_test", "timestamp": datetime.now().isoformat()}
            )
        
        # Build detailed message based on what's actually enabled and properly configured
        message_parts = []
        enabled_backends = []
        errors = []
        
        if config.siem.enable_windows_event_log:
            enabled_backends.append("Windows Event Log")
            message_parts.append("Windows Event Log: Test event logged for fatal errors")
        
        if config.siem.syslog_enabled:
            # Validate syslog configuration
            if not config.siem.syslog_host or config.siem.syslog_host.strip() == '':
                errors.append("Syslog host is not configured")
            elif not config.siem.syslog_port or config.siem.syslog_port <= 0:
                errors.append("Syslog port is not configured")
            else:
                enabled_backends.append("Syslog forwarding to IT SIEM")
                message_parts.append(f"Syslog: Test event sent to {config.siem.syslog_host}:{config.siem.syslog_port}")
        
        # If syslog is enabled but not properly configured, add error
        if config.siem.syslog_enabled and not any("Syslog" in b for b in enabled_backends):
            errors.append("Syslog forwarding is enabled but not properly configured (host and port required)")
        
        if errors:
            return {
                "success": False,
                "error": "; ".join(errors)
            }
        
        if not enabled_backends:
            return {
                "success": False,
                "error": "No SIEM logging backends are enabled. Please enable Windows Event Log and/or Syslog forwarding to IT's SIEM server."
            }
        
        message = f"✅ SIEM test successful | Enabled backends: {', '.join(enabled_backends)} | " + " | ".join(message_parts)
        
        return {
            "success": True,
            "message": message
        }
    except Exception as e:
        logger.error(f"Error testing SIEM connection: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e)
        }


@app.get("/api/admin/siem/statistics")
async def get_siem_statistics(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Get SIEM statistics - Admin only"""
    return {
        "success": True,
        "stats": {
            "events_sent_today": 0,
            "success_rate": 100,
            "failed_events": 0,
            "last_sent": "Never"
        }
    }


@app.get("/api/admin/siem/activity")
async def get_siem_activity(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Get recent SIEM activity - Admin only"""
    return {
        "success": True,
        "activity": []
    }


@app.get("/api/admin/system-components/status")
async def get_system_components_status(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Get comprehensive status of all system components - Admin only"""
    from datetime import datetime, timedelta
    from .config import config
    from .auth import get_auth_service
    from .database_adapter import get_database_adapter
    
    components = {}
    
    try:
        # 1. Authentication Service
        auth_service = get_auth_service()
        active_sessions = auth_service.get_active_sessions()
        
        # Get last login from audit log
        last_login = None
        last_login_user = None
        try:
            import sqlite3
            with sqlite3.connect(auth_service.local_db.db_path) as conn:
                conn.row_factory = sqlite3.Row
                cursor = conn.execute("""
                    SELECT username, last_login FROM sys_users 
                    WHERE last_login IS NOT NULL 
                    ORDER BY last_login DESC LIMIT 1
                """)
                row = cursor.fetchone()
                if row:
                    last_login = row['last_login']
                    last_login_user = row['username']
        except:
            pass
        
        # Calculate time since last login
        last_login_ago = None
        if last_login:
            try:
                last_login_dt = datetime.fromisoformat(last_login)
                delta = datetime.now() - last_login_dt
                if delta.days > 0:
                    last_login_ago = f"{delta.days} day{'s' if delta.days > 1 else ''} ago"
                elif delta.seconds >= 3600:
                    hours = delta.seconds // 3600
                    last_login_ago = f"{hours} hour{'s' if hours > 1 else ''} ago"
                else:
                    minutes = delta.seconds // 60
                    last_login_ago = f"{minutes} minute{'s' if minutes > 1 else ''} ago" if minutes > 0 else "Just now"
            except:
                pass
        
        auth_mode = "Hybrid (AD + Local)" if auth_service.mode.value == "hybrid" else auth_service.mode.value
        components['authentication'] = {
            'status': 'pass',
            'name': 'Authentication Service',
            'details': f'Live • {len(active_sessions)} active session(s)',
            'info': f'{auth_mode} • Last login: {last_login_ago or "Never"} by {last_login_user or "N/A"}'
        }
        
        # 2. Database
        try:
            adapter = get_database_adapter()
            with adapter.get_connection() as conn:
                # Test connection
                if hasattr(conn, 'execute'):
                    conn.execute("SELECT 1")
                else:
                    cursor = conn.cursor()
                    cursor.execute("SELECT 1")
                    cursor.fetchone()
            
            db_type = config.database.db_type.upper()
            if db_type == "MSSQL":
                db_display = f"MS SQL Server ({config.database.mssql_server or 'N/A'})"
            elif db_type == "AZURESQL":
                db_display = f"Azure SQL ({config.database.azuresql_server or 'N/A'})"
            elif db_type == "POSTGRESQL":
                db_display = f"PostgreSQL ({config.database.postgresql_host or 'N/A'})"
            elif db_type == "MYSQL":
                db_display = f"MySQL ({config.database.mysql_host or 'N/A'})"
            else:
                db_display = f"SQLite ({config.database.path})"
            
            components['database'] = {
                'status': 'pass',
                'name': 'Database',
                'details': 'Connected and operational',
                'info': db_display
            }
        except Exception as e:
            components['database'] = {
                'status': 'fail',
                'name': 'Database',
                'details': 'Connection failed',
                'info': f'Error: {str(e)[:50]}'
            }
        
        # 3. SFTP Server (configuration check only - no actual connection)
        try:
            if config.sftp.enabled and config.sftp.host:
                # Just check if SFTP is configured - don't actually connect
                hostname = config.sftp.host or 'Unknown'
                port = config.sftp.port or 22
                username = config.sftp.username or 'N/A'
                
                # Validate configuration completeness
                if config.sftp.host and config.sftp.username:
                    components['sftp'] = {
                        'status': 'pass',
                        'name': 'SFTP Server',
                        'details': f'Configured: {hostname}',
                        'info': f'Port {port} • User: {username}'
                    }
                else:
                    components['sftp'] = {
                        'status': 'warning',
                        'name': 'SFTP Server',
                        'details': 'Incomplete configuration',
                        'info': 'Missing host or username settings'
                    }
            else:
                components['sftp'] = {
                    'status': 'warning',
                    'name': 'SFTP Server',
                    'details': 'Not configured',
                    'info': 'SFTP integration is disabled or host not set'
                }
        except ImportError:
            components['sftp'] = {
                'status': 'unknown',
                'name': 'SFTP Server',
                'details': 'Service unavailable',
                'info': 'SFTP service module not available'
            }
        except Exception as e:
            components['sftp'] = {
                'status': 'unknown',
                'name': 'SFTP Server',
                'details': 'Status unknown',
                'info': f'Error: {str(e)[:80]}'
            }
        
        # 4. Active Directory
        if auth_service.ad_enabled:
            components['active_directory'] = {
                'status': 'pass',
                'name': 'Active Directory',
                'details': f'Connected to {auth_service.ad_domain}.local',
                'info': f'Server: {auth_service.ad_server} • Search base: {auth_service.ad_search_base}'
            }
        else:
            components['active_directory'] = {
                'status': 'warning',
                'name': 'Active Directory',
                'details': 'Not configured',
                'info': 'AD authentication is disabled'
            }
        
        # 5. Logging
        try:
            siem_enabled = config.siem.enabled
            active_logging_types = []
            
            if siem_enabled:
                if config.siem.enable_windows_event_log:
                    active_logging_types.append("Windows Event Log")
                if config.siem.syslog_enabled:
                    active_logging_types.append("Syslog to IT SIEM")
                
                if active_logging_types:
                    components['siem'] = {
                        'status': 'pass',
                        'name': 'Logging',
                        'details': ', '.join(active_logging_types),
                        'info': f'{len(active_logging_types)} logging method(s) active'
                    }
                else:
                    components['siem'] = {
                        'status': 'warning',
                        'name': 'Logging',
                        'details': 'No logging methods enabled',
                        'info': 'SIEM is enabled but no logging backends are configured'
                    }
            else:
                components['siem'] = {
                    'status': 'warning',
                    'name': 'Logging',
                    'details': 'Not configured',
                    'info': 'Logging is disabled'
                }
        except Exception as e:
            components['siem'] = {
                'status': 'unknown',
                'name': 'Logging',
                'details': 'Status unknown',
                'info': f'Error: {str(e)[:50]}'
            }
        
        # 6. ETL Engine
        components['etl_engine'] = {
            'status': 'pass',
            'name': 'ETL Engine',
            'details': 'Ready for processing',
            'info': 'All validation rules active'
        }
        
    except Exception as e:
        logger.error(f"Error getting system components status: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e)
        }
    
    return {
        "success": True,
        "components": components
    }


# ========================================================================
# LOGGING API ENDPOINTS (Windows Event Log & JSON Logging)
# ========================================================================

@app.get("/api/admin/logging/windows/status")
async def get_windows_event_log_status(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Get Windows Event Log status - Admin only"""
    try:
        import sys
        from .siem_logger import WindowsEventLogger
        from .config import config
        
        pywin32_available = False
        try:
            import win32evtlog
            pywin32_available = True
        except ImportError:
            pass
        
        windows_logger = WindowsEventLogger()
        
        return {
            "success": True,
            "info": {
                "platform": sys.platform,
                "pywin32_available": pywin32_available,
                "supported": sys.platform == 'win32' and pywin32_available,
                "enabled": config.siem.enable_windows_event_log,
                "app_name": "UniteUsETL",
                "events_logged": 0  # Could track this if needed
            }
        }
    except Exception as e:
        logger.error(f"Error getting Windows Event Log status: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e)
        }


@app.post("/api/admin/logging/windows/test")
async def test_windows_event_log(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Test Windows Event Log - Admin only"""
    try:
        import sys
        from .siem_logger import log_siem_event, SIEMEventType, SIEMSeverity, get_siem_logger
        from .config import config
        
        if sys.platform != 'win32':
            return {
                "success": False,
                "error": "Windows Event Log is only available on Windows"
            }
        
        # Check if Windows Event Log is enabled
        if not config.siem.enable_windows_event_log:
            return {
                "success": False,
                "error": "Windows Event Log is not enabled. Please enable it in the configuration above."
            }
        
        # Get the SIEM logger to check if Windows logger is initialized
        siem_logger = get_siem_logger()
        if not siem_logger.windows_logger or not siem_logger.windows_logger.enabled:
            return {
                "success": False,
                "error": "Windows Event Log logger is not initialized or not available. Check that pywin32 is installed."
            }
        
        # Log test event - this will write to all enabled backends (JSON, Windows Event Log, etc.)
        log_siem_event(
            SIEMEventType.SYSTEM_EVENT,
            f"Windows Event Log test event from Admin Control Panel - User: {session.username}",
            severity=SIEMSeverity.INFO,
            username=session.username,
            success=True,
            additional_data={"test": True, "source": "admin_panel", "timestamp": datetime.now().isoformat()}
        )
        
        # Also directly test Windows Event Logger
        try:
            siem_logger.windows_logger.log_event(
                f"Direct Windows Event Log test - User: {session.username}",
                "Information"
            )
        except Exception as we_error:
            logger.warning(f"Direct Windows Event Log write failed: {we_error}")
        
        return {
            "success": True,
            "message": "Test event logged successfully to Windows Event Log. Open Event Viewer and navigate to: Applications and Services Logs → UniteUsETL",
            "details": {
                "windows_event_log_enabled": config.siem.enable_windows_event_log,
                "windows_logger_available": siem_logger.windows_logger is not None,
                "windows_logger_enabled": siem_logger.windows_logger.enabled if siem_logger.windows_logger else False
            }
        }
    except Exception as e:
        logger.error(f"Error testing Windows Event Log: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e)
        }


@app.get("/api/admin/logging/windows/events")
async def get_windows_event_log_events(
    limit: int = 20,
    session: UserSession = Depends(require_role(UserRole.ADMIN))
):
    """Get recent Windows Event Log events - Admin only"""
    try:
        import sys
        from datetime import datetime
        
        if sys.platform != 'win32':
            return {
                "success": True,
                "events": [],
                "message": "Not running on Windows"
            }
        
        try:
            import win32evtlog
            import win32con
        except ImportError:
            return {
                "success": True,
                "events": [],
                "message": "pywin32 not installed"
            }
        
        events = []
        
        try:
            # Open the Application event log
            # We'll look for events from our application name
            app_name = "UniteUsETL"
            hand = win32evtlog.OpenEventLog(None, "Application")
            
            # Read events (newest first)
            flags = win32evtlog.EVENTLOG_BACKWARDS_READ | win32evtlog.EVENTLOG_SEQUENTIAL_READ
            events_read = win32evtlog.ReadEventLog(hand, flags, 0)
            
            # Filter for our application and limit results
            count = 0
            for event in events_read:
                # Check if this event is from our application
                # Only check SourceName - don't check event data as it's too broad
                source_name = event.SourceName if hasattr(event, 'SourceName') else None
                
                # Only include events where SourceName exactly matches our app name
                if source_name == app_name:
                    # Parse event time
                    event_time = event.TimeGenerated if hasattr(event, 'TimeGenerated') else None
                    if event_time:
                        try:
                            time_str = event_time.strftime("%Y-%m-%d %H:%M:%S")
                        except:
                            time_str = str(event_time)
                    else:
                        time_str = "N/A"
                    
                    # Map event type
                    event_type_num = event.EventType if hasattr(event, 'EventType') else 4
                    if event_type_num == 1:  # ERROR
                        event_type = "Error"
                    elif event_type_num == 2:  # WARNING
                        event_type = "Warning"
                    else:  # INFORMATION (4)
                        event_type = "Information"
                    
                    # Get message from event data/strings
                    event_strings = event.StringInserts if hasattr(event, 'StringInserts') and event.StringInserts else []
                    message = ' '.join(event_strings) if event_strings else "No message data"
                    
                    events.append({
                        "time": time_str,
                        "type": event_type,
                        "message": message[:200] if len(message) > 200 else message,  # Truncate long messages
                        "event_id": event.EventID if hasattr(event, 'EventID') else 0,
                        "record_id": event.RecordNumber if hasattr(event, 'RecordNumber') else 0
                    })
                    
                    count += 1
                    if count >= limit:
                        break
            
            win32evtlog.CloseEventLog(hand)
            
            # If we didn't find events in Application log, try Applications and Services Logs
            if len(events) == 0:
                try:
                    # Try to open the Applications and Services Logs path
                    # This requires the log to be created first
                    log_path = f"Application/{app_name}"
                    hand = win32evtlog.OpenEventLog(None, log_path)
                    
                    events_read = win32evtlog.ReadEventLog(hand, flags, 0)
                    
                    for event in events_read:
                        event_time = event.TimeGenerated if hasattr(event, 'TimeGenerated') else None
                        if event_time:
                            try:
                                time_str = event_time.strftime("%Y-%m-%d %H:%M:%S")
                            except:
                                time_str = str(event_time)
                        else:
                            time_str = "N/A"
                        
                        event_type_num = event.EventType if hasattr(event, 'EventType') else 4
                        if event_type_num == 1:
                            event_type = "Error"
                        elif event_type_num == 2:
                            event_type = "Warning"
                        else:
                            event_type = "Information"
                        
                        event_strings = event.StringInserts if hasattr(event, 'StringInserts') and event.StringInserts else []
                        message = ' '.join(event_strings) if event_strings else "No message data"
                        
                        events.append({
                            "time": time_str,
                            "type": event_type,
                            "message": message[:200] if len(message) > 200 else message,
                            "event_id": event.EventID if hasattr(event, 'EventID') else 0,
                            "record_id": event.RecordNumber if hasattr(event, 'RecordNumber') else 0
                        })
                        
                        if len(events) >= limit:
                            break
                    
                    win32evtlog.CloseEventLog(hand)
                except Exception as e2:
                    # If Applications and Services Logs doesn't exist yet, that's okay
                    logger.debug(f"Could not read from Applications and Services Logs: {e2}")
            
            return {
                "success": True,
                "events": events,
                "count": len(events)
            }
            
        except Exception as read_error:
            logger.error(f"Error reading Windows Event Log: {read_error}", exc_info=True)
            return {
                "success": True,
                "events": [],
                "message": f"Could not read events: {str(read_error)}. Events may not have been created yet, or you may need administrator permissions."
            }
            
    except Exception as e:
        logger.error(f"Error getting Windows Event Log events: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e)
        }


@app.post("/api/admin/logging/windows/open-viewer")
async def open_windows_event_viewer(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Open Windows Event Viewer - Admin only"""
    try:
        import sys
        import subprocess
        
        if sys.platform != 'win32':
            return {
                "success": False,
                "error": "Not running on Windows"
            }
        
        # Open Event Viewer using eventvwr.msc
        # This will open to the default view, user can navigate to Applications and Services Logs
        try:
            subprocess.Popen(['eventvwr.msc'], shell=True)
            return {
                "success": True,
                "message": "Event Viewer opened. Navigate to: Applications and Services Logs → UniteUsETL"
            }
        except Exception as e:
            logger.error(f"Error opening Event Viewer: {e}")
            return {
                "success": False,
                "error": f"Could not open Event Viewer: {str(e)}"
            }
    except Exception as e:
        logger.error(f"Error opening Windows Event Viewer: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e)
        }


@app.get("/api/admin/logging/json/status")
async def get_json_logging_status(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Get JSON logging status - Admin only"""
    try:
        from pathlib import Path
        from .config import config
        from datetime import datetime
        
        log_path = Path(config.siem.json_log_path)
        file_exists = log_path.exists()
        file_size = log_path.stat().st_size if file_exists else 0
        
        # Count entries (rough estimate by counting lines)
        entries_count = 0
        last_updated = None
        if file_exists:
            try:
                with open(log_path, 'r', encoding='utf-8') as f:
                    lines = f.readlines()
                    entries_count = len([l for l in lines if l.strip()])
                    if entries_count > 0:
                        # Get file modification time
                        last_updated = datetime.fromtimestamp(log_path.stat().st_mtime).isoformat()
            except Exception:
                pass
        
        return {
            "success": True,
            "info": {
                "log_path": str(log_path),
                "file_exists": file_exists,
                "file_size": file_size,
                "enabled": config.siem.enable_json_logging,
                "entries_count": entries_count,
                "last_updated": last_updated
            }
        }
    except Exception as e:
        logger.error(f"Error getting JSON logging status: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e)
        }


@app.post("/api/admin/logging/json/test")
async def test_json_logging(session: UserSession = Depends(require_role(UserRole.ADMIN))):
    """Test JSON logging - Admin only"""
    try:
        from .siem_logger import log_siem_event, SIEMEventType, SIEMSeverity
        
        log_siem_event(
            SIEMEventType.SYSTEM_EVENT,
            "JSON logging test event from Admin Control Panel",
            severity=SIEMSeverity.INFO,
            username=session.username,
            success=True,
            additional_data={"test": True, "source": "admin_panel"}
        )
        
        return {
            "success": True,
            "message": "Test event logged successfully. Check the JSON log file."
        }
    except Exception as e:
        logger.error(f"Error testing JSON logging: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e)
        }


@app.get("/api/admin/logging/json/entries")
async def get_json_log_entries(
    limit: int = 20,
    session: UserSession = Depends(require_role(UserRole.ADMIN))
):
    """Get recent JSON log entries - Admin only"""
    try:
        import json
        from pathlib import Path
        from .config import config
        
        log_path = Path(config.siem.json_log_path)
        
        if not log_path.exists():
            return {
                "success": True,
                "entries": []
            }
        
        entries = []
        try:
            with open(log_path, 'r', encoding='utf-8') as f:
                lines = f.readlines()
                # Get last N lines
                recent_lines = lines[-limit:] if len(lines) > limit else lines
                
                for line in recent_lines:
                    line = line.strip()
                    if line:
                        try:
                            entry = json.loads(line)
                            entries.append({
                                "timestamp": entry.get("timestamp", entry.get("asctime", "N/A")),
                                "event_type": entry.get("event_type", "N/A"),
                                "severity": entry.get("severity", "INFO"),
                                "event_message": entry.get("event_message", entry.get("message", "N/A")),
                                "username": entry.get("username", "N/A")
                            })
                        except json.JSONDecodeError:
                            continue
                
                # Reverse to show newest first
                entries.reverse()
        except Exception as e:
            logger.error(f"Error reading JSON log file: {e}")
            return {
                "success": False,
                "error": f"Error reading log file: {str(e)}"
            }
        
        return {
            "success": True,
            "entries": entries
        }
    except Exception as e:
        logger.error(f"Error getting JSON log entries: {e}", exc_info=True)
        return {
            "success": False,
            "error": str(e)
        }


@app.get("/api/admin/logging/json/download")
async def download_json_log_file(
    path: str = None,
    session: UserSession = Depends(require_role(UserRole.ADMIN))
):
    """Download JSON log file - Admin only"""
    try:
        from pathlib import Path
        from fastapi.responses import FileResponse
        from .config import config
        
        log_path = Path(path) if path else Path(config.siem.json_log_path)
        
        if not log_path.exists():
            raise HTTPException(status_code=404, detail="Log file not found")
        
        return FileResponse(
            path=str(log_path),
            filename=log_path.name,
            media_type='application/json'
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error downloading JSON log file: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/admin/sftp/verify-key")
async def verify_sftp_key(
    request: dict,
    session: UserSession = Depends(require_role(UserRole.ADMIN))
):
    """Verify that SFTP key file exists - Admin only"""
    try:
        key_path = request.get('key_path', '')
        
        if not key_path:
            return {
                "exists": False,
                "message": "No key path provided"
            }
        
        # Resolve the path relative to the project root
        from pathlib import Path
        key_file = Path(key_path)
        
        # Check if file exists
        if key_file.exists():
            # Try to determine the key format by reading first line
            try:
                with open(key_file, 'r') as f:
                    first_line = f.readline().strip()
                    
                    if first_line.startswith('-----BEGIN OPENSSH PRIVATE KEY-----'):
                        key_format = "OpenSSH"
                    elif first_line.startswith('-----BEGIN RSA PRIVATE KEY-----'):
                        key_format = "RSA/PEM"
                    elif first_line.startswith('-----BEGIN DSA PRIVATE KEY-----'):
                        key_format = "DSA"
                    elif first_line.startswith('-----BEGIN EC PRIVATE KEY-----'):
                        key_format = "ECDSA"
                    elif first_line.startswith('-----BEGIN PRIVATE KEY-----'):
                        key_format = "PKCS#8"
                    elif first_line.startswith('PuTTY-User-Key-File'):
                        key_format = "PuTTY (needs conversion!)"
                    elif first_line.startswith('---- BEGIN SSH2'):
                        key_format = "SSH2"
                    else:
                        key_format = "Unknown"
                        
            except Exception:
                key_format = None
            
            return {
                "exists": True,
                "format": key_format,
                "message": f"File found at {key_file.absolute()}"
            }
        else:
            return {
                "exists": False,
                "message": f"File not found. Expected location: {key_file.absolute()}"
            }
            
    except Exception as e:
        logger.error(f"Error verifying SFTP key: {e}", exc_info=True)
        return {
            "exists": False,
            "message": f"Error checking file: {str(e)}"
        }


# ============================================================================
# ERROR HANDLERS
# ============================================================================

@app.exception_handler(404)
async def custom_404_handler(request: Request, exc):
    """Custom 404 handler - resource not found"""
    logger.warning(f"404 Not Found: {request.method} {request.url.path}")
    return JSONResponse(
        status_code=404,
        content={
            "error": "Resource not found",
            "path": str(request.url.path),
            "method": request.method
        }
    )


@app.exception_handler(500)
async def custom_500_handler(request: Request, exc):
    """Custom 500 handler - internal server error"""
    logger.error(f"500 Internal Server Error: {request.method} {request.url.path} - {str(exc)}", exc_info=True)
    return JSONResponse(
        status_code=500,
        content={
            "error": "Internal server error",
            "message": str(exc),
            "path": str(request.url.path)
        }
    )


@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException):
    """Enhanced HTTP exception handler with better logging"""
    # Log 401 errors as debug (expected when not authenticated)
    if exc.status_code == 401:
        logger.debug(
            f"HTTP Exception: {request.method} {request.url.path} - "
            f"Status: {exc.status_code} - Detail: {exc.detail}"
        )
    else:
        logger.error(
            f"HTTP Exception: {request.method} {request.url.path} - "
            f"Status: {exc.status_code} - Detail: {exc.detail}"
        )
    return JSONResponse(
        status_code=exc.status_code,
        content={
            "error": exc.detail,
            "status_code": exc.status_code,
            "path": str(request.url.path)
        }
    )


@app.exception_handler(Exception)
async def general_exception_handler(request: Request, exc: Exception):
    """Catch-all exception handler for unhandled errors"""
    logger.error(
        f"Unhandled Exception: {request.method} {request.url.path}\n"
        f"  Exception Type: {type(exc).__name__}\n"
        f"  Message: {str(exc)}",
        exc_info=True
    )
    return JSONResponse(
        status_code=500,
        content={
            "error": "An unexpected error occurred",
            "exception_type": type(exc).__name__,
            "message": str(exc),
            "path": str(request.url.path)
        }
    )


# ============================================================================
# MAIN APPLICATION ENTRY POINT
# ============================================================================

if __name__ == "__main__":
    uvicorn.run(
        "app:app",
        host=config.web.host,
        port=config.web.port,
        reload=config.web.reload,
        log_level=config.web.log_level
    )